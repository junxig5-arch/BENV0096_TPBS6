from pathlib import Path
import math

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(r"E:\UCL Final Essay")
P8_ROOT = PROJECT_ROOT / "p8_uncertainty_framework"
TABLE_DIR = P8_ROOT / "tables"
FIG_DIR = P8_ROOT / "figures"
DOC_DIR = P8_ROOT / "documents"
DOC_DIR.mkdir(parents=True, exist_ok=True)

OUT_DOCX = DOC_DIR / "P8_Uncertainty_Framework_Methods_Results_Discussion_Dissertation_Ready.docx"
DOWNLOADS_DOCX = Path(r"C:\Users\888\Downloads\P8_Uncertainty_Framework_Methods_Results_Discussion_Dissertation_Ready.docx")


NAVY = RGBColor(31, 77, 120)
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(35, 35, 35)
MUTED = RGBColor(96, 96, 96)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
BORDER = "D9E2F3"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=DARK, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(str(text))
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color="D9E2F3", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_inches):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)
                tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(int(width * 1440)))
                tc_w.set(qn("w:type"), "dxa")


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Caption"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(9)
    run = p.add_run(text)
    run.font.color.rgb = MUTED
    run.font.size = Pt(9)
    return p


def add_source_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED
    return p


def fmt_value(value, digits=1):
    if pd.isna(value):
        return ""
    if isinstance(value, str):
        return value
    try:
        v = float(value)
    except Exception:
        return str(value)
    if abs(v) >= 100:
        return f"{v:.0f}"
    return f"{v:.{digits}f}"


def add_df_table(doc, df, columns, headers, widths, font_size=8.4, header_fill=LIGHT_GRAY, max_rows=None):
    display_df = df[columns].copy()
    if max_rows:
        display_df = display_df.head(max_rows)
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True, color=DARK, size=font_size)
        set_cell_shading(hdr[i], header_fill)
    for _, row in display_df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            set_cell_text(cells[i], fmt_value(row[col]), bold=False, size=font_size)
    set_table_borders(table, BORDER, "4")
    set_cell_margins(table)
    set_table_width(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


def add_callout(doc, title, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [6.5])
    set_table_borders(table, "C9D6EA", "6")
    set_cell_margins(table, top=130, start=160, bottom=130, end=160)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = NAVY
    r.font.size = Pt(10.5)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    r2 = p2.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = DARK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_figure(doc, img_path, caption, width=6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(width))
    add_caption(doc, caption)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.color.rgb = MUTED
    caption.paragraph_format.space_after = Pt(8)


def set_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "P8 Uncertainty Framework | Dissertation Draft Material"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.runs[0].font.size = Pt(8.5)
    hp.runs[0].font.color.rgb = MUTED

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = "Junxi Gao | Draft analytical chapter material"
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.runs[0].font.size = Pt(8.5)
    fp.runs[0].font.color.rgb = MUTED


def add_title_block(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("P8 Uncertainty Framework, Scenario Design and Synthesis")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Dissertation-ready Methods, Results and Discussion Draft")
    r.font.size = Pt(12.5)
    r.font.color.rgb = MUTED

    rows = [
        ("Purpose", "Translate P8-1 to P8-3 into dissertation-ready text and evidence."),
        ("Scope", "Historical delivered-rate benchmark; sectoral uncertainty and linkages; 2x2 scenario matrix; near-miss and offsets stress tests."),
        ("Status", "Draft material for integration into Methods, Results and Discussion."),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_width(table, [1.15, 5.35])
    set_cell_margins(table, top=80, start=120, bottom=80, end=120)
    set_table_borders(table, "D9E2F3", "4")
    for i, (label, value) in enumerate(rows):
        set_cell_text(table.cell(i, 0), label, bold=True, color=NAVY, size=9.5)
        set_cell_shading(table.cell(i, 0), LIGHT_BLUE)
        set_cell_text(table.cell(i, 1), value, size=9.5)
    doc.add_paragraph()


def load_tables():
    return {
        "hist_rates": pd.read_csv(TABLE_DIR / "p8_1_historical_rate_windows.csv"),
        "future_rates": pd.read_csv(TABLE_DIR / "p8_1_future_pathway_rate_comparison.csv"),
        "sector_metrics": pd.read_csv(TABLE_DIR / "p8_2_sector_rate_anchor_metrics.csv"),
        "sector_band": pd.read_csv(TABLE_DIR / "p8_2_sector_uncertainty_band_2050.csv"),
        "driver_matrix": pd.read_csv(TABLE_DIR / "p8_2_sector_linkage_driver_matrix.csv"),
        "scenario_matrix": pd.read_csv(TABLE_DIR / "p8_3_2x2_scenario_matrix.csv"),
        "mini_population": pd.read_csv(TABLE_DIR / "p8_3_mini_scenario_population.csv"),
        "near_miss": pd.read_csv(TABLE_DIR / "p8_3_near_miss_summary.csv"),
        "credit_stress": pd.read_csv(TABLE_DIR / "p8_3_negative_emissions_credit_stress_test.csv"),
    }


def main():
    data = load_tables()
    hist = data["hist_rates"]
    future = data["future_rates"]
    sector_metrics = data["sector_metrics"]
    sector_band = data["sector_band"]
    scenario_matrix = data["scenario_matrix"]
    mini = data["mini_population"]
    near = data["near_miss"]

    hist_1990 = float(hist.loc[hist["window"].eq("1990-2025"), "avg_annual_reduction_MtCO2e_per_year"].iloc[0])
    hist_2018 = float(hist.loc[hist["window"].eq("2018-2025"), "avg_annual_reduction_MtCO2e_per_year"].iloc[0])
    desnz_late = float(future[(future["pathway"].str.contains("including IAS")) & (future["period"].eq("2035-2050"))]["avg_annual_reduction_MtCO2e_per_year"].iloc[0])
    ccc_late = float(future[(future["pathway"].str.contains("CCC Seventh")) & (future["period"].eq("2035-2050"))]["avg_annual_reduction_MtCO2e_per_year"].iloc[0])
    u1 = float(mini.loc[mini["quadrant_id"].eq("U1"), "net_2050_after_default_offsets_MtCO2e"].median())
    u4 = float(mini.loc[mini["quadrant_id"].eq("U4"), "net_2050_after_default_offsets_MtCO2e"].median())

    doc = Document()
    configure_styles(doc)
    set_header_footer(doc)
    add_title_block(doc)

    add_callout(
        doc,
        "Central finding",
        (
            f"The uncertainty analysis strengthens the dissertation argument rather than replacing it. "
            f"Historical delivery rates show that the CCC7 late-period reduction requirement ({ccc_late:.1f} MtCO2e/year) "
            f"is demanding but historically recognisable, while the DESNZ baseline after 2035 ({desnz_late:.1f} MtCO2e/year) "
            f"represents a delivery slowdown. The 2x2 mini-scenarios then show that supportive external conditions are useful, "
            f"but weak domestic delivery still leaves large residual gaps."
        ),
    )

    doc.add_heading("1. Role of P8 in the Dissertation", level=1)
    doc.add_paragraph(
        "P8 extends the earlier baseline-benchmark and sectoral results into an explicit uncertainty framework. "
        "P5 established the aggregate gap between the DESNZ current-policy baseline and the CCC target-consistent benchmark. "
        "P6 showed that the gap is not sectorally uniform. P8 asks how far that gap is sensitive to historical delivery rates, "
        "sectoral linkages, domestic policy delivery and external technology or fuel conditions."
    )
    doc.add_paragraph(
        "The analytical purpose is not to forecast probabilities. Instead, P8 provides a transparent structure for judging whether the "
        "UK moves closer to a target-consistent transition, remains a near miss, or stays materially off track under delayed delivery."
    )

    doc.add_heading("2. Draft Methods Subsection: Uncertainty and Scenario Design", level=1)
    doc.add_heading("2.1 Historical delivered-rate benchmark", level=2)
    doc.add_paragraph(
        "The first uncertainty anchor compares future pathway reduction rates with historical UK territorial greenhouse gas reduction rates. "
        f"The historical data indicate average reductions of {hist_1990:.1f} MtCO2e/year over 1990-2025 and {hist_2018:.1f} MtCO2e/year over 2018-2025. "
        "These rates are used as delivered-rate benchmarks, not as forecasts, because the composition of past reductions differs from the composition of future abatement."
    )
    add_df_table(
        doc,
        hist,
        ["window", "start_emissions_MtCO2e", "end_emissions_MtCO2e", "avg_annual_reduction_MtCO2e_per_year", "implied_zero_year_if_linear_continues"],
        ["Window", "Start MtCO2e", "End MtCO2e", "Avg reduction / year", "Linear zero year"],
        [1.15, 1.25, 1.25, 1.45, 1.4],
        font_size=8.3,
    )
    add_source_note(doc, "Source: DESNZ final/provisional UK greenhouse gas emissions statistics; author calculations.")

    doc.add_heading("2.2 Sectoral uncertainty and linkage framework", level=2)
    doc.add_paragraph(
        "The second uncertainty anchor uses the cleaned DESNZ TES-sector time series and the P6 sectoral bridge. "
        "For each sector, the analysis compares projected DESNZ 2023-2050 reduction rates with historical delivered rates and classifies whether the projection is comparable, slower or a reversal. "
        "A qualitative linkage matrix then identifies sectors that can move together through clean power and electrification and sectors that remain relatively independent."
    )
    add_bullet(doc, "Coupled electrification cluster: electricity supply, domestic transport, and buildings/product uses.")
    add_bullet(doc, "Partly linked hard-to-abate cluster: industry and fuel supply, especially where hydrogen, CCUS and low-carbon fuels matter.")
    add_bullet(doc, "Residual and near-miss risk sectors: agriculture, IAS, LULUCF and waste, where clean electricity alone is insufficient.")

    doc.add_heading("2.3 2x2 scenario matrix and mini-scenarios", level=2)
    doc.add_paragraph(
        "The third uncertainty anchor combines two axes: UK domestic policy delivery and external conditions. "
        "Domestic delivery captures implementation strength across buildings, transport, power, industry and residual sectors. "
        "External conditions capture technology learning, clean-energy costs, fossil fuel exposure, supply-chain constraints and international action. "
        "The four quadrants are represented by twelve mini-scenarios, three per quadrant."
    )
    add_df_table(
        doc,
        scenario_matrix,
        ["quadrant_id", "scenario_name", "uk_policy_delivery", "external_conditions", "default_removals_credit_available_MtCO2e"],
        ["ID", "Scenario", "UK delivery", "External conditions", "Default offsets MtCO2e"],
        [0.55, 2.35, 1.1, 1.55, 0.95],
        font_size=8.2,
    )
    add_source_note(doc, "Source: Author scenario construction using P8-1, P8-2 and NESO FES 2025 evidence.")

    doc.add_heading("2.4 Negative emissions and credits stress test", level=2)
    doc.add_paragraph(
        "Because residual emissions remain in several sectors, the scenario analysis tests 0, 10 and 50 MtCO2e offsets in 2050. "
        "This follows the supervisor's suggestion to ask where negative emissions or credit purchases would enter the bridge between current-policy and target-consistent pathways. "
        "The stress test is interpreted cautiously: offsets can reduce residual gaps, but they do not replace mitigation in high-residual sectors."
    )

    doc.add_page_break()
    doc.add_heading("3. Results Synthesis", level=1)
    doc.add_heading("3.1 DESNZ slowdown relative to historical delivery", level=2)
    doc.add_paragraph(
        f"The P8-1 result clarifies the aggregate gap. The DESNZ current-policy pathway including IAS falls by only {desnz_late:.1f} MtCO2e/year over 2035-2050, "
        f"whereas the CCC Seventh Carbon Budget Balanced Pathway falls by about {ccc_late:.1f} MtCO2e/year over the same period. "
        "The comparison with historical delivered rates suggests that the DESNZ baseline is not only high in 2050; it slows markedly after 2035."
    )
    add_figure(
        doc,
        FIG_DIR / "p8_1_historical_vs_future_pathways.jpg",
        "Figure P8-1. Historical UK territorial GHG emissions compared with DESNZ and CCC future pathways.",
        width=6.4,
    )
    add_figure(
        doc,
        FIG_DIR / "p8_1_rate_benchmark_comparison.jpg",
        "Figure P8-2. Historical average annual reductions compared with 2035-2050 pathway rates.",
        width=6.4,
    )

    doc.add_heading("3.2 Sectoral uncertainty and linkages", level=2)
    doc.add_paragraph(
        "P8-2 adds sector-level interpretation. The largest DESNZ 2050 residual sectors are buildings and product uses, domestic transport, industry and electricity supply. "
        "However, the uncertainty interpretation depends on the type of sector. Buildings and LULUCF are projected to reverse after historical decline; domestic transport, industry, electricity, IAS, fuel supply and waste show slower projected delivery than historical rate anchors."
    )
    compact_sector = sector_metrics[[
        "tes_sector",
        "projected_rate_2023_2050_MtCO2e_per_year",
        "best_positive_historical_rate_MtCO2e_per_year",
        "rate_anchor_interpretation",
        "rank_2050_residual",
    ]].copy()
    compact_sector["tes_sector"] = compact_sector["tes_sector"].replace({"Buildings and product uses": "Buildings/product uses", "Domestic Transport": "Domestic transport"})
    add_df_table(
        doc,
        compact_sector,
        ["rank_2050_residual", "tes_sector", "projected_rate_2023_2050_MtCO2e_per_year", "best_positive_historical_rate_MtCO2e_per_year", "rate_anchor_interpretation"],
        ["Rank", "Sector", "Projected rate", "Best historical rate", "Interpretation"],
        [0.45, 1.45, 1.05, 1.15, 2.4],
        font_size=7.7,
    )
    add_source_note(doc, "Source: DESNZ EEP sector tables and author calculations.")
    add_figure(
        doc,
        FIG_DIR / "p8_2_sector_uncertainty_band_2050.jpg",
        "Figure P8-3. Stylised sectoral 2050 uncertainty band with CCC7 mapped benchmark.",
        width=6.35,
    )
    add_figure(
        doc,
        FIG_DIR / "p8_2_sector_linkage_driver_matrix.jpg",
        "Figure P8-4. Sector linkage driver matrix: clean power, electrification, hydrogen/CCUS, demand, land/non-CO2, international fuel/technology and removals/credits.",
        width=6.35,
    )

    doc.add_heading("3.3 2x2 outcomes and near-miss interpretation", level=2)
    doc.add_paragraph(
        f"The aligned-transition quadrant has the lowest median net 2050 residual after default offsets, at approximately {u1:.1f} MtCO2e. "
        f"The delayed-transition quadrant has the highest median net residual, at approximately {u4:.1f} MtCO2e. "
        "The spread between these quadrants is the key P8 finding: supportive external conditions reduce transition difficulty, but the domestic delivery axis remains decisive."
    )
    add_figure(
        doc,
        FIG_DIR / "p8_3_quadrant_outcomes.jpg",
        "Figure P8-5. 2x2 scenario matrix: median 2050 residual after default offsets.",
        width=5.8,
    )
    add_figure(
        doc,
        FIG_DIR / "p8_3_mini_scenario_residual_gaps.jpg",
        "Figure P8-6. Mini-scenario residual gaps after default offsets.",
        width=6.35,
    )

    near50 = near[near["offsets_available_MtCO2e"].eq(50)][["quadrant_id", "median_net_2050_MtCO2e", "median_delay_years", "near_miss_share"]].copy()
    add_df_table(
        doc,
        near50,
        ["quadrant_id", "median_net_2050_MtCO2e", "median_delay_years", "near_miss_share"],
        ["Quadrant", "Median net residual", "Median delay years", "Near-miss share"],
        [1.0, 1.8, 1.6, 1.5],
        font_size=8.3,
    )
    add_source_note(doc, "Source: P8-3 mini-scenario stress test. Near-miss is defined as on-time or residual <= 70 MtCO2e after offsets.")
    add_figure(
        doc,
        FIG_DIR / "p8_3_negative_emissions_credit_sensitivity.jpg",
        "Figure P8-7. Sensitivity of quadrant-level median residuals to 0, 10 and 50 MtCO2e offsets.",
        width=6.35,
    )

    doc.add_page_break()
    doc.add_heading("4. Draft Discussion Outline", level=1)
    doc.add_heading("4.1 Interpretation of the main finding", level=2)
    doc.add_paragraph(
        "The central interpretation is that the UK net-zero gap is a delivery problem under uncertainty, not only an accounting endpoint problem. "
        "DESNZ's current-policy pathway leaves a large residual gap because post-2035 reductions slow considerably. CCC7 and several NESO pathways show that target-consistent or near-target outcomes are possible in stylised terms, but only when sectoral delivery remains strong."
    )
    doc.add_heading("4.2 Domestic delivery versus external conditions", level=2)
    add_numbered(doc, "Supportive global clean-technology conditions lower the difficulty of electrification, clean power and low-carbon fuels, but do not automatically close the residual gap.")
    add_numbered(doc, "Weak UK policy delivery remains a binding risk even under favourable external conditions, especially in buildings, transport and industry.")
    add_numbered(doc, "Constrained external conditions make high domestic delivery harder, but they do not make domestic implementation irrelevant.")

    doc.add_heading("4.3 Near-miss outcomes", level=2)
    doc.add_paragraph(
        "Near-miss outcomes are analytically useful because they avoid treating every missed pathway as equally unsuccessful. "
        "A pathway with residual emissions of around 50-70 MtCO2e in 2050 has a different policy meaning from one with residuals above 250-300 MtCO2e. "
        "The discussion should therefore distinguish on-time delivery, near-miss residual gaps, moderate misses and delayed-transition risk."
    )
    doc.add_heading("4.4 Negative emissions and credits", level=2)
    doc.add_paragraph(
        "Negative emissions or credit purchases can reduce the residual gap, but P8 suggests they should be treated as a complement to mitigation, not as a substitute for it. "
        "The 50 MtCO2e stress test moves the aligned-transition quadrant into near-miss territory for most mini-scenarios, but it does not rescue weak-delivery or delayed-transition cases."
    )
    doc.add_heading("4.5 Limitations to state explicitly", level=2)
    add_bullet(doc, "The mini-scenarios are stylised and should not be interpreted as probabilities or forecasts.")
    add_bullet(doc, "Historical reduction rates are useful delivery benchmarks, but past reductions were partly enabled by sectoral changes that may not be repeatable.")
    add_bullet(doc, "DESNZ and CCC sector mappings are not always one-to-one, especially for buildings/product uses, IAS, LULUCF and engineered removals.")
    add_bullet(doc, "Offsets and negative emissions are treated as stress-test quantities rather than as confirmed future availability.")

    doc.add_heading("5. Source and Citation Note", level=1)
    doc.add_paragraph(
        "The P8 analysis uses official UK government statistics, DESNZ projections, CCC benchmark pathway data, NESO FES scenario evidence and international energy-cost evidence. "
        "The following references should be retained in the dissertation reference list."
    )
    refs = [
        "Department for Energy Security and Net Zero (DESNZ) (2026) Final UK greenhouse gas emissions statistics: 1990 to 2024. London: DESNZ.",
        "Department for Energy Security and Net Zero (DESNZ) (2026) Provisional UK greenhouse gas emissions statistics: 2025. London: DESNZ.",
        "Department for Energy Security and Net Zero (DESNZ) (2024) Energy and emissions projections: 2024 to 2050. London: DESNZ.",
        "Climate Change Committee (CCC) (2025) The Seventh Carbon Budget. London: Climate Change Committee.",
        "National Energy System Operator (NESO) (2025) Future Energy Scenarios 2025. Warwick: NESO.",
        "International Energy Agency (IEA) (2024) World Energy Investment 2024. Paris: IEA.",
        "International Energy Agency (IEA) (2024) Renewables 2024. Paris: IEA.",
        "International Renewable Energy Agency (IRENA) (2025) Renewable Power Generation Costs in 2024. Abu Dhabi: IRENA.",
    ]
    for ref in refs:
        add_bullet(doc, ref)

    doc.add_heading("6. Integration Checklist", level=1)
    add_bullet(doc, "Methods: insert Section 2 as the uncertainty and scenario-design subsection.")
    add_bullet(doc, "Results: insert Sections 3.1-3.3 after the P5/P6 baseline and sectoral results.")
    add_bullet(doc, "Discussion: use Section 4 to structure domestic delivery, external conditions, near-miss outcomes and offsets/removals.")
    add_bullet(doc, "Appendix: place full CSV outputs and detailed sector mini-scenario tables if the main dissertation becomes too dense.")

    doc.save(OUT_DOCX)
    doc.save(DOWNLOADS_DOCX)
    print(f"Saved: {OUT_DOCX}")
    print(f"Copied: {DOWNLOADS_DOCX}")


if __name__ == "__main__":
    main()
