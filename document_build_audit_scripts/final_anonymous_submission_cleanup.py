# -*- coding: utf-8 -*-
from __future__ import annotations

from copy import deepcopy
from io import BytesIO
from pathlib import Path
import math
import re
import shutil
import tempfile
import zipfile

from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import numpy as np
import pandas as pd
from PIL import Image, ImageDraw, ImageFont


SRC = Path(r"C:\Users\888\Desktop\Final Draft\Dissertation_Final_Draft_Junxi_Gao.docx")
OUT = Path(r"C:\Users\888\Desktop\Final Draft\TBPS6_Dissertation_Final_Draft_Cleaned.docx")

PROJECT_ROOT = Path(r"E:\UCL Final Essay")
TABLE_DIR = PROJECT_ROOT / "p8_uncertainty_framework" / "tables"
FIG_DIR = PROJECT_ROOT / "p8_uncertainty_framework" / "figures" / "submission_cleaned"
FIG_DIR.mkdir(parents=True, exist_ok=True)

OFFICIAL_INDEX_OLD = PROJECT_ROOT / "Reference" / "Official_data_sources" / "Official_Data_Sources_Index_Harvard_Junxi_Gao.docx"
OFFICIAL_INDEX_NEW = PROJECT_ROOT / "Reference" / "Official_data_sources" / "Official_Data_Sources_Index_Harvard_TBPS6.docx"

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
    "dc": "http://purl.org/dc/elements/1.1/",
    "dcterms": "http://purl.org/dc/terms/",
}


def get_font(size: int, bold: bool = False):
    candidates = []
    if bold:
        candidates += [
            Path(r"C:\Windows\Fonts\arialbd.ttf"),
            Path(r"C:\Windows\Fonts\calibrib.ttf"),
            Path(r"C:\Windows\Fonts\segoeuib.ttf"),
        ]
    candidates += [
        Path(r"C:\Windows\Fonts\arial.ttf"),
        Path(r"C:\Windows\Fonts\calibri.ttf"),
        Path(r"C:\Windows\Fonts\segoeui.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def hex_to_rgb(value: str):
    value = value.lstrip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def lighten(rgb, amount=0.85):
    return tuple(int(c + (255 - c) * amount) for c in rgb)


def draw_text(draw, xy, text, font, fill=(31, 31, 31), anchor=None):
    draw.text(xy, str(text), font=font, fill=fill, anchor=anchor)


def save_jpg(img: Image.Image, path: Path):
    img.convert("RGB").save(path, quality=96, optimize=True)


FONT_TITLE = get_font(42, True)
FONT_SUBTITLE = get_font(24)
FONT_LABEL = get_font(22)
FONT_SMALL = get_font(18)
FONT_TINY = get_font(15)
FONT_AXIS = get_font(19)
FONT_BOLD = get_font(22, True)

BG = (255, 255, 255)
INK = (31, 31, 31)
MUTED = (92, 92, 92)
GRID = (218, 218, 218)

TYPOLOGY_COLORS = {
    "Coupled electrification block": "#2E7D32",
    "Industrial hydrogen-CCUS block": "#1F77B4",
    "Land and non-CO2 residual block": "#8E6C00",
    "International fuel and technology block": "#B23A48",
    "Lower-coupling residual block": "#666666",
}

SECTOR_SHORT = {
    "Buildings and product uses": "Buildings/products",
    "Domestic Transport": "Domestic transport",
    "Industry": "Industry",
    "Electricity supply": "Electricity",
    "Agriculture": "Agriculture",
    "IAS": "IAS",
    "Fuel supply": "Fuel supply",
    "Waste": "Waste",
    "LULUCF": "LULUCF",
}


def generate_replacement_figures():
    metrics = pd.read_csv(TABLE_DIR / "p8_4_sector_linkage_deepening_metrics.csv")
    similarity = pd.read_csv(TABLE_DIR / "p8_4_sector_dependency_similarity_matrix.csv", index_col=0)

    fig15 = FIG_DIR / "figure_4_15_linkage_score_vs_residual_submission_cleaned.jpg"
    fig16 = FIG_DIR / "figure_4_16_dependency_similarity_submission_cleaned.jpg"
    fig17 = FIG_DIR / "figure_4_17_linkage_weighted_ranking_submission_cleaned.jpg"

    # Figure 4.15: wider canvas and manual label offsets to prevent overlap.
    img = Image.new("RGB", (2200, 1400), BG)
    draw = ImageDraw.Draw(img)
    draw_text(draw, (85, 55), "Linkage score and 2050 residual emissions", FONT_TITLE)
    draw_text(
        draw,
        (85, 112),
        "Higher scores indicate stronger dependence on shared transition drivers; bubble size reflects 2050 residual share.",
        FONT_SUBTITLE,
        MUTED,
    )
    plot_df = metrics.sort_values("rank_2050_residual").copy()
    left, right, top, bottom = 255, 160, 175, 215
    plot_w, plot_h = 2200 - left - right, 1400 - top - bottom
    x_min, x_max = 0, 10
    y_min, y_max = 0, 110
    for tick in range(0, 11, 2):
        x = left + (tick - x_min) / (x_max - x_min) * plot_w
        draw.line((x, top, x, top + plot_h), fill=GRID, width=1)
        draw_text(draw, (x, top + plot_h + 24), str(tick), FONT_AXIS, MUTED, anchor="ma")
    for tick in range(0, 111, 20):
        y = top + plot_h - (tick - y_min) / (y_max - y_min) * plot_h
        draw.line((left, y, left + plot_w, y), fill=GRID, width=1)
        draw_text(draw, (left - 18, y), str(tick), FONT_AXIS, MUTED, anchor="rm")
    draw.line((left, top, left, top + plot_h), fill=INK, width=2)
    draw.line((left, top + plot_h, left + plot_w, top + plot_h), fill=INK, width=2)
    draw_text(draw, (left + plot_w / 2, 1298), "System linkage score", FONT_BOLD, anchor="ma")
    draw_text(draw, (left + 10, top - 34), "2050 residual (MtCO2e)", FONT_BOLD, MUTED)
    offsets = {
        "Buildings and product uses": (30, -58),
        "Domestic Transport": (34, 8),
        "Industry": (34, -50),
        "Electricity supply": (34, 24),
        "Agriculture": (34, -40),
        "IAS": (34, 22),
        "Fuel supply": (34, 24),
        "Waste": (34, -36),
        "LULUCF": (34, 16),
    }
    for _, row in plot_df.iterrows():
        x = left + (row["system_linkage_score"] - x_min) / (x_max - x_min) * plot_w
        y = top + plot_h - (row["desnz_central_2050_MtCO2e"] - y_min) / (y_max - y_min) * plot_h
        radius = 20 + max(3, row["share_of_2050_inc_IAS_total_pct"]) * 1.25
        color = hex_to_rgb(TYPOLOGY_COLORS[row["linkage_typology"]])
        draw.ellipse((x - radius, y - radius, x + radius, y + radius), fill=lighten(color, 0.35), outline=color, width=4)
        dx, dy = offsets[row["tes_sector"]]
        draw_text(draw, (x + dx, y + dy), SECTOR_SHORT[row["tes_sector"]], FONT_LABEL)
    legend_x, legend_y = 85, 1190
    for i, (label, color_hex) in enumerate(TYPOLOGY_COLORS.items()):
        y = legend_y + i * 30
        color = hex_to_rgb(color_hex)
        draw.rectangle((legend_x, y - 10, legend_x + 22, y + 10), fill=color)
        draw_text(draw, (legend_x + 36, y), label, FONT_SMALL, MUTED, anchor="lm")
    draw_text(
        draw,
        (2130, 1352),
        "Source: Author calculations from P8 sector residual and linkage-driver tables. Diagnostic scores, not causal estimates.",
        FONT_TINY,
        MUTED,
        anchor="ra",
    )
    save_jpg(img, fig15)

    # Figure 4.16: shorter, wider heatmap; wording avoids causal or observed co-movement claim.
    heat_df = similarity.copy()
    sectors = list(heat_df.index)
    short_labels = [SECTOR_SHORT.get(s, s) for s in sectors]
    n = len(sectors)
    img = Image.new("RGB", (2300, 1580), BG)
    draw = ImageDraw.Draw(img)
    draw_text(draw, (85, 55), "Dependency-profile similarity between residual sectors", FONT_TITLE)
    draw_text(
        draw,
        (85, 112),
        "Cosine similarity based on seven transition-driver scores. Values near 1 indicate similar dependency profiles.",
        FONT_SUBTITLE,
        MUTED,
    )
    left, top, cell = 400, 295, 120
    low, high = np.array([247, 251, 255]), np.array([8, 48, 107])
    for i, label in enumerate(short_labels):
        y = top + i * cell + cell / 2
        draw_text(draw, (left - 22, y), label, FONT_AXIS, INK, anchor="rm")
    for j, label in enumerate(short_labels):
        x = left + j * cell + cell / 2
        lines = re.sub("Domestic transport", "Domestic\ntransport", label).replace("Buildings/products", "Buildings/\nproducts").split("\n")
        for k, line in enumerate(lines):
            draw_text(draw, (x, top - 88 + k * 23), line, FONT_TINY, INK, anchor="ma")
    for i in range(n):
        for j in range(n):
            value = float(heat_df.iloc[i, j])
            rgb = tuple(np.round(low + (high - low) * value).astype(int))
            x0, y0 = left + j * cell, top + i * cell
            draw.rectangle((x0, y0, x0 + cell, y0 + cell), fill=rgb, outline=(245, 245, 245), width=2)
            txt_color = (255, 255, 255) if value >= 0.62 else INK
            draw_text(draw, (x0 + cell / 2, y0 + cell / 2), f"{value:.2f}", FONT_AXIS, txt_color, anchor="mm")
    draw.rectangle((left, top, left + n * cell, top + n * cell), outline=INK, width=2)
    draw_text(
        draw,
        (85, 1518),
        "Interpretation: shared dependency structure only; this is not evidence of observed causal co-movement.",
        FONT_SMALL,
        MUTED,
    )
    save_jpg(img, fig16)

    # Figure 4.17: wider right margin, shorter annotations and unclipped right-side labels.
    bar_df = metrics.sort_values("linkage_weighted_residual_index", ascending=True).reset_index(drop=True)
    img = Image.new("RGB", (2300, 1400), BG)
    draw = ImageDraw.Draw(img)
    draw_text(draw, (85, 55), "Linkage-weighted residual risk by sector", FONT_TITLE)
    draw_text(
        draw,
        (85, 112),
        "Index = DESNZ central 2050 residual emissions multiplied by the sector's system linkage score.",
        FONT_SUBTITLE,
        MUTED,
    )
    left, right, top, bottom = 380, 430, 175, 170
    plot_w, plot_h = 2300 - left - right, 1400 - top - bottom
    max_index = float(bar_df["linkage_weighted_residual_index"].max())
    row_h = plot_h / len(bar_df)
    for tick in np.linspace(0, max_index, 6):
        x = left + tick / max_index * plot_w
        draw.line((x, top, x, top + plot_h), fill=GRID, width=1)
        draw_text(draw, (x, top + plot_h + 22), f"{tick:.0f}", FONT_AXIS, MUTED, anchor="ma")
    for i, row in bar_df.iterrows():
        y = top + i * row_h + row_h * 0.5
        bar_len = row["linkage_weighted_residual_index"] / max_index * plot_w
        color = hex_to_rgb(TYPOLOGY_COLORS[row["linkage_typology"]])
        draw_text(draw, (left - 24, y), SECTOR_SHORT[row["tes_sector"]], FONT_AXIS, INK, anchor="rm")
        draw.rounded_rectangle(
            (left, y - row_h * 0.27, left + bar_len, y + row_h * 0.27),
            radius=8,
            fill=lighten(color, 0.25),
            outline=color,
            width=2,
        )
        annotation = (
            f"index {row['linkage_weighted_residual_index']:.0f}; "
            f"residual {row['desnz_central_2050_MtCO2e']:.1f}; "
            f"score {row['system_linkage_score']:.0f}"
        )
        draw_text(draw, (left + bar_len + 18, y), annotation, FONT_SMALL, MUTED, anchor="lm")
    draw.line((left, top + plot_h, left + plot_w, top + plot_h), fill=INK, width=2)
    draw_text(draw, (left + plot_w / 2, 1302), "Linkage-weighted residual index", FONT_BOLD, anchor="ma")
    draw_text(
        draw,
        (2230, 1352),
        "Source: Author calculations from P8 sector residual and linkage-driver tables.",
        FONT_TINY,
        MUTED,
        anchor="ra",
    )
    save_jpg(img, fig17)

    return {
        "media/image16.jpg": fig15,
        "media/image17.jpg": fig16,
        "media/image18.jpg": fig17,
    }


def replace_paragraph_text(paragraph, new_text):
    for child in list(paragraph._p):
        paragraph._p.remove(child)
    run = paragraph.add_run(new_text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def edit_docx_content():
    if OUT.exists():
        OUT.unlink()
    shutil.copy2(SRC, OUT)
    if OFFICIAL_INDEX_OLD.exists() and not OFFICIAL_INDEX_NEW.exists():
        shutil.copy2(OFFICIAL_INDEX_OLD, OFFICIAL_INDEX_NEW)

    doc = Document(str(OUT))
    for par in doc.paragraphs:
        if par.text.startswith("Third, the 2x2 framework cannot represent every uncertainty"):
            replace_paragraph_text(
                par,
                "The 2x2 framework cannot represent every uncertainty or assign probabilities (Trutnevyte et al., 2016; Lo Piano et al., 2021). "
                "The linkage extension should also be read as shared-dependency evidence. Testing causal spillovers or policy effects would require "
                "policy-level identification or a linked energy-system model, which is outside the scope of this pathway-diagnostic dissertation.",
            )
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    if "Junxi_Gao" in par.text:
                        replace_paragraph_text(par, par.text.replace("Junxi_Gao", "TBPS6"))
                    if "Junxi Gao" in par.text:
                        replace_paragraph_text(par, par.text.replace("Junxi Gao", "TBPS6"))

    props = doc.core_properties
    props.author = "TBPS6"
    props.last_modified_by = "TBPS6"
    props.title = "TBPS6 MSc ESDA Dissertation"
    props.subject = ""
    props.keywords = ""
    props.comments = ""
    props.category = ""
    props.content_status = ""
    doc.save(str(OUT))


def set_keep_with_next_for_figures():
    doc = Document(str(OUT))
    figure_numbers = {"Figure 4.15.", "Figure 4.16.", "Figure 4.17."}
    for idx, par in enumerate(doc.paragraphs):
        if any(par.text.strip().startswith(f) for f in figure_numbers):
            par.paragraph_format.keep_with_next = True
            if idx > 0:
                doc.paragraphs[idx - 1].paragraph_format.keep_with_next = True
    doc.save(str(OUT))


def scrub_core_and_replace_media(replacements):
    tmp = Path(tempfile.mkdtemp(prefix="tbps6_docx_"))
    with zipfile.ZipFile(OUT, "r") as zin:
        zin.extractall(tmp)

    # Replace embedded figure media.
    for target, new_path in replacements.items():
        dest = tmp / "word" / target
        dest.write_bytes(new_path.read_bytes())

    # Update displayed extents for the three replaced figures.
    doc_xml = tmp / "word" / "document.xml"
    tree = etree.parse(str(doc_xml))
    root = tree.getroot()
    rels = etree.parse(str(tmp / "word" / "_rels" / "document.xml.rels")).getroot()
    rid_by_target = {}
    for rel in rels:
        rid_by_target[rel.get("Id")] = rel.get("Target")

    width_emu = int(6.15 * 914400)
    media_aspect = {
        "media/image16.jpg": Image.open(replacements["media/image16.jpg"]).size,
        "media/image17.jpg": Image.open(replacements["media/image17.jpg"]).size,
        "media/image18.jpg": Image.open(replacements["media/image18.jpg"]).size,
    }
    for blip in root.xpath(".//a:blip[@r:embed]", namespaces=NS):
        rid = blip.get(f"{{{NS['r']}}}embed")
        target = rid_by_target.get(rid)
        if target in media_aspect:
            px_w, px_h = media_aspect[target]
            height_emu = int(width_emu * px_h / px_w)
            inline = blip.xpath("./ancestor::wp:inline[1]", namespaces=NS)
            if inline:
                extent = inline[0].find("./wp:extent", namespaces=NS)
                if extent is not None:
                    extent.set("cx", str(width_emu))
                    extent.set("cy", str(height_emu))
            ext = blip.xpath("./ancestor::a:graphic[1]//a:xfrm/a:ext", namespaces=NS)
            for node in ext:
                node.set("cx", str(width_emu))
                node.set("cy", str(height_emu))

    tree.write(str(doc_xml), encoding="UTF-8", xml_declaration=True, standalone=True)

    # Scrub core metadata directly in OOXML to ensure no personal name remains.
    core_xml = tmp / "docProps" / "core.xml"
    core_tree = etree.parse(str(core_xml))
    core_root = core_tree.getroot()
    metadata_values = {
        "dc:title": "TBPS6 MSc ESDA Dissertation",
        "dc:subject": "",
        "dc:creator": "TBPS6",
        "cp:keywords": "",
        "dc:description": "",
        "cp:lastModifiedBy": "TBPS6",
        "cp:category": "",
        "cp:contentStatus": "",
    }
    for xpath, value in metadata_values.items():
        nodes = core_root.xpath(f"./{xpath}", namespaces=NS)
        for node in nodes:
            node.text = value
    core_tree.write(str(core_xml), encoding="UTF-8", xml_declaration=True, standalone=True)

    with zipfile.ZipFile(OUT, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for file in tmp.rglob("*"):
            if file.is_file():
                zout.write(file, file.relative_to(tmp).as_posix())


def validate_images(replacements):
    rows = []
    for target, path in replacements.items():
        img = Image.open(path).convert("RGB")
        arr = np.asarray(img)
        mask = np.any(arr < 248, axis=2)
        ys, xs = np.where(mask)
        bbox = (int(xs.min()), int(ys.min()), int(xs.max()), int(ys.max())) if len(xs) else None
        margins = None
        if bbox:
            margins = {
                "left": bbox[0],
                "top": bbox[1],
                "right": img.width - bbox[2] - 1,
                "bottom": img.height - bbox[3] - 1,
            }
        rows.append(
            {
                "target": target,
                "path": str(path),
                "pixels": f"{img.width}x{img.height}",
                "non_white_bbox": str(bbox),
                "margins_px": str(margins),
                "edge_touching": bool(margins and min(margins.values()) < 10),
            }
        )
    return rows


def main():
    edit_docx_content()
    replacements = generate_replacement_figures()
    checks = validate_images(replacements)
    scrub_core_and_replace_media(replacements)
    set_keep_with_next_for_figures()
    print(f"output={OUT}")
    print(f"official_index_copy={OFFICIAL_INDEX_NEW} exists={OFFICIAL_INDEX_NEW.exists()}")
    for row in checks:
        print(row)


if __name__ == "__main__":
    main()
