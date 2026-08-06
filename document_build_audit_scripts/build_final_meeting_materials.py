from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd


ROOT = Path(r"E:\UCL Final Essay")
OUT_DIR = Path(r"C:\Users\888\Desktop\Final Draft\TBPS6_Meeting_Materials")
OUT_DIR.mkdir(parents=True, exist_ok=True)

THESIS = Path(r"C:\Users\888\Desktop\Final Draft\TBPS6_Dissertation_Final_Draft_Cleaned.docx")
EVIDENCE_PACK = OUT_DIR / "TBPS6_Dissertation_Meeting_Evidence_Pack_for_Neil.docx"
ORAL_SCRIPT = OUT_DIR / "TBPS6_One_Hour_Meeting_Oral_Script_Cued_Bilingual.docx"
GUIDE = OUT_DIR / "TBPS6_Dissertation_Detailed_Understanding_Guide_Bilingual.docx"

FIGURES = {
    "baseline": ROOT / r"meeting_preparation\figures\docx_safe_images\p5_cleaned_desnz_vs_ccc_annual_pathways_docx_safe.jpg",
    "gap": ROOT / r"meeting_preparation\figures\docx_safe_images\p5_cleaned_annual_gap_vs_ccc7_docx_safe.jpg",
    "sector_rank": ROOT / r"meeting_preparation\figures\docx_safe_images\p6_desnz_2050_residual_emissions_ranking_docx_safe.jpg",
    "historical_rate": ROOT / r"p8_uncertainty_framework\figures\p8_1_rate_benchmark_comparison.jpg",
    "sector_band": ROOT / r"p8_uncertainty_framework\figures\p8_2_sector_uncertainty_band_2050.jpg",
    "linkage_score": ROOT / r"p8_uncertainty_framework\figures\submission_cleaned\figure_4_15_linkage_score_vs_residual_submission_cleaned.jpg",
    "similarity": ROOT / r"p8_uncertainty_framework\figures\submission_cleaned\figure_4_16_dependency_similarity_submission_cleaned.jpg",
    "weighted": ROOT / r"p8_uncertainty_framework\figures\submission_cleaned\figure_4_17_linkage_weighted_ranking_submission_cleaned.jpg",
    "quadrant": ROOT / r"p8_uncertainty_framework\figures\p8_3_quadrant_outcomes.jpg",
    "mini": ROOT / r"p8_uncertainty_framework\figures\p8_3_mini_scenario_residual_gaps.jpg",
    "offset": ROOT / r"p8_uncertainty_framework\figures\p8_3_negative_emissions_credit_sensitivity.jpg",
    "neso_cost": ROOT / r"p7_neso_uncertainty\figures\p7_2050_cost_sensitivity_range.jpg",
}

DATA = {
    "p4": ROOT / r"p4_p5_local_reproduction\tables\p4_final_metrics_table.csv",
    "p5": ROOT / r"p4_p5_local_reproduction\tables\p5_final_benchmark_metrics_table.csv",
    "rate": ROOT / r"p8_uncertainty_framework\tables\p8_1_future_pathway_rate_comparison.csv",
    "sector_band": ROOT / r"p8_uncertainty_framework\tables\p8_2_sector_uncertainty_band_2050.csv",
    "scenario": ROOT / r"p8_uncertainty_framework\tables\p8_3_2x2_scenario_matrix.csv",
    "near": ROOT / r"p8_uncertainty_framework\tables\p8_3_near_miss_summary.csv",
    "linkage": ROOT / r"p8_uncertainty_framework\tables\p8_4_sector_linkage_deepening_metrics.csv",
}


def load_csv(key):
    p = DATA[key]
    return pd.read_csv(p) if p.exists() else pd.DataFrame()


P4 = load_csv("p4")
P5 = load_csv("p5")
RATE = load_csv("rate")
SECTOR = load_csv("sector_band")
SCEN = load_csv("scenario")
NEAR = load_csv("near")
LINKAGE = load_csv("linkage")


def metric(df, metric_id, col):
    if df.empty:
        return ""
    row = df.loc[df.get("metric_id", "") == metric_id]
    if row.empty:
        return ""
    return row.iloc[0].get(col, "")


CB6_GAP = float(metric(P4, "CB6_official_gap", "value_MtCO2e") or 737.469409)
GAP_2030 = float(metric(P5, "2030_gap_vs_CCC7", "gap_MtCO2e") or 72.211962)
GAP_2035 = float(metric(P5, "2035_gap_vs_CCC7", "gap_MtCO2e") or 154.795312)
GAP_2050 = float(metric(P5, "2050_gap_vs_CCC7", "gap_MtCO2e") or 325.398436)
CUM_GAP = float(metric(P5, "cumulative_2025_2050_gap_vs_CCC7", "gap_MtCO2e") or 4550.725308)


def rate_value(pathway, period):
    if RATE.empty:
        return None
    row = RATE[(RATE["pathway"].str.contains(pathway, regex=False)) & (RATE["period"] == period)]
    if row.empty:
        return None
    return float(row.iloc[0]["avg_annual_reduction_MtCO2e_per_year"])


DESNZ_2035_2050_RATE = rate_value("including IAS", "2035-2050") or 1.166685
CCC7_2035_2050_RATE = rate_value("CCC Seventh", "2035-2050") or 12.540226


def fmt(x, digits=1):
    try:
        return f"{float(x):,.{digits}f}"
    except Exception:
        return str(x)


def set_doc_defaults(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.82)
    sec.right_margin = Inches(0.82)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.3)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(5)
    for style_name, size, before, after in [
        ("Heading 1", 16, 14, 7),
        ("Heading 2", 12.5, 10, 5),
        ("Heading 3", 11, 7, 3),
    ]:
        st = styles[style_name]
        st.font.name = "Arial"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
    for sec in doc.sections:
        sec.footer.paragraphs[0].text = "TBPS6"
        sec.footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT


def set_props(doc, title):
    props = doc.core_properties
    props.title = title
    props.author = "TBPS6"
    props.last_modified_by = "TBPS6"
    props.subject = ""
    props.keywords = ""
    props.comments = ""
    props.category = ""


def add_title(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(18)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.name = "Arial"
        r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r2.font.size = Pt(10.5)
    doc.add_paragraph()


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=110, bottom=80, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def cell_text(cell, text, bold=False, size=9.2):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(str(text))
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(size)
    r.bold = bold


def set_widths(table, widths):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(width))
        grid.append(gc)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None, header_fill="F2F2F2", font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell_text(table.rows[0].cells[i], h, bold=True, size=font_size)
        shade(table.rows[0].cells[i], header_fill)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value, size=font_size)
    if widths:
        set_widths(table, widths)
    doc.add_paragraph()
    return table


def add_note(doc, title, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    r = p.add_run(title + " ")
    r.bold = True
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.name = "Arial"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r2.font.size = Pt(10)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Inches(0.26)
        p.paragraph_format.first_line_indent = Inches(-0.16)
        p.paragraph_format.space_after = Pt(3)
        p.add_run("- " + item)


def add_picture(doc, path, caption, width=6.1):
    path = Path(path)
    if not path.exists():
        p = doc.add_paragraph(f"[Figure file not found: {path.name}]")
        p.runs[0].italic = True
        return
    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = cap.add_run(caption)
    rr.italic = True
    rr.font.size = Pt(8.8)
    rr.font.name = "Arial"
    rr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_page_break(doc):
    doc.add_page_break()


def make_evidence_pack():
    doc = Document()
    set_doc_defaults(doc)
    set_props(doc, "TBPS6 Dissertation Meeting Evidence Pack")
    add_title(
        doc,
        "Dissertation Meeting Evidence Pack",
        "Assessing the UK's Net-Zero Delivery Gap: DESNZ projections, CCC benchmarks, sectoral drivers and uncertainty to 2050 | Candidate TBPS6",
    )
    add_note(
        doc,
        "Meeting purpose:",
        "This pack is intended to support a first-draft discussion. The main aim is to test whether the dissertation argument is now clear enough, evidence-led enough and appropriately cautious for final revision.",
    )
    add_table(
        doc,
        ["Item", "Current position"],
        [
            ["Draft status", "A complete first-draft structure is now in place: Introduction, Literature Review, Data and Methods, Results, Discussion, Conclusion, References and Appendices."],
            ["Main contribution", "The work has moved from separate pathway comparisons into a delivery-gap diagnosis: size, timing, sectoral concentration, uncertainty and governance implications."],
            ["Evidence base", "DESNZ current-policy projections, CCC Sixth/Seventh Carbon Budget evidence, NESO FES 2025 material, DESNZ historical statistics and selected peer-reviewed literature."],
            ["Word count", "Formal count on the current cleaned draft: 9,870 words under the course inclusion/exclusion rule."],
        ],
        widths=[2300, 7060],
    )

    doc.add_heading("1. Dissertation Argument in One Page", level=1)
    add_bullets(
        doc,
        [
            f"The DESNZ current-policy baseline is not target-consistent against the CCC Seventh Carbon Budget benchmark: the including-IAS 2050 gap is about {fmt(GAP_2050)} MtCO2e.",
            f"The problem appears before 2050: the official Sixth Carbon Budget gap is about {fmt(CB6_GAP)} MtCO2e, and the DESNZ baseline slows to about {fmt(DESNZ_2035_2050_RATE)} MtCO2e/year over 2035-2050 while CCC7 continues at about {fmt(CCC7_2035_2050_RATE)} MtCO2e/year.",
            "Residual emissions are not evenly distributed. Buildings/product uses, domestic transport, industry, electricity supply, agriculture and IAS explain most of the pressure, but each sector has a different uncertainty profile.",
            "The sector-linkage analysis suggests that electricity, buildings and transport should be read as a coupled electrification block, while industry, IAS, agriculture and land-related residuals require more sector-specific interpretation.",
            "The 2x2 scenario matrix separates domestic policy delivery from international conditions. Within this bounded scenario design, domestic delivery produces the larger change in residual emissions, while external costs, fuel exposure and technology conditions still matter.",
        ],
    )

    add_table(
        doc,
        ["Research question", "How the draft answers it"],
        [
            ["Main RQ", "The UK delivery gap is large under DESNZ current-policy projections and should be interpreted through benchmark alignment, timing, residual-sector drivers and uncertainty, not as a single endpoint number."],
            ["RQ1", "DESNZ remains above CCC7 in every benchmark year: around 72.2 MtCO2e in 2030, 154.8 MtCO2e in 2035 and 325.4 MtCO2e in 2050 including IAS."],
            ["RQ2", "Residual emissions concentrate in a small group of sectors. Sectoral uncertainty and linkage scores change the interpretation from a simple ranking into a residual-risk profile."],
            ["RQ3", "Scenario results show near-miss outcomes only under strong domestic delivery plus sufficient offset/removal availability; delayed delivery leaves large residuals even when external conditions are more supportive."],
        ],
        widths=[1700, 7660],
    )

    add_page_break(doc)
    doc.add_heading("2. Evidence Design and Method Boundaries", level=1)
    add_table(
        doc,
        ["Evidence block", "Role in the dissertation", "Interpretation boundary"],
        [
            ["DESNZ EEP 2024-2050", "Current-policy baseline for annual gaps, carbon-budget comparison and sectoral residuals.", "Used as a baseline, not as an inevitable forecast."],
            ["CCC6/CCC7", "Target-consistent benchmark and statutory carbon-budget context.", "Used as benchmark evidence, not a neutral forecast."],
            ["Historical DESNZ statistics", "Delivered-rate benchmark and plausibility anchor.", "Past rates are diagnostic; the sectoral mix may not repeat."],
            ["NESO FES 2025", "External cost, fuel exposure and pathway context.", "Used to interpret conditions, not to run a full energy-system model."],
            ["Author-constructed indicators", "Sector uncertainty bands, linkage scores, 2x2 scenario matrix and near-miss tests.", "Transparent stress tests; no probabilities or causal coefficients are assigned."],
        ],
        widths=[2100, 3900, 3360],
    )
    add_note(
        doc,
        "Method stance:",
        "The draft deliberately avoids presenting the results as a forecast or a least-cost model. It is a pathway-diagnostic dissertation: it asks whether current-policy projections look aligned with target-consistent benchmarks, and where the delivery risk is concentrated.",
    )

    doc.add_heading("3. Result 1: Aggregate Baseline-Benchmark Gap", level=1)
    add_table(
        doc,
        ["Metric", "Value", "Interpretation"],
        [
            ["2030 DESNZ-CCC7 annual gap", f"{fmt(GAP_2030)} MtCO2e", "The gap is visible before the 2030s are over."],
            ["2035 DESNZ-CCC7 annual gap", f"{fmt(GAP_2035)} MtCO2e", "The gap widens around the CB6 delivery period."],
            ["2050 DESNZ-CCC7 annual gap", f"{fmt(GAP_2050)} MtCO2e", "The current-policy baseline remains far above a target-consistent pathway."],
            ["Cumulative 2025-2050 gap", f"{fmt(CUM_GAP)} MtCO2e", "A supplementary indicator of the scale of pathway divergence."],
            ["Official CB6 gap", f"{fmt(CB6_GAP)} MtCO2e", "The issue is not only a distant 2050 endpoint problem."],
        ],
        widths=[3100, 1900, 4360],
    )
    add_picture(doc, FIGURES["baseline"], "Figure A. DESNZ current-policy baseline compared with CCC target-consistent pathways.", 6.25)
    add_picture(doc, FIGURES["gap"], "Figure B. Annual DESNZ including-IAS gap relative to CCC7 Balanced Pathway.", 6.25)

    doc.add_heading("4. Result 2: Historical Delivered-Rate Reality Check", level=1)
    add_table(
        doc,
        ["Pathway period", "Average annual reduction", "Interpretation"],
        [
            ["DESNZ including IAS, 2035-2050", f"{fmt(DESNZ_2035_2050_RATE)} MtCO2e/year", "Current-policy decarbonisation nearly flattens after 2035."],
            ["CCC7 Balanced Pathway, 2035-2050", f"{fmt(CCC7_2035_2050_RATE)} MtCO2e/year", "Demanding, but close to the long-run historical delivered-rate scale."],
            ["Main inference", "Rate gap", "The issue is not simply that CCC7 is ambitious; it is that the DESNZ baseline loses momentum while the target-consistent pathway continues."],
        ],
        widths=[3300, 2300, 3760],
    )
    add_picture(doc, FIGURES["historical_rate"], "Figure C. Historical average annual reductions compared with DESNZ and CCC future pathway rates.", 6.2)

    doc.add_heading("5. Result 3: Residual Sectors, Uncertainty and Linkages", level=1)
    top_linkage = []
    if not LINKAGE.empty:
        for _, r in LINKAGE.sort_values("linkage_priority_rank").head(6).iterrows():
            top_linkage.append([
                str(int(r["linkage_priority_rank"])),
                r["tes_sector"],
                r["linkage_typology"],
                fmt(r["desnz_central_2050_MtCO2e"]),
                str(int(r["system_linkage_score"])),
                f"{float(r['linkage_weighted_residual_share_pct']):.1f}%",
            ])
    add_table(
        doc,
        ["Rank", "Sector", "Linkage family", "2050 residual", "Score", "Weighted share"],
        top_linkage,
        widths=[700, 2200, 2600, 1350, 850, 1660],
        font_size=8.8,
    )
    add_picture(doc, FIGURES["sector_band"], "Figure D. Sectoral 2050 uncertainty band with CCC7 mapped benchmark.", 6.2)
    add_picture(doc, FIGURES["linkage_score"], "Figure E. Linkage score compared with 2050 residual emissions.", 6.0)
    add_picture(doc, FIGURES["weighted"], "Figure F. Linkage-weighted residual-risk ranking.", 6.0)
    add_note(
        doc,
        "Important interpretation:",
        "The linkage score is dependency evidence, not causal attribution. It helps identify sectors that depend on shared transition systems such as clean power, electrification infrastructure, hydrogen/CCUS, behaviour change, land/non-CO2 constraints and international technology or fuel conditions.",
    )

    doc.add_heading("6. Result 4: Scenario Matrix, Near-Miss Timing and External Conditions", level=1)
    scenario_rows = []
    if not NEAR.empty:
        default = NEAR[NEAR["offsets_available_MtCO2e"] == 50]
        for _, r in default.iterrows():
            scenario_rows.append([
                r["quadrant_id"],
                r["scenario_name"],
                f"{float(r['median_net_2050_MtCO2e']):.1f}",
                f"{float(r['median_delay_years']):.1f}",
                f"{float(r['near_miss_share'])*100:.0f}%",
            ])
    add_table(
        doc,
        ["Quadrant", "Scenario", "Median net 2050 residual", "Median delay", "Near-miss share"],
        scenario_rows,
        widths=[1000, 3700, 1900, 1300, 1460],
        font_size=8.7,
    )
    add_picture(doc, FIGURES["quadrant"], "Figure G. 2x2 scenario matrix: median 2050 residual after default offsets.", 6.0)
    add_picture(doc, FIGURES["offset"], "Figure H. Sensitivity of quadrant-level median residuals to offset/removal assumptions.", 6.0)
    add_picture(doc, FIGURES["neso_cost"], "Figure I. NESO 2050 total energy cost sensitivity ranges used as external-condition context.", 6.0)
    add_note(
        doc,
        "Near-miss framing:",
        "The draft now treats target-year delays cautiously. It highlights near-miss and short-delay cases rather than giving undue interpretive weight to far-delayed stress-test years.",
    )

    doc.add_heading("7. Discussion Points for the Meeting", level=1)
    add_bullets(
        doc,
        [
            "Is the main dissertation argument now clear: gap size -> timing -> sectoral drivers -> uncertainty -> policy interpretation?",
            "Is the current treatment of historical delivered rates sufficiently cautious, or should it be framed even more explicitly as a plausibility anchor?",
            "Does the sector-linkage analysis look useful as a dissertation-level diagnostic, given that it is not a causal model?",
            "Is the 2x2 uncertainty framework at the right level of detail for the dissertation, or should one part be moved to an appendix?",
            "For final revision, should the priority be tightening the Discussion, strengthening the Literature Review, improving presentation of Results, or compressing appendices?",
        ],
    )
    doc.save(EVIDENCE_PACK)


def make_oral_script():
    doc = Document()
    set_doc_defaults(doc)
    set_props(doc, "TBPS6 One-Hour Meeting Oral Script")
    add_title(
        doc,
        "One-Hour Meeting Oral Script",
        "Cued to the dissertation meeting evidence pack | English script with brief Chinese guidance | Candidate TBPS6",
    )
    add_note(
        doc,
        "How to use this script:",
        "Do not read every word mechanically. Use the English script as a speaking base, use the Chinese notes to remind yourself of the purpose, and pause at the suggested questions so the meeting becomes a discussion rather than a monologue.",
    )
    add_table(
        doc,
        ["Time", "Evidence-pack cue", "Meeting purpose"],
        [
            ["0-5 min", "Cover and Section 1", "Open the meeting and state what kind of feedback you need."],
            ["5-12 min", "Section 1", "Explain the research question and main argument."],
            ["12-20 min", "Section 2", "Defend data, method and boundaries."],
            ["20-32 min", "Sections 3-4", "Present aggregate gap, CB6 timing and historical rate benchmark."],
            ["32-44 min", "Section 5", "Present sectoral residuals, uncertainty and linkages."],
            ["44-53 min", "Section 6", "Present 2x2 uncertainty, near-miss and external conditions."],
            ["53-60 min", "Section 7", "Discuss final revision priorities with Neil."],
        ],
        widths=[1200, 2800, 5360],
    )

    sections = [
        (
            "0-5 min | Opening",
            "Evidence pack cue: Cover page and Section 1, Dissertation Argument in One Page.",
            "中文提示：先说明这次不是简单汇报进度，而是请 Neil 检验你的初稿主论证是否已经站稳。",
            [
                "Thanks for meeting with me today. I have brought a more complete first-draft evidence pack rather than a separate set of small analysis notes. What I would like to do is walk through the dissertation argument as it now stands, then use the second half of the meeting to test where it still needs strengthening.",
                "The draft now has the full structure in place: introduction, literature review, data and methods, results, discussion, conclusion, references and appendices. The word count is also now within the expected range. The main change since the earlier work is that I am no longer treating the DESNZ-CCC comparison, the sector analysis and the uncertainty framework as separate pieces. I have tried to turn them into one connected argument about the UK's net-zero delivery gap.",
                "The argument I want to test today is this: the UK delivery gap should not be interpreted only as a 2050 endpoint difference between DESNZ and CCC. It is a pathway problem. It appears in the Sixth Carbon Budget period, it is visible in the post-2035 reduction-rate slowdown, it is concentrated in a small set of residual sectors, and it is sensitive to domestic delivery and international conditions.",
                "I will keep the presentation structured, but I would like to pause after each major part and ask you whether the framing is academically defensible. I am especially interested in whether the sector-linkage and 2x2 uncertainty framework are at the right level for an MSc dissertation, or whether I should simplify or move more detail into the appendix.",
                "In practical terms, I would like to use the meeting in three layers. First, I will explain the dissertation as a coherent argument. Second, I will show the main evidence behind that argument, especially the aggregate gap, the timing result, the sectoral residuals and the scenario tests. Third, I would like to ask for your judgement on final revision priorities, because I am now at the stage where the largest risk is not lack of material, but deciding what should remain central and what should be compressed.",
                "I have also tried to make the draft more cautious than some earlier versions. For example, I avoid saying that a pathway definitely fails by a precise far-future year, because that can give too much weight to a simple extrapolation. Instead, I focus on whether a case is on-time, near-miss, moderate miss or large miss. I also avoid saying that sector linkages prove causality. The linkage work is meant to support interpretation of dependency, not to claim that I have identified causal policy effects.",
                "So the way I would like to frame today's discussion is: does this now look like a credible dissertation-level diagnostic of the UK's net-zero delivery gap, and what are the two or three changes that would most improve it before the next submitted version?",
            ],
            "Question to Neil: Does this look like the right kind of discussion agenda for a first-draft meeting, or would you prefer me to focus first on Results and only then on structure?",
        ),
        (
            "5-12 min | Research Question and Main Claim",
            "Evidence pack cue: Section 1, the research-question table and the bullet list of the dissertation argument.",
            "中文提示：这一部分讲清楚研究问题不是“英国能不能净零”这么宽，而是“官方当前政策路径与目标一致路径之间的 delivery gap 如何理解”。",
            [
                "My main research question is: how large is the UK's projected net-zero delivery gap under DESNZ current-policy projections relative to target-consistent benchmarks, and how should that gap be interpreted through timing, sectors and uncertainty?",
                "The three sub-questions then break this down. RQ1 asks how the DESNZ current-policy baseline compares with CCC target-consistent pathways and carbon-budget benchmarks to 2050. RQ2 asks which sectors drive the residual gap and how uncertainty and sectoral linkages affect that interpretation. RQ3 asks how domestic delivery and external conditions change the interpretation of target achievement, near-miss cases and residual-emissions risk.",
                "The main claim is deliberately narrower than a full net-zero policy evaluation. I am not trying to build a least-cost energy model or forecast exactly when the UK will hit net zero. Instead, I am using official pathways and datasets to diagnose whether the current-policy baseline is aligned with a target-consistent pathway, where the gap appears, and what kind of delivery risks explain it.",
                f"The headline number is that the DESNZ including-IAS current-policy baseline is about {fmt(GAP_2050)} MtCO2e above the CCC7 Balanced Pathway in 2050. But I am trying not to leave the dissertation at that number. The more important interpretation is that the gap emerges earlier, especially around the Sixth Carbon Budget, and then becomes a sectoral and governance problem.",
                "This is why the Results chapter moves from aggregate gap metrics to carbon-budget and reduction-rate diagnostics, then to residual sectors, then to historical delivered-rate checks, sector uncertainty and the scenario matrix. The sequence is meant to make the dissertation feel cumulative, rather than like several unrelated notebook outputs.",
                "The reason I have kept the wording around 'delivery gap' is that it captures the difference between a legal or policy target and the pathway implied by current measures. A target can be legally fixed, but its credibility depends on interim delivery, sectoral implementation and the assumptions that sit behind target-consistent pathways. That is the conceptual link between the literature review and the empirical chapters.",
                "One thing I have tried to improve is the distinction between a baseline and a benchmark. DESNZ is not being treated as a bad forecast, and CCC is not being treated as perfect truth. They answer different questions. DESNZ asks: if current policies are projected forward, what does the emissions pathway look like? CCC asks: what does a target-consistent route look like under a policy-relevant benchmark? The gap between them is therefore not simply a prediction error. It is a measure of alignment pressure.",
                "The sub-questions are designed to stop the dissertation from being too narrow. RQ1 gives the scale and timing. RQ2 asks what is inside the gap. RQ3 asks how robust the interpretation is when delivery and external conditions vary. In that sense, the dissertation moves from measurement to explanation to conditional judgement.",
                "If I had to say the thesis in one sentence, I would say: the UK's projected net-zero delivery gap is large, visible before 2050, concentrated in residual sectors that are partly linked through shared transition systems, and unlikely to be closed by favourable external conditions unless domestic implementation is also strong.",
            ],
            "Question to Neil: Is this research question still too broad, or is the current version now focused enough for the dissertation?",
        ),
        (
            "12-20 min | Data and Method Boundaries",
            "Evidence pack cue: Section 2, Evidence Design and Method Boundaries.",
            "中文提示：重点解释为什么 DESNZ、CCC、NESO 的角色不同，以及为什么不能把结果说成 forecast 或 causal proof。",
            [
                "The evidence design is built around a clear separation between source roles. DESNZ is the current-policy baseline. CCC is the target-consistent benchmark. Historical DESNZ statistics provide delivered-rate anchors. NESO is used for external-condition context, especially costs, fuel exposure and pathway assumptions.",
                "A key methodological point is that I do not average these pathways together. I use them for different purposes. DESNZ answers: what does the current-policy baseline look like? CCC answers: what would a target-consistent benchmark imply? NESO helps interpret what external technology, cost and fuel conditions might mean for transition risk.",
                "The draft also tries to control accounting boundaries. Including-IAS, excluding-IAS, carbon-budget accounting, LULUCF and removals are not mixed as if they are identical. In the sector chapter, I use DESNZ TES sectors for internal consistency, then treat the CCC sector bridge as broad diagnostic evidence rather than exact one-to-one attribution.",
                "The author-constructed indicators are intentionally compact. The historical rate benchmark is a plausibility anchor, not a claim that the past can simply repeat. The sector-linkage score is dependency evidence, not causal attribution. The 2x2 framework is a bounded stress-test design, not a probability model.",
                "I have tried to make these boundaries explicit because they help protect the dissertation from overclaiming. At the same time, I think they still allow useful analytical judgement: even with these caveats, the DESNZ baseline remains misaligned, the gap is sectorally concentrated, and weak domestic delivery produces a much larger residual in the scenario tests.",
                "The cleaning and harmonisation work is also important for the credibility of the results. The DESNZ and CCC annual pathways are placed on a common annual comparison basis, while the statutory carbon-budget result is kept on the official comparison basis. I do not combine including-IAS and excluding-IAS values without saying which basis is being used. That is why the dissertation sometimes reports more than one value: the purpose is to avoid a false precision problem.",
                "For the sectoral work, I made a cautious choice. I use DESNZ TES sectors for the central residual ranking because that keeps the internal baseline consistent. Where I compare to CCC, I call it a broad sector bridge rather than exact reconciliation. This matters because CCC and DESNZ do not always draw the same boundaries around IAS, LULUCF, F-gases, product uses and removals.",
                "For the uncertainty work, the indicators are intentionally transparent. The sector uncertainty band uses accelerated and delayed anchors rather than a black-box model. The linkage score is a simple 0-1-2 scoring system across shared transition drivers. The 2x2 matrix uses domestic delivery and external conditions because those two dimensions match the dissertation's central argument: some risks are about UK implementation, while others are about technology costs, international fuels and supply-chain context.",
                "I have added robustness discussion because some settings are judgement-sensitive. For example, changing a near-miss threshold can move a few boundary cases, but it does not remove the contrast between aligned transition and weak or delayed delivery. Changing offset assumptions mechanically lowers residuals, but it does not make the large-miss quadrants target-consistent. That is the sort of limited but useful robustness claim I am making.",
                "The method chapter therefore tries to be honest about what the dissertation can and cannot do. It can quantify and interpret official-pathway misalignment. It can identify sectors and conditions that matter. It cannot estimate the abatement effect of each policy instrument or prove causal sector spillovers. I think that distinction is important for keeping the dissertation academically defensible.",
            ],
            "Question to Neil: Are the method boundaries stated strongly enough, or do you think the dissertation still risks sounding more causal or predictive than the evidence supports?",
        ),
        (
            "20-32 min | RQ1: Aggregate Gap, CB6 and Historical Reality Check",
            "Evidence pack cue: Sections 3 and 4; Figures A-C.",
            "中文提示：这里是论文最硬的基础结果。讲数字，但不要只堆数字，要强调 timing 和 momentum。",
            [
                "The first empirical block answers RQ1. The DESNZ baseline is above the CCC7 benchmark in every benchmark year. The including-IAS gap is about 72.2 MtCO2e in 2030, 154.8 MtCO2e in 2035 and 325.4 MtCO2e in 2050. So the gap is not something that suddenly appears at the endpoint. It widens through the 2030s and 2040s.",
                f"The official Sixth Carbon Budget comparison is also important. The projected official CB6 gap is about {fmt(CB6_GAP)} MtCO2e. This gives the dissertation a nearer-term governance point: the problem is not just whether the UK reaches net zero in 2050, but whether the pathway is credible through the interim statutory budget period.",
                f"The reduction-rate result helps explain the divergence. For the 2035-2050 period, the DESNZ including-IAS baseline reduces at only about {fmt(DESNZ_2035_2050_RATE)} MtCO2e per year, while CCC7 reduces at about {fmt(CCC7_2035_2050_RATE)} MtCO2e per year. That is why I describe the DESNZ pathway as an inertia baseline. It does not continue decarbonisation at the pace implied by the target-consistent benchmark.",
                "The historical delivered-rate benchmark adds another layer. It shows that CCC7 is demanding, especially in the early period, but it is not arbitrary. Some parts of the pathway are broadly comparable with historical delivered-rate scales, while DESNZ becomes much slower than recent historical delivery after 2035.",
                "I think this makes the RQ1 conclusion stronger than a simple DESNZ-versus-CCC gap. The issue is the combination of endpoint residual emissions, interim carbon-budget misalignment and post-2035 momentum loss.",
                "If you look at the first pathway figure, the visual message is that the DESNZ and CCC pathways separate steadily rather than just at a single point. The CCC7 pathway moves down towards net zero, while DESNZ remains much flatter in the long run. The annual-gap figure then makes the divergence explicit by showing the widening MtCO2e gap year by year.",
                "The carbon-budget result changes the policy interpretation. A 2050 gap can sometimes sound like a distant long-term problem. The CB6 result makes it nearer-term and statutory. It shows that the current-policy pathway is already misaligned with an interim governance checkpoint, so the dissertation is not relying only on a speculative 2050 endpoint comparison.",
                "The historical-rate result is where I have tried to respond to the concern that a CCC pathway might simply be ambitious by definition. I compare future required rates with achieved historical reductions. I am careful not to say that historical reductions can simply be repeated, because the UK has already taken many easier reductions, especially in power generation. But the comparison still helps distinguish an ambitious target-consistent pathway from a baseline that slows too much.",
                "In the oral explanation, I would emphasise that historical achieved reductions are being used to judge realism in a limited way. They are a benchmark of scale, not a proof of feasibility. A reduction rate similar to the historical average is not automatically easy, because the remaining sectors are harder. But a rate much slower than historical delivery suggests that the baseline is not aligned with the momentum required for net zero.",
                "So for RQ1, I would summarise the result as three connected findings: first, the 2050 gap is large; second, the gap is already a CB6 governance issue; third, the long-run DESNZ pathway slows at exactly the time when a target-consistent pathway still needs sustained reductions.",
            ],
            "Question to Neil: Is the historical delivered-rate comparison useful enough in the main Results chapter, or should I move more of it to the appendix and keep only the headline interpretation?",
        ),
        (
            "32-44 min | RQ2: Sectoral Residuals, Uncertainty and Linkage",
            "Evidence pack cue: Section 5; Figures D-F and the linkage-priority table.",
            "中文提示：这里要回应 Neil 关心的部门联系问题。核心是 buildings, transport, electricity 不是孤立部门；industry/IAS/agriculture 是另一类残余风险。",
            [
                "The second Results block answers RQ2. The residual gap is concentrated. Buildings and product uses are the largest 2050 residual in the DESNZ projection, followed by domestic transport, industry, electricity supply, agriculture and IAS. The top residual sectors account for most of the including-IAS 2050 total.",
                "The sectoral uncertainty band is useful because it prevents the Results from relying only on central estimates. Domestic transport, for example, is strongly implementation-sensitive: accelerated delivery can pull the 2050 residual down, while delayed delivery leaves it much higher. Agriculture and IAS are different. They retain residual emissions even under more favourable assumptions, which means they need a different interpretation.",
                "The sector-linkage analysis is my attempt to address the fact that residual sectors are not independent. Buildings and transport both depend heavily on clean power and electrification infrastructure. Electricity supply itself is part of that system, because decarbonised power is an enabling condition for heat and transport electrification. Industry has a different linkage profile, with stronger hydrogen, CCUS and international-technology exposure.",
                "This is why I classify the sectors into linkage families. Buildings, domestic transport and electricity sit in a coupled electrification block. Industry and fuel supply are more connected to hydrogen, CCUS and industrial clusters. Agriculture, LULUCF and waste sit closer to land, non-CO2 and removals constraints. IAS is especially exposed to international fuel and technology conditions.",
                "The important caveat is that this is dependency evidence, not causal attribution. I am not claiming that the figure proves, for example, that a delay in clean power quantitatively causes a given MtCO2e residual in buildings. The point is more modest: the sectors share transition dependencies, so the residual-risk interpretation should not treat them as isolated policy silos.",
                "The reason I think this sectoral block is important is that the aggregate gap by itself does not tell us what kind of policy problem the UK faces. A 325 MtCO2e gap could come from many different configurations. In this draft, the residual ranking shows that the pressure is concentrated in sectors with different delivery mechanisms: heat and buildings, transport stock turnover, industrial process emissions, electricity system decarbonisation, land and non-CO2 emissions, and international aviation and shipping.",
                "The uncertainty band then changes the reading of the ranking. A large residual sector with high delivery sensitivity is different from a residual sector that remains large under both accelerated and delayed assumptions. For transport, the interpretation is about implementation pace, fleet turnover, infrastructure and demand. For agriculture and IAS, the interpretation is more about persistent residual emissions, harder mitigation options and the role of removals or credits after gross mitigation.",
                "The linkage figures are meant to make the sector story more system-aware. Buildings do not decarbonise in isolation if heat electrification depends on electricity system readiness and local networks. Road transport does not decarbonise in isolation if EV uptake depends on charging infrastructure, clean power and vehicle stock turnover. Industry has a different dependency structure because hydrogen, CCUS, industrial clusters and international competitiveness matter more.",
                "The similarity heatmap is also useful here, but I would be careful in the meeting not to oversell it. It shows which sectors have similar dependency profiles in the constructed scoring framework. It does not show observed co-movement in emissions, and it does not prove that one sector causes another to fail. It is a way of making the residual-sector interpretation less siloed.",
                "My interpretation is that the linked electrification block is the clearest example of sectors that can move together if enabling systems are delivered. In contrast, agriculture, IAS and some industrial emissions are more independent or hard-to-abate in the sense that they cannot be solved only by clean power. They need different instruments, including demand-side measures, technology innovation, standards, land-use choices, or post-processing through removals and credits.",
                "This means the policy implication is not one-size-fits-all. For buildings and transport, delivery indicators should focus on heat pumps, retrofit, EV uptake, chargers and grid readiness. For electricity, indicators should focus on clean generation, flexibility and network connection. For industry, they should focus on CCUS and hydrogen cluster delivery, low-carbon fuel use and industrial investment certainty. For agriculture and IAS, indicators need to be more cautious because residuals are less easily eliminated.",
            ],
            "Question to Neil: Does this linkage analysis add enough value as a dissertation-level diagnostic, or should I simplify the scoring and use it more as a qualitative interpretation aid?",
        ),
        (
            "44-53 min | RQ3: 2x2 Uncertainty, Near-Miss and External Conditions",
            "Evidence pack cue: Section 6; Figures G-I and the scenario table.",
            "中文提示：不要说“domestic delivery 是最强决定因素”，要说“在本文 bounded scenario design 里，domestic delivery 带来更大的 residual change”。",
            [
                "RQ3 brings the uncertainty pieces together. The 2x2 matrix separates UK domestic policy delivery from external conditions. The domestic axis captures whether implementation is high or low. The external axis captures whether technology costs, fuel exposure and supply-chain conditions are supportive or constrained.",
                "Under default offset assumptions, the aligned-transition quadrant has a median 2050 net residual of about 52.6 MtCO2e and a median delay of about 4.8 years. That is the only quadrant where near-miss outcomes become common in the tested mini-scenarios. By contrast, weak or delayed domestic delivery leaves much larger residuals even when external conditions are supportive.",
                "I am being careful with the wording here. Within the bounded scenario design used in the dissertation, domestic delivery produces the larger change in residual emissions. That does not mean I have proven a general macroeconomic law. It means that, under the tested assumptions, weak domestic implementation has a larger residual effect than supportive external conditions can offset.",
                "The external conditions still matter. NESO cost and fuel-exposure indicators help explain why global technology costs, fossil-price exposure and supply-chain conditions affect UK transition risk. The point is not to run a full welfare or macroeconomic analysis, but to show that international conditions can change the difficulty of delivery.",
                "The near-miss framing has also been tightened. I no longer want to emphasise far-delayed target years, because very late extrapolated dates are not very informative. The useful part is the contrast between near-miss cases and large-miss cases. Near-miss cases show where a pathway is close enough that timing, offsets and delivery acceleration matter; large-miss cases show where the issue is still structural mitigation, not just accounting or timing.",
                "The matrix has four quadrants. U1 is aligned transition: high domestic delivery and supportive external conditions. U2 is domestic delivery under external constraints: UK delivery is still high, but external costs, supply chains or fuel conditions are less favourable. U3 is weak delivery despite supportive conditions: global conditions help, but UK implementation is weak. U4 is delayed transition risk: both domestic delivery and external conditions are weak.",
                "The most useful comparison is not the exact target year for every pathway. It is the difference in residual scale across quadrants. U1 can move into near-miss territory under strong delivery and sufficient offsets. U2 is still a miss, but it is less severe than weak-delivery cases. U3 and U4 show that favourable external conditions cannot compensate fully for weak domestic implementation in the tested scenario space.",
                "This is where the international conditions discussion fits. Falling clean-technology costs can help the UK, particularly where electrification and clean power are central. Lower technology costs can reduce delivery difficulty and make faster uptake more plausible. But the dissertation does not claim that global cost declines automatically deliver UK net zero. Domestic policy still has to convert those conditions into deployment, infrastructure, finance, standards and behavioural change.",
                "The offset and removals sensitivity is deliberately treated as post-processing. I do not allow removals to erase the need for gross mitigation. Instead, the sensitivity asks: if a residual remains in 2050, how much could different offset or removal values change the classification from large miss to moderate miss or near miss? The answer is that offsets improve some boundary cases, especially U1, but they do not rescue weak or delayed delivery pathways.",
                "If Neil asks why I use a near-miss threshold, I would explain that it is a diagnostic scale marker, not a relaxation of net zero. It is there to distinguish cases where the residual is small enough that timing and offsets become relevant from cases where the residual is so large that the pathway is structurally misaligned. I can also say that sensitivity checks show that changing the threshold changes some labels, but not the main contrast between quadrants.",
                "Overall, RQ3 is not trying to produce a final probability-weighted forecast. It gives a disciplined way to discuss uncertainty: domestic delivery, external conditions, near-miss timing, and post-processing through removals. That gives the Discussion chapter a stronger basis than simply saying 'there is uncertainty'.",
            ],
            "Question to Neil: Is the near-miss threshold and offset treatment useful as a compact stress test, or should I reduce the quantitative emphasis and present it more cautiously?",
        ),
        (
            "53-60 min | Discussion, Final Revision and Questions",
            "Evidence pack cue: Section 7, Discussion Points for the Meeting.",
            "中文提示：最后不要问太泛的问题。请 Neil 帮你排优先级：Results clarity、Discussion、LR linkage、Appendix compression。",
            [
                "The Discussion chapter now tries to move from results to interpretation. The main policy message is not simply that the UK has a gap, but that the gap has different mechanisms across sectors. Buildings and transport require stock turnover, electrification infrastructure and clean-power readiness. Industry requires cluster delivery, CCUS or hydrogen infrastructure and investment certainty. Agriculture, IAS and land-related residuals are harder to eliminate fully and may require separate treatment, including removals or credits as post-processing rather than as a substitute for mitigation.",
                "The limitations chapter is also important. The dissertation does not directly model policy instruments, it does not assign probabilities to scenarios, and it does not prove causal spillovers between sectors. I have tried to make these limitations part of the academic discipline of the dissertation rather than something hidden at the end.",
                "For final revision, I would like your guidance on priority. My own sense is that the empirical Results are now the strongest part. The Discussion has improved, but it may still need tightening so that it does not repeat the Results. The Literature Review has enough material for an initial draft, but I may need to sharpen the critical comparison between official pathways, scenario literature and implementation credibility.",
                "The main decision I need from this meeting is where to spend the remaining revision time. Should I prioritise clearer Results presentation, stronger literature linkage, more cautious wording around uncertainty, or a more polished Discussion and Conclusion?",
                "My plan after this meeting is to revise the draft around your comments, then do a final audit of references, word count, appendices, figure clarity and consistency between the research questions and the conclusion.",
                "If we have time, I would also like to ask about the balance between main text and appendices. The draft now contains a lot of evidence, and I have tried to keep the key results in the main body while putting source audits and detailed tables in appendices. But I am still not fully sure whether the Results chapter is too dense. I would value your view on whether the key figures are doing enough work, or whether some tables should be shortened.",
                "Another final-revision question is the Literature Review. I have tried to make it problem-driven: target credibility, pathway design, sector coupling, uncertainty and negative-emissions governance. I think it supports the analysis, but it is still probably less strong than the Results. If the examiner expects a more critical literature review, I can strengthen the comparison between current-policy projections, target-consistent pathways and implementation credibility.",
                "The Conclusion now tries to say what the dissertation changes in understanding. It is not just 'there is a gap'. It is that the gap should be read as timed, sectorally concentrated and uncertainty-conditioned. But I would like to know whether that contribution sounds convincing or whether it should be expressed more modestly.",
                "So my final questions are quite practical. First, is the main Results sequence clear enough? Second, is the sector-linkage analysis worth keeping in the main body? Third, does the 2x2 scenario matrix make the uncertainty discussion stronger, or does it risk looking too constructed? Fourth, if I have limited revision time, what would you want me to improve first?",
                "After the meeting, I will revise the dissertation in that order. I will also do a final submission clean-up: Harvard consistency, figure captions, appendix references, word count, metadata and file anonymity.",
            ],
            "Question to Neil: If I only have time to substantially improve two parts before the next version, which two would you choose?",
        ),
    ]

    for title, cue, chinese, paras, question in sections:
        doc.add_heading(title, level=1)
        add_note(doc, "Cue:", cue)
        add_note(doc, "中文说明:", chinese)
        doc.add_heading("English speaking script", level=2)
        for para in paras:
            doc.add_paragraph(para)
        add_note(doc, "Pause / ask Neil:", question)

    doc.save(ORAL_SCRIPT)


def make_guide():
    doc = Document()
    set_doc_defaults(doc)
    set_props(doc, "TBPS6 Dissertation Detailed Understanding Guide")
    add_title(
        doc,
        "Detailed Dissertation Understanding Guide",
        "Bilingual explanation of the current thesis logic, content and defence points | Candidate TBPS6",
    )
    add_note(
        doc,
        "Purpose:",
        "This guide explains what each part of the current dissertation is doing. It is for preparation and understanding, not for submission.",
    )

    bilingual_rows = [
        (
            "Whole thesis logic",
            "The dissertation starts from a simple official-pathway comparison, but it does not stop there. It builds a layered diagnosis: first the DESNZ-CCC benchmark gap, then carbon-budget timing and reduction rates, then residual sectors, then historical realism, sectoral uncertainty, linkage and scenario uncertainty.",
            "整篇论文不是只做 DESNZ 和 CCC 的差值比较，而是把 gap 一层层解释清楚：先量化差距，再看碳预算时期和减排速度，再看哪些部门造成残余排放，最后加入历史现实性、部门不确定性、部门联系和情景不确定性。",
        ),
        (
            "Main research question",
            "The central question asks how large the UK's projected net-zero delivery gap is under current-policy projections, and how this gap should be interpreted through benchmark alignment, timing, sectoral concentration and uncertainty.",
            "核心研究问题关注的是：在现行政策预测下，英国距离目标一致路径有多大 delivery gap，以及这个 gap 应该如何通过时间、部门和不确定性来解释。",
        ),
        (
            "Chapter 1: Introduction",
            "The Introduction defines the policy problem, narrows the scope and presents the main RQ and three sub-questions. Its job is to show that the dissertation is not asking whether net zero is desirable, but whether current-policy projections look aligned with target-consistent benchmarks.",
            "第一章负责界定问题和研究范围。重点不是讨论净零目标好不好，而是判断英国当前政策路径是否和目标一致路径相符，以及为什么可能不相符。",
        ),
        (
            "Chapter 2: Literature Review",
            "The Literature Review is used as an analytical bridge. It connects target credibility, scenario/pathway choice, sectoral coupling, uncertainty and negative-emissions governance to the dissertation's research design.",
            "第二章不是泛泛综述所有净零文献，而是把文献变成分析框架：目标可信度、路径选择、部门联动、不确定性和负排放治理如何支撑本文的方法。",
        ),
        (
            "Chapter 3: Data and Methods",
            "The Methods chapter explains the evidence hierarchy: DESNZ as current-policy baseline, CCC as target-consistent benchmark, historical statistics as delivered-rate anchors, and NESO as external-condition context. It also states the boundaries: no forecast, no least-cost model, no causal proof from linkage scores.",
            "第三章说明数据角色：DESNZ 是当前政策基线，CCC 是目标一致基准，历史统计用于现实性锚点，NESO 用于外部条件解释。同时明确限制：不是预测模型、不是最低成本模型、部门联系不是因果证明。",
        ),
        (
            "Chapter 4: Results",
            f"The Results chapter contains the strongest empirical contribution. It reports the {fmt(GAP_2050)} MtCO2e 2050 DESNZ-CCC7 gap, the {fmt(CB6_GAP)} MtCO2e official CB6 gap, the DESNZ post-2035 slowdown, residual-sector concentration, historical delivered-rate checks, sector uncertainty bands, linkage-weighted risk and the 2x2 scenario matrix.",
            f"第四章是论文最核心的实证部分：包括 2050 年约 {fmt(GAP_2050)} MtCO2e 的 DESNZ-CCC7 差距，约 {fmt(CB6_GAP)} MtCO2e 的 CB6 官方差距，DESNZ 2035 年后的减速，残余部门集中度，历史减排速度对照，部门不确定性，联动风险，以及 2x2 情景矩阵。",
        ),
        (
            "Chapter 5: Discussion",
            "The Discussion interprets the results rather than repeating them. It argues that the delivery gap is a governance and implementation problem: buildings and transport depend on electrification and clean power; industry depends more on hydrogen/CCUS and cluster delivery; IAS/agriculture/land-related emissions are harder residual categories.",
            "第五章把结果转化为解释：delivery gap 不只是数字差距，而是治理和执行问题。建筑和交通依赖电气化与清洁电力，工业更依赖氢能、CCUS 和产业集群，农业、IAS、土地相关排放则更接近难减排残余。",
        ),
        (
            "Chapter 6: Conclusion",
            "The Conclusion directly answers the research questions and states what the dissertation changes in understanding: the UK net-zero gap should be read as a timed, sectorally concentrated and uncertainty-conditioned delivery problem, not as a single 2050 number.",
            "第六章直接回答研究问题，并总结本文贡献：英国净零 gap 应被理解为带有时间性、部门集中性和不确定性条件的 delivery problem，而不是一个孤立的 2050 端点数字。",
        ),
    ]
    doc.add_heading("1. Thesis Logic, Chapter by Chapter", level=1)
    add_table(
        doc,
        ["Part", "English explanation", "中文说明"],
        bilingual_rows,
        widths=[1900, 3780, 3680],
        font_size=8.8,
    )

    doc.add_heading("2. How the Three Research Questions Work Together", level=1)
    rq_rows = [
        [
            "RQ1",
            "What is the aggregate baseline-benchmark gap?",
            "This gives the dissertation its empirical foundation: DESNZ is materially above CCC7, and the gap is visible before 2050.",
            "先证明总体差距存在，而且不是 2050 年才突然出现。",
        ],
        [
            "RQ2",
            "Which sectors explain the residual gap?",
            "This prevents the dissertation from staying at economy-wide level. It identifies residual-sector concentration and cross-sector dependency.",
            "进一步说明差距来自哪些部门，以及这些部门之间是否存在转型依赖。",
        ],
        [
            "RQ3",
            "How do uncertainty and conditions change interpretation?",
            "This adds judgement: target achievement depends on domestic delivery, external conditions, near-miss cases and offsets/removals.",
            "加入不确定性判断：能否接近目标取决于国内执行、国际条件、near miss 和负排放/抵消。",
        ],
    ]
    add_table(doc, ["RQ", "Question role", "English explanation", "中文说明"], rq_rows, widths=[700, 2300, 3350, 3010], font_size=8.8)

    doc.add_heading("3. Key Results You Need to Be Able to Explain", level=1)
    key_rows = [
        [
            "2050 DESNZ-CCC7 gap",
            f"{fmt(GAP_2050)} MtCO2e including IAS.",
            "This is the headline evidence that the current-policy baseline is not target-consistent.",
            "这是最核心的 headline 数字，说明当前政策路径离目标一致路径很远。",
        ],
        [
            "Sixth Carbon Budget gap",
            f"{fmt(CB6_GAP)} MtCO2e official projected gap.",
            "This moves the argument from a distant endpoint to near-term statutory governance.",
            "说明问题不是 2050 年才发生，而是在 CB6 阶段已经出现治理压力。",
        ],
        [
            "Post-2035 reduction-rate divergence",
            f"DESNZ around {fmt(DESNZ_2035_2050_RATE)} MtCO2e/year vs CCC7 around {fmt(CCC7_2035_2050_RATE)} MtCO2e/year.",
            "This explains the momentum loss in the current-policy baseline.",
            "解释 DESNZ 为什么会落后：2035 年后减排速度几乎放缓。",
        ],
        [
            "Sectoral residual concentration",
            "Buildings/product uses, transport, industry, electricity, agriculture and IAS dominate the residual profile.",
            "This shows where the delivery problem sits in the economy.",
            "说明 gap 不是平均分布，而集中在若干关键残余部门。",
        ],
        [
            "Sector linkage",
            "Buildings, transport and electricity form a coupled electrification block; industry has a hydrogen/CCUS profile; agriculture/IAS/land are harder residual categories.",
            "This turns the sector analysis into a system-risk interpretation.",
            "把部门分析从简单排名提升为系统依赖和残余风险分析。",
        ],
        [
            "2x2 scenario matrix",
            "Aligned transition can become near-miss under favourable assumptions; weak/delayed domestic delivery leaves large residuals.",
            "This is the core RQ3 uncertainty result.",
            "这是 RQ3 的核心：国内执行强弱会显著影响 residual 结果，但仍需谨慎表述为 bounded scenario design。",
        ],
    ]
    add_table(doc, ["Result", "Number / finding", "English meaning", "中文理解"], key_rows, widths=[1900, 2600, 2850, 2010], font_size=8.6)

    doc.add_heading("4. Defence Points: What You Should Say If Challenged", level=1)
    defence_rows = [
        [
            "Is this a forecast?",
            "No. The dissertation compares official current-policy and target-consistent pathways. It diagnoses alignment and delivery risk; it does not forecast the actual future.",
            "不是预测。本文比较官方当前政策路径与目标一致路径，诊断一致性和 delivery risk。",
        ],
        [
            "Are CCC pathways treated as truth?",
            "No. CCC is used as the primary target-consistent benchmark because it is policy-relevant and transparent, while CCC6 is retained as sensitivity context.",
            "不是把 CCC 当作绝对真理，而是作为目标一致基准，且保留 CCC6 作为稳健性参考。",
        ],
        [
            "Do linkage scores prove causality?",
            "No. They are dependency evidence, not causal attribution. Their value is to prevent isolated sector-silo interpretation.",
            "不证明因果。它们是 dependency evidence，帮助避免把部门孤立看待。",
        ],
        [
            "Why use historical delivered rates?",
            "They give a plausibility anchor. They do not imply that past reductions can simply be repeated, but they make future pathway speed easier to interpret.",
            "历史减排速度是现实性锚点，不代表过去可直接重复，但可以帮助解释未来路径速度是否合理。",
        ],
        [
            "Why include offsets/removals?",
            "Offsets/removals are treated as post-processing sensitivity, not as a substitute for gross mitigation.",
            "负排放/抵消是后处理敏感性，不是替代实际减排。",
        ],
    ]
    add_table(doc, ["Challenge", "English answer", "中文提示"], defence_rows, widths=[2100, 4400, 2860], font_size=8.8)

    doc.add_heading("5. Final Revision Priorities After the Meeting", level=1)
    add_bullets(
        doc,
        [
            "Use Neil's feedback to decide whether the sector-linkage scoring should stay in the main Results or partly move to the appendix.",
            "Tighten the Discussion so it interprets the Results without repeating every number.",
            "Sharpen the Literature Review where it explains pathway credibility, historical realism and sectoral implementation risk.",
            "Check that every major conclusion directly answers the Main RQ, RQ1, RQ2 or RQ3.",
            "Run the final submission audit again: word count, references, appendix consistency, metadata, figure clarity and Harvard style.",
        ],
    )

    doc.add_heading("6. Figure-by-Figure Speaking Logic", level=1)
    fig_rows = [
        [
            "Figures 4.1-4.2",
            "Use these to show the basic pathway divergence. Figure 4.1 shows DESNZ and CCC pathways; Figure 4.2 turns that into annual gap values.",
            "这两张图负责证明最基本的路径分化：一张看路径，一张看年度 gap。",
        ],
        [
            "Figure 4.3 / 4.11 / 4.12",
            "Use these to explain timing and historical realism. The key point is the DESNZ post-2035 slowdown relative to CCC7 and historical delivered-rate scales.",
            "这组图说明时间性和现实性：DESNZ 2035 年后减速，而 CCC7 仍需要持续减排。",
        ],
        [
            "Figures 4.4-4.5",
            "Use these for the central residual-sector story. They show which sectors remain large in 2050 and how their trajectories differ.",
            "这组图回答哪些部门留下最多残余排放，以及它们的轨迹是否相同。",
        ],
        [
            "Figure 4.13",
            "Use this to show that sector analysis cannot rely only on central estimates. Some sectors are implementation-sensitive; others remain residual even under better assumptions.",
            "这张图说明部门分析必须加入不确定性，不能只看 central estimate。",
        ],
        [
            "Figures 4.15-4.17",
            "Use these to explain sector linkages. The correct wording is dependency evidence, not causal attribution.",
            "这三张图说明部门之间的转型依赖，但不能说是因果证明。",
        ],
        [
            "Figures 4.18-4.20",
            "Use these for RQ3. They show that near-miss cases mainly appear under strong domestic delivery, while weak delivery leaves large residuals even with favourable external conditions.",
            "这组图是 RQ3 核心：强国内执行才可能 near miss，弱执行会留下大 residual。",
        ],
    ]
    add_table(doc, ["Figures", "English explanation", "中文说明"], fig_rows, widths=[1600, 4300, 3460], font_size=8.7)

    doc.add_heading("7. The Dissertation's Intellectual Development", level=1)
    development_rows = [
        [
            "Initial stage",
            "The work began as a DESNZ-versus-CCC comparison. That was useful, but it risked becoming a single gap number.",
            "最开始只是 DESNZ vs CCC 差距比较，虽然有用，但容易变成一个单薄数字。",
        ],
        [
            "Deepening stage",
            "Carbon-budget timing, CCC benchmark cleaning, sectoral residuals and NESO external-context analysis were added to make the gap more policy-relevant.",
            "之后加入 CB6 时间点、CCC 清洗、部门残余和 NESO 外部条件，让 gap 更有政策解释力。",
        ],
        [
            "First-draft stage",
            "Historical delivered-rate benchmarks, sectoral uncertainty bands, linkage scoring and the 2x2 matrix turned the work into a coherent delivery-risk diagnosis.",
            "再后来加入历史减排速度、部门不确定性、部门联动和 2x2 矩阵，使论文形成完整 delivery-risk diagnosis。",
        ],
        [
            "Current stage",
            "The latest draft focuses on caution and integration: no overclaiming, no causal wording for linkage, no overemphasis on far-delayed target years, and clearer references and appendices.",
            "最新版重点是谨慎和整合：不夸大、不把 linkage 说成因果、不强调几百年后达标，并统一引用和附录。",
        ],
    ]
    add_table(doc, ["Stage", "English explanation", "中文说明"], development_rows, widths=[1600, 4300, 3460], font_size=8.7)

    doc.add_heading("8. What You Must Not Overclaim", level=1)
    caution_rows = [
        [
            "Do not say: DESNZ proves the UK will fail.",
            "Say instead: DESNZ current-policy projections are not aligned with the CCC target-consistent benchmark under the current evidence base.",
            "不要说 DESNZ 证明英国一定失败；要说当前政策预测与 CCC 目标一致基准不一致。",
        ],
        [
            "Do not say: Domestic delivery is universally the strongest determinant.",
            "Say instead: within the bounded scenario design, domestic delivery produces the larger change in residual emissions.",
            "不要泛化说 domestic delivery 是最强决定因素；要限定在本文情景设计内。",
        ],
        [
            "Do not say: Linkage scores prove sector causality.",
            "Say instead: linkage scores show shared dependency structures and help interpret residual-risk concentration.",
            "不要说 linkage scores 证明因果；要说它们展示共同依赖结构。",
        ],
        [
            "Do not say: Offsets solve the gap.",
            "Say instead: offsets/removals can improve boundary cases, but delayed or weak delivery still requires gross mitigation.",
            "不要说抵消解决 gap；要说抵消只改善边界案例，无法替代实际减排。",
        ],
    ]
    add_table(doc, ["Avoid", "Better wording", "中文提示"], caution_rows, widths=[2850, 4050, 2460], font_size=8.7)

    doc.save(GUIDE)


def main():
    make_evidence_pack()
    make_oral_script()
    make_guide()
    print(EVIDENCE_PACK)
    print(ORAL_SCRIPT)
    print(GUIDE)


if __name__ == "__main__":
    main()
