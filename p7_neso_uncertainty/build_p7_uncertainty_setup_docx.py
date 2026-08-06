# -*- coding: utf-8 -*-
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"E:\UCL Final Essay")
P45 = ROOT / "p4_p5_local_reproduction"
P6 = ROOT / "p6_sector_analysis"
P7 = ROOT / "p7_neso_uncertainty"
TABLE_DIR = P7 / "tables"
OUT_DIR = P7 / "documents"
OUT_DIR.mkdir(parents=True, exist_ok=True)
TABLE_DIR.mkdir(parents=True, exist_ok=True)

OUT_PATH = OUT_DIR / "P7_Uncertainty_Setup_2x2_Matrix_Dissertation_Ready.docx"

MATRIX_CSV = TABLE_DIR / "p7_uncertainty_2x2_scenario_matrix.csv"
DRIVERS_CSV = TABLE_DIR / "p7_uncertainty_driver_table.csv"
ANCHORS_CSV = TABLE_DIR / "p7_uncertainty_quantitative_anchors.csv"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 89, 89)
INK = RGBColor(22, 22, 22)
LIGHT_BLUE = "E8F1FB"
LIGHT_GRAY = "F2F4F7"
LIGHT_GOLD = "FFF4DE"
LIGHT_GREEN = "EAF5EF"
CONTENT_WIDTH_DXA = 9360


def fmt(x, digits=1):
    if pd.isna(x):
        return "n/a"
    return f"{float(x):,.{digits}f}"


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, color in [
        ("Heading 1", 16, DARK_BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11.5, DARK_BLUE),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True

    header = section.header.paragraphs[0]
    header.text = "P7 uncertainty setup | 2x2 matrix"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=8.8, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "Draft dissertation support note"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer.runs:
        set_run_font(run, size=8.8, color=MUTED)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=100, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for i, width in enumerate(widths):
            tc_pr = row.cells[i]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    for run in p.runs:
        set_run_font(
            run,
            size=16 if level == 1 else 13 if level == 2 else 11.5,
            color=DARK_BLUE if level == 1 else BLUE,
            bold=True,
        )
    return p


def add_para(doc, text, size=10.5, color=INK, bold=False, italic=False, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.rows[0].cells[0]
    shade_cell(cell, fill)
    set_cell_margins(cell, top=110, bottom=110, start=130, end=130)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(label)
    set_run_font(r1, size=10.3, color=DARK_BLUE, bold=True)
    r2 = p.add_run(" " + text)
    set_run_font(r2, size=10.3, color=INK)
    doc.add_paragraph()


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=8.6, color=MUTED, italic=True)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY, font_size=7.8, last_col_left=True):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, header_fill)
        set_cell_margins(cell, top=85, bottom=85, start=90, end=90)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(h))
        set_run_font(r, size=font_size, color=DARK_BLUE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_margins(cells[i], top=75, bottom=75, start=90, end=90)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if last_col_left and i == len(row) - 1 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(val))
            set_run_font(r, size=font_size, color=INK)
    doc.add_paragraph()
    return table


def load_anchors():
    p5 = pd.read_csv(P45 / "tables" / "p5_cleaned_benchmark_year_gap_metrics.csv")
    p4 = pd.read_csv(P45 / "tables" / "p4_final_metrics_table.csv")
    cb6 = pd.read_csv(P45 / "tables" / "p4_p5_deepened_cb6_linkage_metrics.csv")
    rates = pd.read_csv(P45 / "tables" / "p4_p5_deepened_trajectory_reduction_rates.csv")
    p6_rank = pd.read_csv(P6 / "tables" / "p6_desnz_2050_residual_emissions_ranking.csv")
    neso = pd.read_csv(P7 / "tables" / "p7_desnz_ccc_neso_2050_emissions_comparison.csv")
    return p5, p4, cb6, rates, p6_rank, neso


def build_tables():
    matrix_rows = [
        {
            "scenario_id": "U1",
            "scenario_name": "Aligned transition",
            "uk_policy_delivery": "High: domestic policies are delivered on time with credible sector implementation.",
            "external_conditions": "Supportive: technology learning, costs, supply chains and international climate action are favourable.",
            "expected_gap_interpretation": "Smallest residual gap; closest to target-consistent transition pathways.",
            "dissertation_use": "Optimistic boundary; helps interpret what strong delivery would need to resemble.",
        },
        {
            "scenario_id": "U2",
            "scenario_name": "Domestic delivery under external constraints",
            "uk_policy_delivery": "High: UK implementation is credible and timely.",
            "external_conditions": "Constrained: slower global technology learning, higher costs or weaker supply chains.",
            "expected_gap_interpretation": "Moderate residual gap; domestic delivery helps, but external constraints slow system change.",
            "dissertation_use": "Tests whether strong domestic policy can offset difficult external conditions.",
        },
        {
            "scenario_id": "U3",
            "scenario_name": "Weak delivery despite supportive conditions",
            "uk_policy_delivery": "Low: delayed or partial implementation, weak sector delivery and credibility gaps.",
            "external_conditions": "Supportive: technology and international conditions are favourable.",
            "expected_gap_interpretation": "Large residual gap; favourable external conditions cannot fully compensate for weak UK delivery.",
            "dissertation_use": "Policy credibility case; links directly to the DESNZ current-policy gap.",
        },
        {
            "scenario_id": "U4",
            "scenario_name": "Delayed transition risk",
            "uk_policy_delivery": "Low: policies are delayed, incomplete or weakly implemented.",
            "external_conditions": "Constrained: higher costs, slower learning and weaker international action.",
            "expected_gap_interpretation": "Largest residual gap; closest to lower-delivery pathways such as NESO Falling Behind.",
            "dissertation_use": "Risk boundary for discussion; shows why current-policy gaps should not be interpreted as harmless.",
        },
    ]
    matrix = pd.DataFrame(matrix_rows)
    matrix.to_csv(MATRIX_CSV, index=False)

    driver_rows = [
        {
            "driver": "UK policy delivery and credibility",
            "why_it_matters": "Determines whether announced sector policies translate into emissions reductions on schedule.",
            "evidence_link": "P4 CB6 gap; P5 2030s divergence; CCC progress evidence; Rogelj et al. credibility framing.",
            "affected_results": "CB6 gap, annual DESNZ-CCC gap, sectoral residuals.",
            "planned_treatment": "Core axis in 2x2 matrix; discuss qualitatively, with DESNZ current-policy as lower-delivery evidence.",
        },
        {
            "driver": "International action, supply chains and technology costs",
            "why_it_matters": "Affects technology deployment costs, availability and learning rates for electrification, hydrogen, CCS and removals.",
            "evidence_link": "NESO pathways; Waisman et al. pathway-design framing; CCC supplementary analysis.",
            "affected_results": "Feasibility interpretation of CCC7/NESO target-consistent pathways.",
            "planned_treatment": "Second axis in 2x2 matrix; use NESO indicators as supporting context rather than full modelling.",
        },
        {
            "driver": "Power-sector decarbonisation and flexibility",
            "why_it_matters": "Electrification of transport, buildings and industry requires low-carbon electricity and system flexibility.",
            "evidence_link": "NESO power intensity, storage and hydrogen dispatchable capacity indicators; P6 electricity residual.",
            "affected_results": "Interpretation of electricity supply residuals and feasibility of electrification-heavy pathways.",
            "planned_treatment": "Use selected NESO indicators in P7 note; keep detailed power-system modelling outside scope.",
        },
        {
            "driver": "Transport and heat electrification",
            "why_it_matters": "P6 identifies buildings/product uses and domestic transport as leading residual sectors.",
            "evidence_link": "P6 sector ranking; NESO EV stock and heat-pump electricity demand indicators.",
            "affected_results": "Sectoral interpretation of residual emissions and policy-delivery challenge.",
            "planned_treatment": "Discuss as sectoral feasibility context, not as a new sector model.",
        },
        {
            "driver": "Demand reduction and behaviour",
            "why_it_matters": "Lower demand can reduce reliance on difficult technology deployment and residual emissions.",
            "evidence_link": "Barrett et al. demand-side literature; P6 sector concentration.",
            "affected_results": "Magnitude and difficulty of residual emissions in transport/buildings/industry.",
            "planned_treatment": "Discuss in Literature Review and uncertainty section; do not quantify separately unless needed.",
        },
        {
            "driver": "Accounting scope and removals",
            "why_it_matters": "Including IAS, excluding IAS, LULUCF and engineered removals affect numerical gap size and interpretation.",
            "evidence_link": "P4 accounting note; P5 CCC7 near-net-zero value; P6 LULUCF/IAS caveats.",
            "affected_results": "2050 gap, CB6 comparison, sector alignment caveats.",
            "planned_treatment": "Methods caveat and sensitivity anchor; not a main uncertainty axis.",
        },
    ]
    drivers = pd.DataFrame(driver_rows)
    drivers.to_csv(DRIVERS_CSV, index=False)

    p5, p4, cb6, rates, p6_rank, neso = load_anchors()
    p5_2050 = p5.loc[p5["year"] == 2050].iloc[0]
    p5_2030 = p5.loc[p5["year"] == 2030].iloc[0]
    p5_2035 = p5.loc[p5["year"] == 2035].iloc[0]
    cb6_gap = cb6.loc[cb6["metric"] == "DESNZ official emissions minus CCC7 sum", "value_MtCO2e"].iloc[0]
    official_cb6 = p4.loc[p4["metric_id"] == "CB6_official_gap", "rounded_value_for_writing"].iloc[0]
    excl_cb6 = p4.loc[p4["metric_id"] == "CB6_excl_IAS_gap", "rounded_value_for_writing"].iloc[0]
    post2035 = rates.loc[rates["period"] == "2035-2050"].iloc[0]
    top3 = p6_rank.sort_values("rank_2050_residual").head(3)
    top3_sum = top3["2050"].sum()
    top3_share = top3_sum / p6_rank["2050"].sum() * 100
    neso_falling = neso.loc[neso["pathway_or_case"] == "Falling Behind", "value_2050_MtCO2e"].iloc[0]
    neso_ht = neso.loc[neso["pathway_or_case"] == "Holistic Transition", "value_2050_MtCO2e"].iloc[0]

    anchor_rows = [
        {
            "anchor": "P5 2050 DESNZ-CCC7 gap",
            "value": f"{fmt(p5_2050['gap_inc_IAS_DESNZ_minus_CCC7_MtCO2e'])} MtCO2e",
            "interpretation": "Main aggregate benchmark gap; should remain the headline quantitative result.",
        },
        {
            "anchor": "P5 interim-year divergence",
            "value": f"{fmt(p5_2030['gap_inc_IAS_DESNZ_minus_CCC7_MtCO2e'])} MtCO2e in 2030; {fmt(p5_2035['gap_inc_IAS_DESNZ_minus_CCC7_MtCO2e'])} MtCO2e in 2035",
            "interpretation": "Gap is visible before 2050 and becomes material around the 2030s.",
        },
        {
            "anchor": "CB6 linkage",
            "value": f"{cb6_gap} MtCO2e DESNZ above CCC7 over 2033-2037",
            "interpretation": "Connects P5 benchmark divergence to a policy-relevant carbon-budget period.",
        },
        {
            "anchor": "Accounting sensitivity",
            "value": f"{official_cb6} official CB6 gap; {excl_cb6} excluding-IAS caveat",
            "interpretation": "Accounting scope changes the number but not the direction of the conclusion.",
        },
        {
            "anchor": "Post-2035 reduction-rate contrast",
            "value": f"DESNZ {fmt(post2035['DESNZ_avg_annual_reduction'])} vs CCC7 {fmt(post2035['CCC7_avg_annual_reduction'])} MtCO2e/year",
            "interpretation": "DESNZ flattens after 2035 while CCC7 continues reducing much faster.",
        },
        {
            "anchor": "P6 sector concentration",
            "value": f"Top three residual sectors: {fmt(top3_sum)} MtCO2e, {fmt(top3_share)}% of DESNZ 2050 total",
            "interpretation": "Residual baseline is concentrated in buildings/product uses, domestic transport and industry.",
        },
        {
            "anchor": "NESO external context",
            "value": f"Holistic Transition {fmt(neso_ht)} MtCO2e; Falling Behind {fmt(neso_falling)} MtCO2e in 2050",
            "interpretation": "Target-consistent NESO pathways reach near net zero; Falling Behind remains materially above net zero.",
        },
    ]
    anchors = pd.DataFrame(anchor_rows)
    anchors.to_csv(ANCHORS_CSV, index=False)
    return matrix, drivers, anchors


def build_doc():
    matrix, drivers, anchors = build_tables()

    doc = Document()
    configure_doc(doc)

    add_heading(doc, "P7 Uncertainty Setup: 2x2 Matrix and Sensitivity Anchors", 1)
    add_para(
        doc,
        "This draft note turns the agreed uncertainty idea into a dissertation-ready framework. It does not add a new energy-system model. Instead, it uses the measured DESNZ-CCC gap, the DESNZ sectoral residual diagnosis and the NESO supporting comparison to structure how uncertainty should be discussed.",
        size=9.8,
        color=MUTED,
    )
    add_callout(
        doc,
        "Working decision:",
        "Uncertainty should be treated as a bounded interpretation framework. The dissertation can use a 2x2 matrix to organise policy-delivery and external-condition uncertainty, while keeping CCC7 as the main benchmark and NESO as supporting context.",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "1. Role in the Dissertation", 2)
    add_para(
        doc,
        "P5 establishes the size and timing of the benchmark gap, P6 identifies where residual emissions remain concentrated in the DESNZ projection, and P7/P8 explain how strongly those results should be interpreted under different delivery and external-condition assumptions. The uncertainty framework therefore sits after the empirical Results, as a bridge into Discussion.",
    )
    add_table(
        doc,
        ["Evidence block", "What it provides", "How uncertainty uses it"],
        [
            ["P5 benchmark comparison", "DESNZ current-policy baseline is far above CCC7 by 2050 and diverges during the 2030s.", "Provides the quantitative gap that uncertainty needs to interpret."],
            ["P6 sectoral diagnosis", "Residual emissions are concentrated in buildings/product uses, domestic transport and industry.", "Identifies which sectors are most exposed to delivery uncertainty."],
            ["P7 NESO comparison", "Target-consistent NESO pathways reach near net zero, while Falling Behind remains materially above net zero.", "Provides external scenario context for delivery and technology uncertainty."],
        ],
        [1900, 3900, 3560],
        font_size=7.7,
        header_fill=LIGHT_GRAY,
    )

    add_heading(doc, "2. Proposed 2x2 Scenario Matrix", 2)
    add_para(
        doc,
        "The matrix uses two axes. The first axis is UK policy delivery: whether domestic policies are implemented credibly and on time. The second axis is external and technology-cost conditions: whether international action, supply chains and technology learning support or constrain the UK transition. This structure keeps the uncertainty treatment simple while still connecting directly to the research question.",
    )
    matrix_rows = matrix[
        [
            "scenario_id",
            "scenario_name",
            "uk_policy_delivery",
            "external_conditions",
            "expected_gap_interpretation",
        ]
    ].values.tolist()
    add_table(
        doc,
        ["Scenario", "Name", "UK policy delivery", "External conditions", "Expected gap interpretation"],
        matrix_rows,
        [650, 1600, 2650, 2650, 1810],
        font_size=6.9,
        header_fill=LIGHT_GREEN,
    )
    add_caption(doc, "Table 1. Proposed 2x2 uncertainty matrix for policy-delivery and external-condition uncertainty.")

    add_heading(doc, "3. Uncertainty Driver Table", 2)
    add_para(
        doc,
        "The driver table specifies what each uncertainty factor affects and how it should be treated before the final dissertation. The purpose is to prevent uncontrolled expansion: not every uncertainty becomes a new quantitative model.",
    )
    driver_rows = drivers[
        ["driver", "why_it_matters", "evidence_link", "planned_treatment"]
    ].values.tolist()
    add_table(
        doc,
        ["Driver", "Why it matters", "Evidence link", "Planned treatment"],
        driver_rows,
        [1750, 2550, 2650, 2410],
        font_size=6.9,
        header_fill=LIGHT_GRAY,
    )
    add_caption(doc, "Table 2. Uncertainty drivers and planned treatment.")

    add_heading(doc, "4. Quantitative Anchors Already Available", 2)
    add_para(
        doc,
        "Although the 2x2 framework is not a new quantitative model, it should be anchored in the quantitative results already produced. These anchors make the uncertainty discussion empirical rather than purely conceptual.",
    )
    anchor_rows = anchors[["anchor", "value", "interpretation"]].values.tolist()
    add_table(
        doc,
        ["Anchor", "Value", "Interpretation"],
        anchor_rows,
        [2350, 2850, 4160],
        font_size=7.2,
        header_fill=LIGHT_GOLD,
    )
    add_caption(doc, "Table 3. Quantitative anchors for the uncertainty discussion.")

    add_heading(doc, "5. How This Should Be Written in the Dissertation", 2)
    add_para(
        doc,
        "The dissertation should avoid claiming that the uncertainty framework produces a full probabilistic range for the emissions gap. A stronger and more defensible claim is that the framework identifies which assumptions would make the measured gap easier or harder to close. In this framing, DESNZ represents the current-policy baseline, CCC7 represents the main target-consistent benchmark, NESO provides external transition-context pathways, and the 2x2 matrix structures the interpretation of delivery and external-condition uncertainty.",
    )
    add_para(
        doc,
        "This treatment also keeps RQ3 focused. The question is not whether every uncertainty can be modelled quantitatively. The question is how sensitive the interpretation of the projected gap is to policy delivery, technology conditions, sectoral transition pace and accounting scope. The dissertation can answer this through a bounded matrix, selected quantitative anchors and careful discussion of where uncertainty is greatest.",
    )

    add_callout(
        doc,
        "Recommended wording:",
        "The uncertainty analysis is used to interpret the robustness and policy relevance of the measured DESNZ-CCC gap, rather than to replace the benchmark comparison with a new scenario model.",
        fill=LIGHT_GOLD,
    )

    add_heading(doc, "Source Note", 2)
    add_para(
        doc,
        "Primary evidence base: P4 DESNZ baseline and accounting metrics; P5 CCC benchmark comparison; P6 DESNZ sectoral residual-emissions diagnosis; P7 NESO FES 2025 compact indicator extraction; selected literature on pathway feasibility, policy credibility, demand reduction and uncertainty in energy-system modelling.",
        size=9.2,
        color=MUTED,
    )

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_doc())
