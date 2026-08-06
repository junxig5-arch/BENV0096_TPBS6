# -*- coding: utf-8 -*-
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"E:\UCL Final Essay")
P7 = ROOT / "p7_neso_uncertainty"
P45 = ROOT / "p4_p5_local_reproduction"
TABLE_DIR = P7 / "tables"
OUT_DIR = P7 / "documents"
OUT_DIR.mkdir(parents=True, exist_ok=True)

OUT_PATH = OUT_DIR / "P7_NESO_Short_Comparison_Note_Dissertation_Ready.docx"
COMPARISON_CSV = TABLE_DIR / "p7_desnz_ccc_neso_2050_emissions_comparison.csv"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 89, 89)
INK = RGBColor(22, 22, 22)
LIGHT_BLUE = "E8F1FB"
LIGHT_GRAY = "F2F4F7"
LIGHT_GOLD = "FFF4DE"
LIGHT_GREEN = "EAF5EF"
CONTENT_WIDTH_DXA = 9360


def fmt(x, digits=1):
    if pd.isna(x):
        return "n/a"
    return f"{float(x):,.{digits}f}"


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, color in [
        ("Heading 1", 16, DARK_BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11.5, DARK_BLUE),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True

    header = section.header.paragraphs[0]
    header.text = "P7 NESO comparison note"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=8.8, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "Draft dissertation support note"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer.runs:
        set_run_font(run, size=8.8, color=MUTED)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=100, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for i, width in enumerate(widths):
            tc_pr = row.cells[i]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    for run in p.runs:
        set_run_font(
            run,
            size=16 if level == 1 else 13 if level == 2 else 11.5,
            color=DARK_BLUE if level == 1 else BLUE,
            bold=True,
        )
    return p


def add_para(doc, text, size=10.5, color=INK, bold=False, italic=False, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.rows[0].cells[0]
    shade_cell(cell, fill)
    set_cell_margins(cell, top=110, bottom=110, start=130, end=130)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(label)
    set_run_font(r1, size=10.3, color=DARK_BLUE, bold=True)
    r2 = p.add_run(" " + text)
    set_run_font(r2, size=10.3, color=INK)
    doc.add_paragraph()


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=8.6, color=MUTED, italic=True)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY, font_size=8.0, last_col_left=True):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, header_fill)
        set_cell_margins(cell, top=85, bottom=85, start=90, end=90)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(h))
        set_run_font(r, size=font_size, color=DARK_BLUE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_margins(cells[i], top=75, bottom=75, start=90, end=90)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if last_col_left and i == len(row) - 1 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(val))
            set_run_font(r, size=font_size, color=INK)
    doc.add_paragraph()
    return table


def load_inputs():
    wide = pd.read_csv(TABLE_DIR / "p7_neso_selected_year_indicators_wide.csv")
    qc = pd.read_csv(TABLE_DIR / "p7_neso_extraction_quality_checks.csv")
    p5 = pd.read_csv(P45 / "tables" / "p5_cleaned_benchmark_year_gap_metrics.csv")
    return wide, qc, p5


def value(wide, indicator_id, pathway, year):
    row = wide[(wide["indicator_id"] == indicator_id) & (wide["pathway"] == pathway)]
    if row.empty:
        return float("nan")
    return row[str(year)].iloc[0]


def build_emissions_comparison(wide, p5):
    p5_2050 = p5[p5["year"] == 2050].iloc[0]
    rows = [
        {
            "source": "DESNZ EEP 2024",
            "pathway_or_case": "Current-policy baseline including IAS",
            "role_in_dissertation": "Main baseline",
            "value_2050_MtCO2e": p5_2050["DESNZ_EEP_2024_inc_IAS_MtCO2e"],
            "interpretation": "Current-policy residual emissions remain high by 2050.",
        },
        {
            "source": "CCC Seventh Carbon Budget",
            "pathway_or_case": "Balanced Pathway",
            "role_in_dissertation": "Main target-consistent benchmark",
            "value_2050_MtCO2e": p5_2050["CCC7_Balanced_Pathway_MtCO2e"],
            "interpretation": "Approximately net zero / net removals by 2050.",
        },
    ]
    for pathway in ["Holistic Transition", "Electric Engagement", "Hydrogen Evolution", "Falling Behind"]:
        emissions = value(wide, "neso_total_emissions", pathway, 2050)
        if pathway == "Falling Behind":
            interpretation = "Remains materially above net zero; useful as lower-delivery context."
        else:
            interpretation = "Reaches approximately net zero; useful as external feasibility context."
        rows.append(
            {
                "source": "NESO FES 2025",
                "pathway_or_case": pathway,
                "role_in_dissertation": "Supporting external scenario context",
                "value_2050_MtCO2e": emissions,
                "interpretation": interpretation,
            }
        )
    df = pd.DataFrame(rows)
    df.to_csv(COMPARISON_CSV, index=False)
    return df


def build_doc():
    wide, qc, p5 = load_inputs()
    comparison = build_emissions_comparison(wide, p5)

    doc = Document()
    configure_doc(doc)

    add_heading(doc, "P7 NESO Short Comparison Note", 1)
    add_para(
        doc,
        "This note converts the P7 NESO extraction into dissertation-ready interpretation. It is designed as a supporting Results/Discussion note rather than a replacement for the P5 CCC benchmark comparison.",
        size=9.8,
        color=MUTED,
    )
    add_callout(
        doc,
        "Role decision:",
        "NESO Future Energy Scenarios 2025 should be used as supporting external modelling context. CCC7 remains the main target-consistent benchmark, while NESO helps interpret pathway feasibility, electrification, power-sector decarbonisation, heat, transport and flexibility assumptions.",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "Data Extraction Check", 2)
    pass_count = int((qc["status"] == "PASS").sum())
    add_para(
        doc,
        f"The local P7 notebook extracted seven compact indicators from six NESO data tables: WS2, ED1, ES1, ED5, ED7 and FLX1. All core extraction checks passed ({pass_count} PASS checks). The Ten Year Forecast pathway is retained where available, but is treated as optional because it does not provide complete 2050 values for the selected indicators.",
    )

    qc_rows = qc[["check", "status", "details"]].values.tolist()
    add_table(doc, ["Check", "Status", "Details"], qc_rows, [3100, 850, 5410], font_size=7.3, header_fill=LIGHT_GREEN)
    add_caption(doc, "Table 1. P7 NESO extraction quality checks.")

    add_heading(doc, "Simple DESNZ/CCC/NESO 2050 Emissions Comparison", 2)
    add_para(
        doc,
        "Table 2 provides a simple comparison of 2050 economy-wide emissions values across DESNZ, CCC7 and selected NESO FES pathways. This table should be interpreted cautiously because the sources are not identical in purpose or accounting design. Its value is not to replace the P5 benchmark gap calculation, but to show that NESO target-consistent transition pathways broadly sit close to net zero by 2050, whereas a lower-delivery NESO pathway remains materially above net zero.",
    )
    comp_rows = []
    for _, row in comparison.iterrows():
        comp_rows.append(
            [
                row["source"],
                row["pathway_or_case"],
                row["role_in_dissertation"],
                fmt(row["value_2050_MtCO2e"]),
                row["interpretation"],
            ]
        )
    add_table(
        doc,
        ["Source", "Pathway/case", "Role", "2050 MtCO2e", "Interpretation"],
        comp_rows,
        [1600, 2200, 2100, 1100, 2360],
        font_size=7.2,
        header_fill=LIGHT_GRAY,
    )
    add_caption(doc, "Table 2. Simple 2050 emissions comparison across DESNZ, CCC7 and NESO FES 2025 pathways.")

    add_heading(doc, "Selected NESO Energy-System Indicators", 2)
    add_para(
        doc,
        "The selected NESO indicators also clarify what kinds of system changes sit behind the scenario differences. The target-consistent NESO pathways combine high electricity demand, very low power-sector carbon intensity, substantial heat electrification, material storage expansion and hydrogen dispatchable capacity. Falling Behind is useful because it illustrates a lower-delivery context: emissions remain high, power-sector intensity is higher, and flexibility/hydrogen build-out is much weaker.",
    )

    indicator_labels = [
        ("neso_total_electricity_demand", "Total electricity demand", "TWh"),
        ("neso_power_intensity_excluding_beccs", "Power intensity excluding BECCS", "gCO2/kWh"),
        ("neso_electric_vehicle_stock", "Electric vehicle stock", "million vehicles"),
        ("neso_heat_pump_electricity_demand", "Heat-pump electricity demand", "TWh"),
        ("neso_storage_connection_capacity", "Storage connection capacity", "GW"),
        ("neso_hydrogen_dispatchable_capacity", "Hydrogen dispatchable capacity", "GW"),
    ]
    pathways = ["Holistic Transition", "Electric Engagement", "Hydrogen Evolution", "Falling Behind"]
    rows = []
    for indicator_id, label, unit in indicator_labels:
        row = [label, unit]
        for pathway in pathways:
            row.append(fmt(value(wide, indicator_id, pathway, 2050)))
        rows.append(row)
    add_table(
        doc,
        ["Indicator", "Unit", "Holistic Transition", "Electric Engagement", "Hydrogen Evolution", "Falling Behind"],
        rows,
        [2200, 900, 1550, 1550, 1550, 1610],
        font_size=7.2,
        header_fill=LIGHT_GOLD,
        last_col_left=False,
    )
    add_caption(doc, "Table 3. Selected NESO FES 2025 2050 indicators extracted from the local P7 notebook.")

    add_heading(doc, "Dissertation Interpretation", 2)
    add_para(
        doc,
        "The NESO comparison supports three dissertation points. First, it reinforces the P5 finding that the DESNZ current-policy baseline is not close to target-consistent pathways by 2050. Second, it provides external scenario context for P6: sectors such as buildings, transport and electricity supply are not only residual-emissions categories but also areas where transition pathways require substantial electrification, low-carbon power and flexibility. Third, it motivates the P7/P8 uncertainty framework: differences between Holistic Transition, Electric Engagement, Hydrogen Evolution and Falling Behind can be interpreted through domestic policy delivery, technology/fuel conditions and system-flexibility assumptions.",
    )
    add_para(
        doc,
        "The NESO outputs should therefore be used sparingly in the main dissertation. A short comparison table is sufficient for the Results or Discussion chapter. More detailed NESO indicators can be placed in an appendix if needed. The central analytical chain should remain DESNZ baseline, CCC7 benchmark comparison and DESNZ sectoral residual diagnosis.",
    )

    add_callout(
        doc,
        "Recommended use:",
        "Use NESO as context, not as a second benchmark hierarchy. In the dissertation, cite NESO to support pathway-feasibility and uncertainty discussion, while keeping CCC7 as the main target-consistent comparator.",
        fill=LIGHT_GOLD,
    )

    add_heading(doc, "Source Note", 2)
    add_para(
        doc,
        "Primary sources: NESO Future Energy Scenarios 2025 Data Workbook V006; DESNZ Energy and emissions projections 2024-2050 Annex A; CCC Seventh Carbon Budget dataset. NESO indicators are extracted from WS2, ED1, ES1, ED5, ED7 and FLX1 using the local P7 reproducible notebook.",
        size=9.2,
        color=MUTED,
    )

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_doc())
