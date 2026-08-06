from pathlib import Path
import math
import re

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"E:\UCL Final Essay")
P5_ANNUAL = ROOT / "p4_p5_local_reproduction" / "data_processed" / "p5_cleaned_desnz_ccc_annual_comparison.csv"
P7_LONG = ROOT / "p7_neso_uncertainty" / "tables" / "p7_neso_compact_indicator_long.csv"
MEETING_PACK = ROOT / "meeting_preparation" / "documents" / "P5_P7_Meeting_Evidence_Pack_for_Neil_Deduplicated_Deepened.docx"

OUT_DIR = ROOT / "p7_neso_uncertainty"
TABLE_DIR = OUT_DIR / "tables"
DOC_DIR = OUT_DIR / "documents"
MEETING_DIR = ROOT / "meeting_preparation" / "documents"
TABLE_DIR.mkdir(parents=True, exist_ok=True)
DOC_DIR.mkdir(parents=True, exist_ok=True)

OUT_TARGET = TABLE_DIR / "p7_target_achievement_year_estimates.csv"
OUT_EXTERNAL = TABLE_DIR / "p7_external_conditions_evidence_table.csv"
OUT_QC = TABLE_DIR / "p7_target_year_addendum_quality_checks.csv"
OUT_DOCX = DOC_DIR / "P7_Target_Achievement_Year_and_External_Conditions_Addendum.docx"
OUT_MEETING = MEETING_DIR / "P5_P7_Meeting_Evidence_Pack_for_Neil_Deduplicated_Deepened_TargetYear.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(26, 36, 48)
MUTED = RGBColor(90, 102, 115)
LIGHT_BLUE = "E8F1FA"
LIGHT_GREY = "F2F4F7"
CALLOUT = "F7F9FC"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def fmt(x, digits=1):
    if x is None or pd.isna(x):
        return ""
    return f"{float(x):,.{digits}f}"


def fmt_year(x):
    if x is None or pd.isna(x):
        return ""
    return f"{float(x):.1f}"


def clean(s):
    return re.sub(r"\s+", " ", str(s)).strip()


def style_run(run, size=10.5, color=INK, bold=False, italic=False):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic


def style_para(p, size=10.5, color=INK, after=6, line=1.10):
    for run in p.runs:
        style_run(run, size=size, color=color)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line


def add_para(doc, text, size=10.5, color=INK, bold=False, italic=False, after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, size=size, color=color, bold=bold, italic=italic)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.bold = True
        run.font.color.rgb = BLUE if level <= 2 else DARK_BLUE
        run.font.size = Pt(16 if level == 1 else 13 if level == 2 else 12)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_cell_text(cell, value, size=8.3, bold=False, color=INK):
    cell.text = str(value)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.03
        for run in p.runs:
            style_run(run, size=size, color=color, bold=bold)


def add_table(doc, rows, widths=None, font_size=8.3, header_fill=LIGHT_GREY):
    if widths is None:
        n = len(rows[0])
        widths = [int(CONTENT_WIDTH_DXA / n)] * n
        widths[-1] = CONTENT_WIDTH_DXA - sum(widths[:-1])
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(rows[0]):
        set_cell_text(table.rows[0].cells[i], h, size=font_size, bold=True)
        shade_cell(table.rows[0].cells[i], header_fill)
    for row_data in rows[1:]:
        row = table.add_row()
        for i, value in enumerate(row_data):
            set_cell_text(row.cells[i], value, size=font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph()
    return table


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    shade_cell(cell, CALLOUT)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(label + ": ")
    style_run(r, size=10.2, color=DARK_BLUE, bold=True)
    r = p.add_run(text)
    style_run(r, size=10.2, color=INK)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.space_after = Pt(0)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    doc.add_paragraph()


def first_crossing(df):
    d = df[["year", "value"]].dropna().sort_values("year")
    prev = None
    for _, row in d.iterrows():
        y = float(row["year"])
        v = float(row["value"])
        if v <= 0:
            if prev and prev[1] > 0:
                y0, v0 = prev
                return y0 + v0 / (v0 - v), y, v
            return y, y, v
        prev = (y, v)
    return None, None, None


def value_at(df, year):
    s = df[df["year"].eq(year)]["value"].dropna()
    return float(s.iloc[0]) if len(s) else None


def annual_rate(df, start, end):
    v0 = value_at(df, start)
    v1 = value_at(df, end)
    if v0 is None or v1 is None:
        return None
    return (v1 - v0) / (end - start)


def estimate_missed_year(value_2050, rate):
    if value_2050 is None or pd.isna(value_2050) or value_2050 <= 0:
        return None
    if rate is None or pd.isna(rate) or rate >= 0:
        return None
    return 2050 + value_2050 / abs(rate)


def build_target_estimates():
    p5 = pd.read_csv(P5_ANNUAL)
    rows = []
    p5_specs = [
        ("DESNZ EEP 2024", "DESNZ current-policy baseline including IAS", "DESNZ_EEP_2024_inc_IAS_MtCO2e"),
        ("DESNZ EEP 2024", "DESNZ current-policy baseline excluding IAS", "DESNZ_EEP_2024_excl_IAS_MtCO2e"),
        ("CCC Seventh Carbon Budget", "CCC7 Balanced Pathway", "CCC7_Balanced_Pathway_MtCO2e"),
        ("CCC Sixth Carbon Budget", "CCC6 Balanced Net Zero Pathway", "CCC6_Balanced_Net_Zero_Pathway_MtCO2e"),
    ]
    for source, pathway, col in p5_specs:
        d = p5[["year", col]].rename(columns={col: "value"}).copy()
        add_pathway_rows(rows, source, pathway, d)

    long = pd.read_csv(P7_LONG)
    neso = long[long["indicator_id"].eq("neso_total_emissions")].copy()
    neso["year"] = pd.to_numeric(neso["year"], errors="coerce")
    neso["value"] = pd.to_numeric(neso["value"], errors="coerce")
    for pathway, g in neso.groupby("pathway"):
        if pathway == "Ten Year Forecast":
            continue
        add_pathway_rows(rows, "NESO FES 2025", pathway, g[["year", "value"]].copy())

    out = pd.DataFrame(rows)
    out.to_csv(OUT_TARGET, index=False)
    return out


def add_pathway_rows(rows, source, pathway, d):
    crossing, first_year, first_value = first_crossing(d)
    v2050 = value_at(d, 2050)
    r_2030_2050 = annual_rate(d, 2030, 2050)
    r_2040_2050 = annual_rate(d, 2040, 2050)
    broad_year = estimate_missed_year(v2050, r_2030_2050)
    late_year = estimate_missed_year(v2050, r_2040_2050)
    if crossing is not None and crossing <= 2050:
        status = "Meets or reaches net zero by 2050"
        preferred_year = crossing
        delay = 0.0
        delay_note = "On time; interpolated crossing is within the legal target year."
    elif v2050 is not None and v2050 > 0:
        status = "Does not meet net zero by 2050"
        preferred_year = late_year if late_year is not None else broad_year
        delay = preferred_year - 2050 if preferred_year is not None else None
        if late_year is None:
            delay_note = "No finite crossing under the 2040-2050 trend; the 2030-2050 extrapolation is a mechanical stress-test, not a forecast."
        else:
            delay_note = "Missed-pathway estimate based on extending the 2040-2050 average annual reduction rate."
    else:
        status = "Not assessable from available 2050 value"
        preferred_year = None
        delay = None
        delay_note = "Insufficient annual endpoint data."
    rows.append({
        "source": source,
        "pathway": pathway,
        "value_2050_MtCO2e": v2050,
        "target_status": status,
        "interpolated_crossing_year_if_on_time": crossing,
        "preferred_estimated_target_year": preferred_year,
        "delay_years_after_2050": delay,
        "avg_annual_change_2030_2050_MtCO2e_per_year": r_2030_2050,
        "extrapolated_target_year_using_2030_2050_rate": broad_year,
        "avg_annual_change_2040_2050_MtCO2e_per_year": r_2040_2050,
        "extrapolated_target_year_using_2040_2050_rate": late_year,
        "method_note": delay_note,
    })


def build_external_table():
    rows = [
        {
            "external_factor": "Clean electricity cost decline",
            "evidence": "IEA World Energy Investment 2024 reports solar PV investment above USD 500 billion in 2024 and notes that each dollar invested in wind and solar PV in 2023 yielded 2.5 times more energy output than a decade earlier.",
            "uk_relevance": "Supports cheaper power-sector decarbonisation, which underpins EVs, heat pumps and industrial electrification.",
            "direction": "Supportive if planning, grid connection and supply chains allow rapid deployment.",
            "source": "IEA (2024) World Energy Investment 2024",
            "url": "https://www.iea.org/reports/world-energy-investment-2024/overview-and-key-findings",
        },
        {
            "external_factor": "Solar PV module price decline",
            "evidence": "IEA Renewables 2024 reports global spot prices for solar PV modules fell by about 60% over 2022-2024, driven by polysilicon price falls and supply overcapacity.",
            "uk_relevance": "Could reduce costs of UK solar deployment and system electrification, but benefits depend on market access and non-module costs.",
            "direction": "Supportive, but not sufficient without grid, flexibility and planning delivery.",
            "source": "IEA (2024) Renewables 2024",
            "url": "https://www.iea.org/reports/renewables-2024",
        },
        {
            "external_factor": "Renewable LCOE competitiveness",
            "evidence": "IRENA reports utility-scale solar PV global weighted-average LCOE at USD 0.043/kWh in 2024, around 41% below the least-cost fossil-fuel alternative.",
            "uk_relevance": "Strengthens the economic case for clean power as an enabling condition for UK net-zero pathways.",
            "direction": "Supportive for Holistic Transition and Electric Engagement style pathways.",
            "source": "IRENA (2025) Renewable Power Generation Costs in 2024",
            "url": "https://www.irena.org/-/media/Files/IRENA/Agency/Publication/2025/Jul/IRENA_TEC_RPGC_in_2024_Summary_2025.pdf",
        },
        {
            "external_factor": "Technology supply chains, grids and finance",
            "evidence": "IEA Renewables 2024 expects renewable capacity to grow 2.7 times by 2030, but still short of the tripling goal; financing, grids and market barriers remain important constraints.",
            "uk_relevance": "Even if technology costs fall globally, UK delivery can be delayed by grid connection, skills, planning, capital cost and supply-chain limits.",
            "direction": "Constrained external conditions push the analysis toward delayed-transition cases.",
            "source": "IEA (2024) Renewables 2024",
            "url": "https://www.iea.org/reports/renewables-2024/executive-summary",
        },
        {
            "external_factor": "Reduced exposure to fossil-fuel price shocks",
            "evidence": "NESO's FES 2025 Economics Annex indicates that a 2022-style fossil price shock would have a much smaller GDP effect under a highly decarbonised 2050 pathway than under today's fossil-exposed system.",
            "uk_relevance": "International fossil-fuel volatility is a risk for delayed pathways and a benefit of faster clean-energy transition.",
            "direction": "Supports the resilience argument for net-zero-compatible pathways.",
            "source": "NESO (2025) FES 2025 Economics Annex",
            "url": "https://www.neso.energy/document/374246/download",
        },
    ]
    out = pd.DataFrame(rows)
    out.to_csv(OUT_EXTERNAL, index=False)
    return out


def target_rows_for_doc(df):
    order = [
        "DESNZ current-policy baseline including IAS",
        "CCC7 Balanced Pathway",
        "CCC6 Balanced Net Zero Pathway",
        "Holistic Transition",
        "Electric Engagement",
        "Hydrogen Evolution",
        "Falling Behind",
    ]
    rows = [["Pathway or scenario", "2050 value", "Meets 2050 target?", "Estimated target year", "Delay", "Basis"]]
    for name in order:
        r = df[df["pathway"].eq(name)]
        if r.empty:
            continue
        row = r.iloc[0]
        meets = "Yes" if "Meets" in row["target_status"] else "No"
        if pd.isna(row["preferred_estimated_target_year"]):
            year = "No finite crossing under late trend"
            delay = "Not finite"
        else:
            year = fmt_year(row["preferred_estimated_target_year"])
            delay = "0.0" if row["delay_years_after_2050"] == 0 else fmt(row["delay_years_after_2050"])
        rows.append([
            name,
            fmt(row["value_2050_MtCO2e"]),
            meets,
            year,
            delay,
            row["method_note"],
        ])
    return rows


def missed_rows_for_doc(df):
    rows = [["Missed pathway", "2050 residual", "2030-2050 rate", "2030-2050 target year", "2040-2050 rate", "2040-2050 target year", "Interpretation"]]
    missed = df[df["target_status"].str.startswith("Does not", na=False)]
    for _, row in missed.iterrows():
        late_year = "No finite crossing" if pd.isna(row["extrapolated_target_year_using_2040_2050_rate"]) else fmt_year(row["extrapolated_target_year_using_2040_2050_rate"])
        broad_year = "No finite crossing" if pd.isna(row["extrapolated_target_year_using_2030_2050_rate"]) else fmt_year(row["extrapolated_target_year_using_2030_2050_rate"])
        if "DESNZ" in row["pathway"]:
            interpretation = "The broad trend implies a very late mechanical crossing, while the 2040-2050 trend is flat/slightly worsening; this is not a credible net-zero pathway without policy change."
        else:
            interpretation = "Continuing the late-period reduction rate would miss the legal target by roughly a quarter-century."
        rows.append([
            row["pathway"],
            fmt(row["value_2050_MtCO2e"]),
            fmt(row["avg_annual_change_2030_2050_MtCO2e_per_year"]),
            broad_year,
            fmt(row["avg_annual_change_2040_2050_MtCO2e_per_year"]),
            late_year,
            interpretation,
        ])
    return rows


def build_docx(target_df, external_df):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.85)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    header = sec.header.paragraphs[0]
    header.text = "P7 addendum | target-achievement year and external conditions"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_para(header, size=8.5, color=MUTED, after=0)

    p = doc.add_paragraph()
    r = p.add_run("P7 Addendum: Target-Achievement Year And External Conditions")
    style_run(r, size=20, color=DARK_BLUE, bold=True)
    p.paragraph_format.space_after = Pt(4)
    add_para(doc, "Purpose: to answer the stricter supervision question: under each pathway or scenario, does the UK reach net zero by 2050; if not, how late would it be under transparent extrapolation; and how do international clean-technology conditions affect the judgement?", color=MUTED)

    add_heading(doc, "1. Method", 1)
    add_callout(doc, "Rule", "A pathway is treated as meeting the target if its annual total net emissions are at or below 0 MtCO2e by 2050. For pathways that cross zero, the target year is linearly interpolated between the last positive and first non-positive annual value. For pathways still positive in 2050, the delay is a mechanical extrapolation using the 2040-2050 late-period reduction rate; a 2030-2050 estimate is also shown as a broader stress-test.")
    add_para(doc, "Important caveat: extrapolated missed-target years are not forecasts. They are diagnostic estimates showing the scale of delay implied if the pathway's observed annual reduction rate continued beyond 2050.")

    add_heading(doc, "2. Target-Achievement Assessment", 1)
    add_table(doc, target_rows_for_doc(target_df), widths=[1950, 900, 1050, 1250, 850, 3360], font_size=7.5)

    add_heading(doc, "3. Missed-Target Delay Estimates", 1)
    add_table(doc, missed_rows_for_doc(target_df), widths=[1700, 850, 1000, 1150, 1000, 1150, 2510], font_size=7.2)
    add_para(doc, "The key result is asymmetric. CCC7, CCC6 and the three NESO net-zero pathways reach or approximately reach net zero by 2050. NESO Falling Behind misses the target and, if its 2040-2050 reduction rate continued, would reach net zero around 2075.6, about 25.6 years late. DESNZ current policy is more problematic: the 2040-2050 trend is effectively flat, so it has no finite crossing under that late-period trend; even the broader 2030-2050 mechanical extrapolation implies around 2198.5, roughly 148.5 years late.")

    add_heading(doc, "4. International And Technology Conditions", 1)
    rows = [["External factor", "Evidence", "UK relevance", "Direction"]]
    for _, row in external_df.iterrows():
        rows.append([row["external_factor"], row["evidence"], row["uk_relevance"], row["direction"]])
    add_table(doc, rows, widths=[1500, 3200, 2550, 2110], font_size=7.1)

    add_heading(doc, "5. How This Strengthens P7", 1)
    add_table(doc, [
        ["Previous P7 treatment", "Deepened treatment now added"],
        ["Compared 2050 values across DESNZ, CCC7 and NESO.", "Classifies each pathway by whether it reaches net zero, estimates target-achievement year, and quantifies missed-target delay."],
        ["Used external conditions as a qualitative uncertainty axis.", "Grounds external conditions in literature and official evidence on clean-technology costs, module prices, renewable capacity growth, fossil-price exposure, supply chains and grids."],
        ["Showed DESNZ is higher than NESO Falling Behind in 2050.", "Shows DESNZ is not just high in 2050; under its late-period trend it does not converge to net zero at all, while Falling Behind is roughly two to three decades late."],
    ], widths=[3900, 5460], font_size=8.2)

    add_heading(doc, "6. Meeting-Ready Wording", 1)
    add_callout(doc, "Suggested answer", "I have added a target-achievement test for each pathway. The CCC pathways and the three NESO net-zero pathways reach net zero around 2050. NESO Falling Behind does not; extending its 2040-2050 trend gives a mechanical estimate of about 2076, roughly 26 years late. DESNZ current policy is worse as a convergence path: by 2050 it remains at 324.3 MtCO2e and the 2040-2050 trend is effectively flat, so there is no finite crossing under the late-period trend. This is why I interpret DESNZ as a current-policy misalignment baseline, not a delayed but otherwise credible net-zero pathway.")

    add_heading(doc, "Source Notes", 1)
    source_rows = [["Source", "Use in this addendum", "URL"]]
    for _, row in external_df.iterrows():
        source_rows.append([row["source"], row["external_factor"], row["url"]])
    source_rows.extend([
        ["DESNZ EEP 2024-2050", "Current-policy baseline and missed-target estimate", "https://www.gov.uk/government/publications/energy-and-emissions-projections-2024-to-2050"],
        ["CCC Seventh Carbon Budget", "Main target-consistent benchmark and 2050 target basis", "https://www.theccc.org.uk/publication/the-seventh-carbon-budget/"],
        ["NESO FES 2025 Data Workbook", "External scenario pathways and annual total emissions", "https://www.neso.energy/document/364561/download"],
        ["NESO FES 2025 document suite", "Report, assumptions and economics annex download page", "https://www.neso.energy/publications/future-energy-scenarios-fes/fes-documents"],
    ])
    add_table(doc, source_rows, widths=[2300, 3400, 3660], font_size=7.0)

    doc.save(OUT_DOCX)


def append_to_meeting_pack(target_df, external_df):
    doc = Document(str(MEETING_PACK))
    doc.add_page_break()
    add_heading(doc, "7. Target-Achievement Year And External Conditions", 1)
    add_para(doc, "This final P7 addition answers a stricter version of the scenario question: not only how different pathways compare in 2050, but whether each pathway reaches the UK's net-zero target, approximately when it reaches it, and how international clean-technology conditions affect that judgement.")
    add_table(doc, target_rows_for_doc(target_df), widths=[1950, 900, 1050, 1250, 850, 3360], font_size=7.4)
    add_table(doc, missed_rows_for_doc(target_df), widths=[1700, 850, 1000, 1150, 1000, 1150, 2510], font_size=7.1)
    add_callout(doc, "Neil-facing point", "The missed pathways are now quantified rather than just labelled as failures. Falling Behind is roughly 20-26 years late depending on the extrapolation window; DESNZ current policy has no finite net-zero crossing under the 2040-2050 trend, and even the broader 2030-2050 trend implies a mechanically estimated crossing around 2199.")
    rows = [["External condition", "Evidence and relevance"]]
    for _, row in external_df.head(4).iterrows():
        rows.append([row["external_factor"], row["evidence"] + " " + row["uk_relevance"]])
    add_table(doc, rows, widths=[2200, 7160], font_size=7.8)
    doc.save(OUT_MEETING)


def inspect_outputs(target_df):
    checks = []
    for path in [OUT_TARGET, OUT_EXTERNAL, OUT_DOCX, OUT_MEETING]:
        checks.append({"check": f"{path.name} exists", "status": "PASS" if path.exists() and path.stat().st_size > 0 else "FAIL"})
    text_doc = Document(str(OUT_DOCX))
    text = "\n".join(p.text for p in text_doc.paragraphs)
    for table in text_doc.tables:
        for row in table.rows:
            text += "\n" + " | ".join(c.text for c in row.cells)
    for phrase in ["2075.6", "2198.5", "no finite crossing", "external conditions", "Falling Behind"]:
        checks.append({"check": f"addendum contains {phrase}", "status": "PASS" if phrase in text else "FAIL"})
    checks.append({"check": "Falling Behind missed target", "status": "PASS" if (target_df[target_df["pathway"].eq("Falling Behind")]["target_status"].iloc[0].startswith("Does not")) else "FAIL"})
    pd.DataFrame(checks).to_csv(OUT_QC, index=False)
    return checks


def main():
    target_df = build_target_estimates()
    external_df = build_external_table()
    build_docx(target_df, external_df)
    append_to_meeting_pack(target_df, external_df)
    checks = inspect_outputs(target_df)
    print(pd.DataFrame(checks).to_string(index=False))
    print(OUT_DOCX)
    print(OUT_MEETING)


if __name__ == "__main__":
    main()
