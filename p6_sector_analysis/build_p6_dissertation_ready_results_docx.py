# -*- coding: utf-8 -*-
from pathlib import Path

import pandas as pd
from PIL import Image
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"E:\UCL Final Essay")
P6 = ROOT / "p6_sector_analysis"
TABLE_DIR = P6 / "tables"
FIG_DIR = P6 / "figures"
OUT_DIR = P6 / "documents"
OUT_DIR.mkdir(parents=True, exist_ok=True)
SAFE_IMG_DIR = FIG_DIR / "docx_safe_images"
SAFE_IMG_DIR.mkdir(parents=True, exist_ok=True)

OUT_PATH = OUT_DIR / "P6_Results_Sectoral_Drivers_Dissertation_Ready.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 89, 89)
INK = RGBColor(22, 22, 22)
LIGHT_BLUE = "E8F1FB"
LIGHT_GRAY = "F2F4F7"
LIGHT_GOLD = "FFF4DE"
LIGHT_GREEN = "EAF5EF"
CONTENT_WIDTH_DXA = 9360


def fmt(x, digits=1):
    if pd.isna(x):
        return ""
    return f"{float(x):,.{digits}f}"


def pct(x):
    return f"{float(x):.1f}%"


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, color in [
        ("Heading 1", 16, DARK_BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11.5, DARK_BLUE),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True

    header = section.header.paragraphs[0]
    header.text = "P6 Results | Sectoral drivers of residual emissions"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=8.8, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "Draft dissertation Results subsection"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer.runs:
        set_run_font(run, size=8.8, color=MUTED)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=100, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for i, width in enumerate(widths):
            tc_pr = row.cells[i]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    for run in p.runs:
        set_run_font(
            run,
            size=16 if level == 1 else 13 if level == 2 else 11.5,
            color=DARK_BLUE if level == 1 else BLUE,
            bold=True,
        )
    return p


def add_para(doc, text, size=10.5, color=INK, bold=False, italic=False, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.rows[0].cells[0]
    shade_cell(cell, fill)
    set_cell_margins(cell, top=110, bottom=110, start=130, end=130)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(label)
    set_run_font(r1, size=10.3, color=DARK_BLUE, bold=True)
    r2 = p.add_run(" " + text)
    set_run_font(r2, size=10.3, color=INK)
    doc.add_paragraph()


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=8.7, color=MUTED, italic=True)


def safe_image_for_docx(path):
    source = Path(path)
    target = SAFE_IMG_DIR / f"{source.stem}_docx_safe.jpg"
    with Image.open(source) as im:
        if im.mode in ("RGBA", "LA") or (im.mode == "P" and "transparency" in im.info):
            rgba = im.convert("RGBA")
            bg = Image.new("RGBA", rgba.size, (255, 255, 255, 255))
            bg.alpha_composite(rgba)
            rgb = bg.convert("RGB")
        else:
            rgb = im.convert("RGB")
        rgb.save(target, "JPEG", quality=96, optimize=True)
    return target


def add_picture(doc, path, width_in, caption):
    path = Path(path)
    if not path.exists():
        add_callout(doc, "Missing figure:", str(path), fill=LIGHT_GOLD)
        return
    image_path = safe_image_for_docx(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(image_path), width=Inches(width_in))
    add_caption(doc, caption)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY, font_size=8.0, last_col_left=True):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr[i], header_fill)
        set_cell_margins(hdr[i], top=85, bottom=85, start=90, end=90)
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(h))
        set_run_font(r, size=font_size, color=DARK_BLUE, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_margins(cells[i], top=75, bottom=75, start=90, end=90)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if last_col_left and i == len(row) - 1 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(val))
            set_run_font(r, size=font_size, color=INK)
    doc.add_paragraph()
    return table


def load_data():
    return {
        "rank": pd.read_csv(TABLE_DIR / "p6_desnz_2050_residual_emissions_ranking.csv"),
        "change": pd.read_csv(TABLE_DIR / "p6_sector_change_2023_to_2050.csv"),
        "align": pd.read_csv(TABLE_DIR / "p6_cautious_desnz_ccc_sector_alignment_2050.csv"),
        "qc": pd.read_csv(TABLE_DIR / "p6_sector_analysis_quality_checks.csv"),
    }


def build_doc():
    data = load_data()
    rank = data["rank"].sort_values("rank_2050_residual")
    change = data["change"].sort_values("rank_2050_residual")
    align = data["align"]

    total_2050 = rank["2050"].sum()
    top3 = rank.head(3)
    top3_sum = top3["2050"].sum()
    top3_share = top3_sum / total_2050 * 100

    buildings = rank.loc[rank["tes_sector"] == "Buildings and product uses"].iloc[0]
    transport = rank.loc[rank["tes_sector"] == "Domestic Transport"].iloc[0]
    industry = rank.loc[rank["tes_sector"] == "Industry"].iloc[0]
    electricity = rank.loc[rank["tes_sector"] == "Electricity supply"].iloc[0]
    agriculture = rank.loc[rank["tes_sector"] == "Agriculture"].iloc[0]

    doc = Document()
    configure_doc(doc)

    add_heading(doc, "P6 Results: Sectoral Drivers of Residual Emissions", 1)
    add_para(
        doc,
        "This draft subsection is written for insertion into the dissertation Results chapter. It builds on the aggregate benchmark comparison in P5 by examining where residual emissions remain concentrated within the DESNZ EEP 2024 current-policy projection.",
        size=9.8,
        color=MUTED,
    )

    add_callout(
        doc,
        "Headline finding:",
        f"Residual emissions in the DESNZ current-policy projection are concentrated in a small group of sectors. The three largest residual sectors in 2050 are buildings and product uses ({fmt(buildings['2050'])} MtCO2e), domestic transport ({fmt(transport['2050'])} MtCO2e) and industry ({fmt(industry['2050'])} MtCO2e). Together, they account for {fmt(top3_sum)} MtCO2e, or {pct(top3_share)} of DESNZ projected 2050 emissions including IAS.",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "Sectoral Scope and Interpretation", 2)
    add_para(
        doc,
        "The analysis uses DESNZ territorial emissions statistics (TES) sectors as the primary classification. This is appropriate because the aim of P6 is to diagnose the composition of the DESNZ current-policy residual baseline, not to produce an exact sector-by-sector decomposition of the DESNZ-CCC benchmark gap. The latter would require a more detailed accounting bridge because DESNZ TES sectors and CCC sectors do not always align one-to-one.",
    )
    add_para(
        doc,
        "The P6 results should therefore be interpreted as a residual-emissions diagnosis. CCC sector values can provide cautious supporting context, but the headline evidence is the DESNZ sectoral projection itself: which sectors remain largest by 2050, how much they change from 2023, and how concentrated the residual baseline becomes.",
    )

    add_heading(doc, "Projected Sectoral Trajectories to 2050", 2)
    add_para(
        doc,
        "Figure 1 shows the DESNZ sectoral trajectories from 2023 to 2050. The trajectories indicate that residual emissions do not follow a uniform pattern across the economy. Domestic transport declines substantially from its 2023 level, while buildings and product uses remains high throughout the projection and is slightly higher in 2050 than in 2023. Electricity supply falls sharply in the late 2020s but rises again after the early 2030s, making it important both as a residual source and as an enabling sector for wider electrification.",
    )
    add_picture(
        doc,
        FIG_DIR / "p6_desnz_sector_projection_2023_2050.png",
        6.2,
        "Figure 1. DESNZ EEP 2024 sectoral emissions projections by TES sector, 2023-2050. Source: DESNZ EEP 2024-2050 Annex A.",
    )

    add_heading(doc, "2050 Residual Emissions Ranking", 2)
    add_para(
        doc,
        f"Table 1 and Figure 2 show the 2050 residual-emissions ranking. Buildings and product uses is the largest residual sector at {fmt(buildings['2050'])} MtCO2e, accounting for {pct(buildings['share_of_2050_inc_IAS_total_pct'])} of the 2050 DESNZ including-IAS total. Domestic transport remains the second-largest sector at {fmt(transport['2050'])} MtCO2e, despite a large projected fall from 2023. Industry is third at {fmt(industry['2050'])} MtCO2e, followed by electricity supply at {fmt(electricity['2050'])} MtCO2e and agriculture at {fmt(agriculture['2050'])} MtCO2e.",
    )

    rank_rows = []
    for _, row in rank.iterrows():
        if row["change_2023_2050_MtCO2e"] > 0:
            interpretation = "Higher than 2023"
        elif row["rank_2050_residual"] <= 5:
            interpretation = "Material residual"
        else:
            interpretation = "Smaller residual"
        rank_rows.append([
            int(row["rank_2050_residual"]),
            row["tes_sector"],
            fmt(row["2050"]),
            pct(row["share_of_2050_inc_IAS_total_pct"]),
            fmt(row["change_2023_2050_MtCO2e"]),
            interpretation,
        ])
    add_table(
        doc,
        ["Rank", "DESNZ TES sector", "2050", "Share", "Change 2023-2050", "Interpretation"],
        rank_rows,
        [650, 2300, 900, 950, 1500, 2300],
        font_size=7.5,
        header_fill=LIGHT_GRAY,
    )
    add_caption(
        doc,
        "Table 1. DESNZ projected residual emissions by TES sector in 2050. Source: author calculation from DESNZ EEP 2024-2050 Annex A.",
    )

    add_picture(
        doc,
        FIG_DIR / "p6_desnz_2050_residual_emissions_ranking.png",
        6.0,
        "Figure 2. DESNZ projected residual emissions by TES sector in 2050. Source: author calculation from DESNZ EEP 2024-2050 Annex A.",
    )

    add_heading(doc, "Concentration of Residual Emissions", 2)
    add_para(
        doc,
        f"The ranking indicates that the DESNZ 2050 residual baseline is concentrated rather than evenly distributed. The top three sectors account for {fmt(top3_sum)} MtCO2e, equivalent to {pct(top3_share)} of the total DESNZ 2050 emissions including IAS. This concentration matters for interpretation because the aggregate P5 benchmark gap is not simply an economy-wide abstraction: it is linked to persistent residual emissions in a limited number of sectors.",
    )
    add_para(
        doc,
        "Buildings and product uses is especially important because it is projected to increase slightly between 2023 and 2050. Domestic transport shows the opposite pattern: it falls sharply over the projection period but remains large because its starting level is high. Industry remains a major residual sector with only a modest projected reduction, while electricity supply remains important because of both its direct emissions and its role in supporting electrification across transport, heat and industry.",
    )

    add_picture(
        doc,
        FIG_DIR / "p6_sector_change_2023_to_2050.png",
        6.0,
        "Figure 3. Projected change in DESNZ sectoral emissions between 2023 and 2050. Source: author calculation from DESNZ EEP 2024-2050 Annex A.",
    )

    add_heading(doc, "Cautious CCC Sector Context", 2)
    add_para(
        doc,
        "The DESNZ residual-emissions ranking can be compared cautiously with CCC sector values, but it should not be presented as an exact decomposition of the DESNZ-CCC gap. Some categories are broad matches, such as agriculture, electricity supply, fuel supply and waste. Others require boundary caution. DESNZ combines buildings and product uses, while CCC separates residential buildings, non-residential buildings and F-gases; DESNZ domestic transport does not map exactly onto CCC surface transport; IAS, LULUCF and engineered removals require explicit accounting treatment.",
    )

    caveat_rows = []
    for sector in ["Buildings and product uses", "Domestic Transport", "Industry", "Electricity supply", "Agriculture", "IAS", "LULUCF", "(no direct DESNZ TES equivalent)"]:
        row = align.loc[align["desnz_tes_sector"] == sector]
        if row.empty:
            continue
        r = row.iloc[0]
        caveat_rows.append([
            r["desnz_tes_sector"],
            r["ccc7_sector_match"],
            r["alignment_type"],
            r["recommended_use"],
        ])
    add_table(
        doc,
        ["DESNZ sector", "CCC comparator", "Alignment", "Recommended use"],
        caveat_rows,
        [1800, 2600, 1550, 3410],
        font_size=7.1,
        header_fill=LIGHT_GOLD,
    )
    add_caption(
        doc,
        "Table 2. Cautious sector-boundary note for interpreting DESNZ residual sectors against CCC sector categories. This table supports the caveat and should not be read as an exact decomposition of the benchmark gap.",
    )

    add_heading(doc, "Interpretation for the Results Chapter", 2)
    add_para(
        doc,
        "Overall, P6 shows that the residual DESNZ current-policy baseline is sectorally concentrated. This strengthens the P5 benchmark result by moving from the size and timing of the aggregate gap to the sectoral composition of the residual baseline. The strongest dissertation claim is therefore not that the 325.4 MtCO2e DESNZ-CCC gap has been exactly decomposed by sector. Rather, the evidence shows that the current-policy residual baseline is dominated by buildings and product uses, domestic transport, industry, electricity supply and agriculture, with the top three sectors alone accounting for more than half of projected 2050 emissions.",
    )

    add_callout(
        doc,
        "Boundary caveat:",
        "This subsection should not claim exact sector-by-sector attribution of the DESNZ-CCC benchmark gap. The robust claim is that the DESNZ current-policy residual baseline is concentrated in a small number of sectors. CCC sector comparisons should remain supporting context unless a full accounting bridge is added.",
        fill=LIGHT_GOLD,
    )

    add_heading(doc, "Source Note", 2)
    add_para(
        doc,
        "Primary data source: DESNZ Energy and emissions projections 2024-2050 Annex A, using territorial emissions statistics sector projections. CCC sector values are used only in the cautious alignment note and are not the basis for the main residual-emissions ranking.",
        size=9.2,
        color=MUTED,
    )

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_doc())
