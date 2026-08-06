from pathlib import Path
import json
import re

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"E:\UCL Final Essay")
P6 = ROOT / "p6_sector_analysis"
TABLE_DIR = P6 / "tables"
FIG_DIR = P6 / "figures" / "docx_safe_images"
DOC_DIR = P6 / "documents"
NB_DIR = P6 / "notebooks"
for d in [TABLE_DIR, FIG_DIR, DOC_DIR, NB_DIR]:
    d.mkdir(parents=True, exist_ok=True)

FINAL_GHG = [
    p for p in (ROOT / "Data_raw").rglob("final-greenhouse-gas-emissions-tables-2023.xlsx")
    if "Final UK greenhouse" in str(p)
][0]
P6_RANKING = TABLE_DIR / "p6_desnz_2050_residual_emissions_ranking.csv"
CCC7 = ROOT / "Data_raw" / "CCC Seventh Carbon Budget  Balanced Pathway data" / "The-Seventh-Carbon-Budget-full-dataset.xlsx"

OUT_SELECTED = TABLE_DIR / "p6_supplement_historical_projected_sector_selected_years.csv"
OUT_TRANSITION = TABLE_DIR / "p6_supplement_historical_projection_transition_metrics.csv"
OUT_BRIDGE = TABLE_DIR / "p6_supplement_desnz_ccc_broad_sector_bridge_2050.csv"
OUT_QC = TABLE_DIR / "p6_supplement_quality_checks.csv"
FIG_TREND = FIG_DIR / "p6_supplement_historical_projected_top_sectors_docx_safe.jpg"
FIG_BRIDGE = FIG_DIR / "p6_supplement_desnz_ccc_broad_sector_bridge_docx_safe.jpg"
DOCX = DOC_DIR / "P6_Supplement_Historical_Trends_and_Sector_Bridge.docx"
NOTEBOOK = NB_DIR / "P6_supplement_historical_sector_bridge_local_reproducible.ipynb"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(31, 31, 31)
MUTED = RGBColor(89, 89, 89)
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"


def fmt(x, digits=1):
    if pd.isna(x):
        return ""
    return f"{float(x):,.{digits}f}"


def read_historical_tes():
    raw = pd.read_excel(FINAL_GHG, sheet_name="1.2", header=None)
    header = raw.iloc[5]
    years = [int(v) for v in header.iloc[3:].dropna().tolist()]
    total_rows = []
    for _, row in raw.iterrows():
        first = row.iloc[0]
        if isinstance(first, str) and (first.lower().endswith(" total") or first == "Grand total"):
            sector = first.replace(" total", "")
            if sector == "Domestic transport":
                sector = "Domestic Transport"
            elif sector == "Grand":
                sector = "Grand total"
            values = row.iloc[3:3 + len(years)].astype(float).tolist()
            for year, value in zip(years, values):
                total_rows.append({"tes_sector": sector, "year": year, "historical_MtCO2e": value})
    hist = pd.DataFrame(total_rows)

    ias_raw = pd.read_excel(FINAL_GHG, sheet_name="5.1", header=None)
    ias_header = ias_raw.iloc[7]
    ias_years = [int(v) for v in ias_header.iloc[2:].dropna().tolist()]
    ias_row = ias_raw[ias_raw.iloc[:, 0].astype(str).eq("International aviation and shipping bunkers total")].iloc[0]
    ias_values = ias_row.iloc[2:2 + len(ias_years)].astype(float).tolist()
    ias = pd.DataFrame(
        [{"tes_sector": "IAS", "year": y, "historical_MtCO2e": v} for y, v in zip(ias_years, ias_values)]
    )
    hist = pd.concat([hist, ias], ignore_index=True)
    hist = hist[hist["tes_sector"] != "Grand total"].copy()
    return hist


def read_projection():
    proj = pd.read_csv(P6_RANKING)
    year_cols = [c for c in ["2023", "2030", "2035", "2050"] if c in proj.columns]
    long = proj.melt(
        id_vars=["tes_sector"],
        value_vars=year_cols,
        var_name="year",
        value_name="projection_MtCO2e",
    )
    long["year"] = long["year"].astype(int)
    long["projection_MtCO2e"] = long["projection_MtCO2e"].astype(float)
    return long


def build_historical_projection_tables():
    hist = read_historical_tes()
    proj = read_projection()
    selected_years = [1990, 2008, 2023, 2030, 2035, 2050]
    sectors = [
        "Buildings and product uses",
        "Domestic Transport",
        "Industry",
        "Electricity supply",
        "Agriculture",
        "IAS",
        "Fuel supply",
        "Waste",
        "LULUCF",
    ]
    rows = []
    for sector in sectors:
        for year in selected_years:
            if year <= 2023:
                value = hist[(hist["tes_sector"] == sector) & (hist["year"] == year)]["historical_MtCO2e"]
                source = "DESNZ final GHG statistics Table 1.2" if sector != "IAS" else "DESNZ final GHG statistics Table 5.1"
            else:
                value = proj[(proj["tes_sector"] == sector) & (proj["year"] == year)]["projection_MtCO2e"]
                source = "DESNZ EEP 2024 Annex A TES projection"
            if not value.empty:
                rows.append({
                    "tes_sector": sector,
                    "year": year,
                    "value_MtCO2e": float(value.iloc[0]),
                    "segment": "historical" if year <= 2023 else "projection",
                    "source": source,
                })
    selected = pd.DataFrame(rows)
    selected.to_csv(OUT_SELECTED, index=False)

    wide = selected.pivot(index="tes_sector", columns="year", values="value_MtCO2e").reset_index()
    metric_rows = []
    for _, row in wide.iterrows():
        sector = row["tes_sector"]
        v1990, v2008, v2023, v2030, v2050 = row.get(1990), row.get(2008), row.get(2023), row.get(2030), row.get(2050)
        hist_abs = v2023 - v1990
        hist_pct = hist_abs / v1990 * 100 if v1990 else None
        proj_abs = v2050 - v2023
        proj_pct = proj_abs / v2023 * 100 if v2023 else None
        if sector == "Buildings and product uses":
            interpretation = "Largest 2050 residual; historical decline does not continue into projection, which rises slightly after 2023."
        elif sector == "Domestic Transport":
            interpretation = "Large historical and projected decline, but 2050 residual remains high because the 2023 base is large."
        elif sector == "Industry":
            interpretation = "Large historical decline, but the projection leaves a persistent 2050 residual."
        elif sector == "Electricity supply":
            interpretation = "Very large historical decline, but residual rises after 2030 and remains strategically important for electrification."
        elif sector == "Agriculture":
            interpretation = "Slow historical and projected decline; persistent non-CO2 residual."
        elif sector == "IAS":
            interpretation = "Memo-item emissions rise historically and decline only modestly in the projection."
        elif sector == "LULUCF":
            interpretation = "Net land-use category is volatile and not directly comparable with engineered removals."
        else:
            interpretation = "Residual is smaller than the leading sectors but still relevant to whole-economy totals."
        metric_rows.append({
            "tes_sector": sector,
            "1990_MtCO2e": v1990,
            "2008_MtCO2e": v2008,
            "2023_MtCO2e": v2023,
            "2030_MtCO2e": v2030,
            "2050_MtCO2e": v2050,
            "historical_change_1990_2023_MtCO2e": hist_abs,
            "historical_change_1990_2023_pct": hist_pct,
            "projected_change_2023_2050_MtCO2e": proj_abs,
            "projected_change_2023_2050_pct": proj_pct,
            "interpretation": interpretation,
        })
    metrics = pd.DataFrame(metric_rows).sort_values("2050_MtCO2e", ascending=False)
    metrics.to_csv(OUT_TRANSITION, index=False)
    return selected, metrics


def build_sector_bridge():
    p6 = pd.read_csv(P6_RANKING)
    desnz_2050 = p6.set_index("tes_sector")["2050"].astype(float).to_dict()
    ccc = pd.read_excel(CCC7, sheet_name="Sector-level data")
    ccc = ccc[
        (ccc["scenario"] == "Balanced Pathway")
        & (ccc["country"] == "United Kingdom")
        & (ccc["variable"] == "Emissions: direct emissions total")
        & (ccc["year"] == 2050)
    ].copy()
    ccc_2050 = ccc.set_index("sector")["value"].astype(float).to_dict()

    bridge_specs = [
        ("Agriculture", "Agriculture", ["Agriculture"], "Direct", "Broadly comparable."),
        ("Buildings and product uses", "Residential + non-residential buildings + F-gases", ["Residential buildings", "Non-residential buildings", "F-gases"], "Aggregate / partial", "DESNZ product uses and CCC F-gases/buildings are only a broad bridge."),
        ("Domestic Transport", "Surface transport", ["Surface transport"], "Partial", "DESNZ domestic transport and CCC surface transport are close but not identical boundaries."),
        ("Electricity supply", "Electricity supply", ["Electricity supply"], "Direct", "Broadly comparable direct sector."),
        ("Fuel supply", "Fuel supply", ["Fuel supply"], "Direct", "Broadly comparable direct sector."),
        ("Industry", "Industry", ["Industry"], "Direct / broad", "Broadly comparable but sector definitions may still differ."),
        ("Waste", "Waste", ["Waste"], "Direct / broad", "Broadly comparable direct sector."),
        ("IAS", "Aviation + Shipping", ["Aviation", "Shipping"], "Aggregate / partial", "IAS bunker definitions and CCC aviation/shipping should be treated cautiously."),
        ("LULUCF", "Land use", ["Land use"], "Partial / not one-to-one", "CCC land use is negative in 2050; removals treatment is a major boundary issue."),
        ("No direct DESNZ TES equivalent", "Engineered removals", ["Engineered removals"], "CCC-only", "CCC engineered removals have no direct DESNZ TES-sector equivalent."),
    ]
    rows = []
    for desnz_sector, ccc_group, ccc_sectors, alignment, note in bridge_specs:
        dval = None if desnz_sector.startswith("No direct") else desnz_2050.get(desnz_sector)
        cval = sum(ccc_2050.get(s, 0.0) for s in ccc_sectors)
        gap = None if dval is None else dval - cval
        rows.append({
            "desnz_tes_sector": desnz_sector,
            "desnz_2050_MtCO2e": dval,
            "ccc7_broad_group": ccc_group,
            "ccc7_2050_MtCO2e": cval,
            "broad_difference_DESNZ_minus_CCC7_MtCO2e": gap,
            "alignment_type": alignment,
            "interpretation_note": note,
        })
    bridge = pd.DataFrame(rows)
    bridge.to_csv(OUT_BRIDGE, index=False)
    return bridge


def load_font(size=26, bold=False):
    for name in ["arialbd.ttf" if bold else "arial.ttf", "calibri.ttf"]:
        try:
            return ImageFont.truetype(name, size)
        except Exception:
            continue
    return ImageFont.load_default()


def make_trend_figure(selected):
    sectors = ["Electricity supply", "Domestic Transport", "Buildings and product uses", "Industry", "Agriculture", "IAS"]
    colors = {
        "Electricity supply": "#2E74B5",
        "Domestic Transport": "#C00000",
        "Buildings and product uses": "#7F3C8D",
        "Industry": "#6B4E16",
        "Agriculture": "#4B7F52",
        "IAS": "#D9822B",
    }
    years = [1990, 2008, 2023, 2030, 2035, 2050]
    width, height = 1800, 1050
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    font_title = load_font(42, True)
    font = load_font(25)
    font_small = load_font(21)
    draw.text((70, 45), "P6 supplement: historical and projected sector trends", fill="#1F4D78", font=font_title)
    draw.text((70, 100), "Historical values use DESNZ final GHG statistics; projected values use DESNZ EEP 2024 TES sectors.", fill="#595959", font=font_small)
    x0, y0, x1, y1 = 150, 180, 1580, 870
    max_y = 220
    for tick in [0, 50, 100, 150, 200]:
        y = y1 - tick / max_y * (y1 - y0)
        draw.line((x0, y, x1, y), fill="#E6E6E6", width=2)
        draw.text((70, y - 12), str(tick), fill="#595959", font=font_small)
    def xp(year):
        return x0 + (year - 1990) / (2050 - 1990) * (x1 - x0)
    def yp(value):
        return y1 - value / max_y * (y1 - y0)
    for year in years:
        x = xp(year)
        draw.line((x, y0, x, y1), fill="#F2F2F2", width=1)
        draw.text((x - 32, y1 + 25), str(year), fill="#595959", font=font_small)
    draw.line((xp(2023), y0, xp(2023), y1), fill="#333333", width=3)
    draw.text((xp(2023) + 10, y0 + 10), "projection starts", fill="#333333", font=font_small)
    for sector in sectors:
        data = selected[selected["tes_sector"] == sector].sort_values("year")
        pts = [(xp(int(r.year)), yp(float(r.value_MtCO2e))) for r in data.itertuples()]
        for a, b in zip(pts, pts[1:]):
            draw.line((a[0], a[1], b[0], b[1]), fill=colors[sector], width=5)
        for x, y in pts:
            draw.ellipse((x - 6, y - 6, x + 6, y + 6), fill=colors[sector])
    legend_x, legend_y = 1160, 190
    for i, sector in enumerate(sectors):
        y = legend_y + i * 45
        draw.rectangle((legend_x, y, legend_x + 30, y + 18), fill=colors[sector])
        draw.text((legend_x + 42, y - 5), sector, fill="#1F1F1F", font=font_small)
    draw.text((70, 940), "Note: values are MtCO2e. This figure is a selected-year bridge, not a complete historical decomposition.", fill="#595959", font=font_small)
    img.save(FIG_TREND, "JPEG", quality=95)


def make_bridge_figure(bridge):
    plot = bridge[
        bridge["desnz_tes_sector"].isin([
            "Buildings and product uses", "Domestic Transport", "Industry", "Electricity supply",
            "Agriculture", "IAS", "Fuel supply", "Waste"
        ])
    ].copy()
    plot = plot.sort_values("desnz_2050_MtCO2e", ascending=True)
    width, height = 1800, 1100
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    font_title = load_font(42, True)
    font = load_font(23)
    font_small = load_font(20)
    draw.text((70, 45), "Cautious broad sector bridge: DESNZ residuals vs CCC7 2050", fill="#1F4D78", font=font_title)
    draw.text((70, 100), "Broad bridge only. Sector boundaries are not exact, especially buildings/product uses, IAS, LULUCF and removals.", fill="#595959", font=font_small)
    left, right = 520, 1630
    top = 190
    row_h = 92
    max_val = max(90, plot["desnz_2050_MtCO2e"].max())
    def xw(v):
        return left + v / max_val * (right - left)
    for tick in [0, 20, 40, 60, 80]:
        x = xw(tick)
        draw.line((x, top - 30, x, top + row_h * len(plot)), fill="#E6E6E6", width=2)
        draw.text((x - 15, top + row_h * len(plot) + 20), str(tick), fill="#595959", font=font_small)
    for i, r in enumerate(plot.itertuples()):
        y = top + i * row_h
        draw.text((70, y + 8), r.desnz_tes_sector, fill="#1F1F1F", font=font)
        dval = float(r.desnz_2050_MtCO2e)
        cval = float(r.ccc7_2050_MtCO2e)
        draw.rounded_rectangle((left, y + 8, xw(dval), y + 35), radius=6, fill="#8B1E3F")
        draw.rounded_rectangle((left, y + 43, xw(max(cval, 0)), y + 70), radius=6, fill="#2E74B5")
        draw.text((xw(dval) + 8, y + 5), fmt(dval), fill="#1F1F1F", font=font_small)
        draw.text((xw(max(cval, 0)) + 8, y + 40), fmt(cval), fill="#1F1F1F", font=font_small)
    draw.rectangle((70, 930, 100, 950), fill="#8B1E3F")
    draw.text((115, 925), "DESNZ EEP 2024 current-policy residual", fill="#1F1F1F", font=font_small)
    draw.rectangle((70, 970, 100, 990), fill="#2E74B5")
    draw.text((115, 965), "CCC7 Balanced Pathway broad group", fill="#1F1F1F", font=font_small)
    img.save(FIG_BRIDGE, "JPEG", quality=95)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_table(table):
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, HEADER_FILL)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = DARK_BLUE if r_idx == 0 else INK
                    if r_idx == 0:
                        run.bold = True


def add_table(doc, rows):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    for i, header in enumerate(rows[0]):
        table.rows[0].cells[i].text = str(header)
    for row_data in rows[1:]:
        row = table.add_row()
        for cell, value in zip(row.cells, row_data):
            cell.text = str(value)
    style_table(table)
    return table


def add_para(doc, text, size=10.5, color=INK):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_image(doc, path, caption, width=6.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED


def build_docx(metrics, bridge):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)
    header = sec.header.paragraphs[0]
    header.text = "P6 Supplement | Historical trends and sector bridge"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED
    for style_name, size, color in [("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 11.5, DARK_BLUE)]:
        st = doc.styles[style_name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True

    title = doc.add_paragraph()
    r = title.add_run("P6 Supplement: Historical Trends And Sector Bridge")
    r.font.name = "Calibri"
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = DARK_BLUE
    add_para(doc, "Purpose: to patch the two weakest P6 areas before the supervision meeting: historical-sector context and a cautious DESNZ-CCC sector bridge.", color=MUTED)

    def mval(sector, col):
        return float(metrics.loc[metrics["tes_sector"].eq(sector), col].iloc[0])

    def bval(sector, col):
        value = bridge.loc[bridge["desnz_tes_sector"].eq(sector), col].iloc[0]
        return None if pd.isna(value) else float(value)

    total_2050 = float(metrics["2050_MtCO2e"].sum())
    top3 = ["Buildings and product uses", "Domestic Transport", "Industry"]
    top3_total = sum(mval(sector, "2050_MtCO2e") for sector in top3)
    top3_share = top3_total / total_2050 * 100
    buildings_gap = bval("Buildings and product uses", "broad_difference_DESNZ_minus_CCC7_MtCO2e")
    transport_gap = bval("Domestic Transport", "broad_difference_DESNZ_minus_CCC7_MtCO2e")
    industry_gap = bval("Industry", "broad_difference_DESNZ_minus_CCC7_MtCO2e")
    electricity_gap = bval("Electricity supply", "broad_difference_DESNZ_minus_CCC7_MtCO2e")

    doc.add_heading("Executive Finding", level=1)
    add_para(
        doc,
        "This supplement makes P6 more defensible by moving it from a simple 2050 sector ranking to a historically "
        "anchored residual-emissions diagnosis. The DESNZ projection still leaves "
        f"{fmt(total_2050)} MtCO2e in 2050 across the TES residual sectors covered here. The three largest residual "
        f"sectors - buildings/product uses ({fmt(mval('Buildings and product uses', '2050_MtCO2e'))} MtCO2e), "
        f"domestic transport ({fmt(mval('Domestic Transport', '2050_MtCO2e'))} MtCO2e), and industry "
        f"({fmt(mval('Industry', '2050_MtCO2e'))} MtCO2e) - account for {fmt(top3_total)} MtCO2e, or "
        f"{fmt(top3_share)}% of the 2050 residual total.",
    )
    add_para(
        doc,
        "The main analytical implication is that the remaining gap is not concentrated in one late-stage modelling "
        "artefact. It is distributed across demand-side heat and product-use emissions, transport decarbonisation, "
        "industrial residuals, and power-sector residual emissions. This supports a sectoral Results subsection that "
        "explains where current-policy emissions persist, while still being careful not to claim a precise causal "
        "decomposition of the DESNZ-CCC national gap.",
    )
    add_table(doc, [
        ["P6 strengthened point", "Evidence now added", "How it should be framed"],
        [
            "Historical grounding",
            "Final UK GHG statistics for 1990, 2008 and 2023 are linked to DESNZ EEP projection points for 2030, 2035 and 2050.",
            "P6 is now a trajectory-based diagnosis, not only a 2050 snapshot.",
        ],
        [
            "Residual concentration",
            f"Top three residual sectors account for {fmt(top3_share)}% of DESNZ 2050 TES residual emissions.",
            "Use this to justify focusing the Results discussion on buildings, transport and industry.",
        ],
        [
            "CCC benchmark bridge",
            f"Broad DESNZ-minus-CCC7 differences are largest for buildings/product uses ({fmt(buildings_gap)} MtCO2e), domestic transport ({fmt(transport_gap)} MtCO2e), industry ({fmt(industry_gap)} MtCO2e) and electricity supply ({fmt(electricity_gap)} MtCO2e).",
            "Use as a cautious bridge, not as an exact sector-by-sector gap decomposition.",
        ],
    ])

    doc.add_heading("1. What This Adds", level=1)
    add_table(doc, [
        ["P6 gap", "Supplement added", "How to use it"],
        ["Historical vs projected sector trends", "DESNZ final GHG statistics 1990-2023 are joined to DESNZ EEP 2024 sector projections for selected years.", "Use as evidence that the P6 ranking is not only a 2050 snapshot."],
        ["CCC sector alignment", "A cautious broad-sector bridge compares DESNZ 2050 residual sectors with CCC7 broad sector groups.", "Use only as an interpretive bridge, not as an exact sectoral gap decomposition."],
        ["Neil-facing caveat", "Boundary limits are explicit for buildings/product uses, IAS, LULUCF and engineered removals.", "Use this caveat proactively if Neil asks about sector comparability."],
    ])

    doc.add_heading("2. Historical-To-Projection Sector Trends", level=1)
    add_para(doc, "The historical data confirm that many sectors have already declined substantially since 1990. However, the DESNZ projection leaves large residual emissions in 2050, especially in buildings/product uses, domestic transport, industry and electricity supply. This strengthens P6 because it shows not only where emissions remain in 2050, but how the projected residual follows from historical and projected trajectories.")
    add_para(
        doc,
        f"Buildings/product uses fall from {fmt(mval('Buildings and product uses', '1990_MtCO2e'))} MtCO2e in 1990 to "
        f"{fmt(mval('Buildings and product uses', '2023_MtCO2e'))} MtCO2e in 2023, but then remain around "
        f"{fmt(mval('Buildings and product uses', '2050_MtCO2e'))} MtCO2e in 2050. Domestic transport declines more "
        f"clearly after 2023, from {fmt(mval('Domestic Transport', '2023_MtCO2e'))} to "
        f"{fmt(mval('Domestic Transport', '2050_MtCO2e'))} MtCO2e, but it remains the second-largest residual category. "
        f"Industry has already fallen sharply from {fmt(mval('Industry', '1990_MtCO2e'))} to "
        f"{fmt(mval('Industry', '2023_MtCO2e'))} MtCO2e, yet the projection shows a persistent 2050 residual of "
        f"{fmt(mval('Industry', '2050_MtCO2e'))} MtCO2e. These patterns give P6 a stronger explanatory basis than the "
        "previous version, which mainly ranked sectors at the endpoint.",
    )
    add_image(doc, FIG_TREND, "Figure S1. Selected historical and projected DESNZ TES sector trends.")
    table_rows = [["Sector", "1990", "2008", "2023", "2030", "2050", "Interpretation"]]
    for r in metrics.head(7).itertuples():
        table_rows.append([
            r.tes_sector,
            fmt(r._2),
            fmt(r._3),
            fmt(r._4),
            fmt(r._5),
            fmt(r._6),
            r.interpretation,
        ])
    add_table(doc, table_rows)

    doc.add_heading("3. Cautious DESNZ-CCC Broad Sector Bridge", level=1)
    add_para(doc, "The bridge below answers the likely supervisor question: can DESNZ sectors be aligned with CCC sectors? The answer is partly yes, but not exactly. Agriculture, electricity supply, fuel supply, industry and waste are broadly comparable. Buildings/product uses, domestic transport, IAS, LULUCF and engineered removals require stronger caveats.")
    add_para(
        doc,
        f"On this broad bridge, DESNZ 2050 residuals sit far above CCC7 values in buildings/product uses "
        f"({fmt(buildings_gap)} MtCO2e), domestic transport ({fmt(transport_gap)} MtCO2e), industry "
        f"({fmt(industry_gap)} MtCO2e) and electricity supply ({fmt(electricity_gap)} MtCO2e). The bridge is useful "
        "because it shows that the same sectors identified by DESNZ residual ranking are also materially distant from "
        "CCC7-consistent endpoint values. However, it must be presented as an indicative alignment because sector "
        "boundaries differ, especially for product uses, aviation and shipping, LULUCF, and engineered removals.",
    )
    add_image(doc, FIG_BRIDGE, "Figure S2. Cautious broad-sector bridge between DESNZ residuals and CCC7 2050 values.")
    table_rows = [["DESNZ TES sector", "DESNZ 2050", "CCC7 broad group", "CCC7 2050", "Alignment", "Note"]]
    for r in bridge.itertuples():
        table_rows.append([
            r.desnz_tes_sector,
            "" if pd.isna(r.desnz_2050_MtCO2e) else fmt(r.desnz_2050_MtCO2e),
            r.ccc7_broad_group,
            fmt(r.ccc7_2050_MtCO2e),
            r.alignment_type,
            r.interpretation_note,
        ])
    add_table(doc, table_rows)

    doc.add_heading("4. Meeting-Ready Interpretation", level=1)
    add_para(doc, "The patched P6 interpretation should be: P6 now provides a DESNZ residual-emissions diagnosis supported by historical-sector context and a cautious CCC bridge. It does not claim exact DESNZ-CCC sectoral decomposition. If Neil wants stricter alignment, the natural next step is a short appendix rather than a major expansion of the main analysis.")
    add_para(
        doc,
        "Suggested spoken framing: 'I realised that my first P6 version was too close to an endpoint ranking, so I have "
        "now added two checks. First, I link the 2050 residual sectors back to the 1990-2023 historical record and the "
        "2030/2035 projection path. Second, I add a careful DESNZ-CCC sector bridge. I am not treating that as an exact "
        "like-for-like decomposition, but it gives a stronger evidence base for explaining which sectors drive the "
        "remaining current-policy residual and where the benchmark-consistent pathway is most demanding.'",
    )
    add_table(doc, [
        ["Likely question", "Recommended answer"],
        ["Does P6 compare historical and projected sectoral trends?", "Yes. The supplement joins DESNZ final GHG statistics to DESNZ EEP sector projections for selected years."],
        ["Can CCC sectors be aligned with DESNZ sectors?", "Partly. Several broad groups align reasonably well, but buildings/product uses, IAS, LULUCF and engineered removals require explicit caveats."],
        ["Does this now explain sectoral drivers of the projected gap?", "It supports sectoral interpretation, but it should still be framed as residual-emissions diagnosis rather than exact causal attribution."],
    ])

    doc.save(DOCX)


def build_notebook():
    cells = [
        {
            "cell_type": "markdown",
            "metadata": {},
            "source": "# P6 supplement: historical sector trends and broad CCC bridge\n\nRun this notebook from the dissertation project folder. It reproduces the P6 supplement tables and JPEG figures."
        },
        {
            "cell_type": "code",
            "execution_count": None,
            "metadata": {},
            "outputs": [],
            "source": "from pathlib import Path\nimport pandas as pd\nfrom PIL import Image, ImageDraw, ImageFont\n\nPROJECT_ROOT = Path.cwd()\nwhile PROJECT_ROOT != PROJECT_ROOT.parent and not (PROJECT_ROOT / 'Data_raw').exists():\n    PROJECT_ROOT = PROJECT_ROOT.parent\nassert (PROJECT_ROOT / 'Data_raw').exists(), 'Run from inside the dissertation project folder.'\nprint(PROJECT_ROOT)"
        },
        {
            "cell_type": "code",
            "execution_count": None,
            "metadata": {},
            "outputs": [],
            "source": "FINAL_GHG = [p for p in (PROJECT_ROOT / 'Data_raw').rglob('final-greenhouse-gas-emissions-tables-2023.xlsx') if 'Final UK greenhouse' in str(p)][0]\nP6_RANKING = PROJECT_ROOT / 'p6_sector_analysis' / 'tables' / 'p6_desnz_2050_residual_emissions_ranking.csv'\nCCC7 = PROJECT_ROOT / 'Data_raw' / 'CCC Seventh Carbon Budget  Balanced Pathway data' / 'The-Seventh-Carbon-Budget-full-dataset.xlsx'\nprint(FINAL_GHG)\nprint(P6_RANKING)\nprint(CCC7)"
        },
        {
            "cell_type": "code",
            "execution_count": None,
            "metadata": {},
            "outputs": [],
            "source": "# Use the generated script for full reproduction.\nimport runpy\nscript = PROJECT_ROOT / 'p6_sector_analysis' / 'build_p6_supplement_historical_bridge.py'\nrunpy.run_path(str(script), run_name='__main__')"
        },
    ]
    nb = {
        "cells": cells,
        "metadata": {
            "kernelspec": {"display_name": "Python 3", "language": "python", "name": "python3"},
            "language_info": {"name": "python", "version": "3.12"},
        },
        "nbformat": 4,
        "nbformat_minor": 5,
    }
    NOTEBOOK.write_text(json.dumps(nb, indent=2), encoding="utf-8")


def inspect_outputs():
    qc_rows = []
    for path in [OUT_SELECTED, OUT_TRANSITION, OUT_BRIDGE, FIG_TREND, FIG_BRIDGE, DOCX, NOTEBOOK]:
        qc_rows.append({"check": f"{path.name} exists", "status": "PASS" if path.exists() and path.stat().st_size > 0 else "FAIL"})
    doc = Document(str(DOCX))
    text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            text += "\n" + " | ".join(c.text for c in row.cells)
    for phrase in ["Historical-To-Projection", "Cautious DESNZ-CCC", "not as an exact sectoral gap decomposition", "Buildings and product uses"]:
        qc_rows.append({"check": f"doc contains {phrase}", "status": "PASS" if phrase in text else "FAIL"})
    pd.DataFrame(qc_rows).to_csv(OUT_QC, index=False)
    return qc_rows


def main():
    selected, metrics = build_historical_projection_tables()
    bridge = build_sector_bridge()
    make_trend_figure(selected)
    make_bridge_figure(bridge)
    build_docx(metrics, bridge)
    build_notebook()
    qc = inspect_outputs()
    print(pd.DataFrame(qc).to_string(index=False))


if __name__ == "__main__":
    main()
