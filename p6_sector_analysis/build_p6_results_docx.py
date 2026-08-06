from pathlib import Path
from zipfile import ZipFile

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"E:\UCL Final Essay")
P6_ROOT = ROOT / "p6_sector_analysis"
TABLES_DIR = P6_ROOT / "tables"
FIGURES_DIR = P6_ROOT / "figures"
DOCS_DIR = P6_ROOT / "documents"
DOCS_DIR.mkdir(parents=True, exist_ok=True)

OUT_DOCX = DOCS_DIR / "P6_Results_Sectoral_Drivers.docx"

CONTENT_WIDTH_DXA = 9360
TABLE_IND_DXA = 120
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 89, 89)
HEADER_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
CAUTION_FILL = "FFF4E5"


def set_cell_text(cell, text, bold=False, color=None, size=8.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(TABLE_IND_DXA))

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)

    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(w))
            set_cell_margins(cell)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_table(doc, headers, rows, widths, font_size=8.3):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.allow_autofit = False
    set_table_geometry(table, widths)
    mark_header_row(table.rows[0])

    for i, h in enumerate(headers):
        shade_cell(table.rows[0].cells[i], HEADER_FILL)
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=font_size, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if i == 0 or isinstance(value, (int, float)) else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], value, size=font_size, align=align)
    set_table_geometry(table, widths)
    doc.add_paragraph()
    return table


def add_callout(doc, label, text, fill=CALLOUT_FILL):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.allow_autofit = False
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    mark_header_row(table.rows[0])
    cell = table.rows[0].cells[0]
    shade_cell(cell, fill)
    set_cell_margins(cell, top=110, bottom=110, start=120, end=120)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.name = "Calibri"
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = DARK_BLUE
    r2 = p.add_run(" " + text)
    r2.font.name = "Calibri"
    r2.font.size = Pt(10.5)
    doc.add_paragraph()
    return table


def add_para(doc, text, style=None, after=6, size=11, color=None, bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = BLUE if level <= 2 else DARK_BLUE
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(13)
        else:
            run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = MUTED
    return p


def add_picture(doc, path, width_in, caption, title, descr):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width_in))
    doc_pr = shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", descr)
    add_caption(doc, caption)
    return shape


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ["top_margin", "bottom_margin", "left_margin", "right_margin"]:
        setattr(section, attr, Inches(1.0))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True

    header_p = section.header.paragraphs[0]
    header_p.text = "P6 Results - Sectoral Drivers"
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header_p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED

    footer_p = section.footer.paragraphs[0]
    footer_p.text = "Draft dissertation results subsection"
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer_p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED


def fmt(value, decimals=1):
    if pd.isna(value):
        return ""
    return f"{float(value):.{decimals}f}"


def build_doc():
    ranking = pd.read_csv(TABLES_DIR / "p6_desnz_2050_residual_emissions_ranking.csv")
    changes = pd.read_csv(TABLES_DIR / "p6_sector_change_2023_to_2050.csv")
    align = pd.read_csv(TABLES_DIR / "p6_cautious_desnz_ccc_sector_alignment_2050.csv")
    qc = pd.read_csv(TABLES_DIR / "p6_sector_analysis_quality_checks.csv")

    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("P6 Results: Sectoral Drivers Of The UK Emissions Gap")
    r.font.name = "Calibri"
    r.font.size = Pt(22)
    r.bold = True
    r.font.color.rgb = INK

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("DESNZ EEP Reference sectoral projections, residual emissions ranking and cautious CCC alignment")
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.color.rgb = MUTED

    top = ranking.iloc[0]
    top3 = ranking.head(3)
    top3_sum = top3["2050"].sum()
    total_2050 = ranking["2050"].sum()
    top3_share = top3_sum / total_2050 * 100

    add_callout(
        doc,
        "Headline result:",
        (
            f"Under the DESNZ EEP Reference projection, 2050 residual emissions are concentrated in a small number of sectors. "
            f"The largest residual sector is {top['tes_sector']} ({fmt(top['2050'])} MtCO2e), while the top three sectors together "
            f"account for {fmt(top3_sum)} MtCO2e, or {fmt(top3_share)}% of DESNZ projected 2050 emissions including IAS."
        ),
    )

    add_heading(doc, "1. Data Basis And Scope", 1)
    add_para(
        doc,
        (
            "This section uses the DESNZ Energy and emissions projections 2024-2050 Annex A TES-sector dataset. "
            "The Reference scenario is used as the current-policy baseline, matching the P4 and P5 treatment of DESNZ EEP as the central projection. "
            "The analysis covers all broad Territorial Emissions Statistics sectors: buildings and product uses, domestic transport, industry, "
            "electricity supply, agriculture, international aviation and shipping (IAS), fuel supply, waste, and LULUCF."
        ),
    )
    add_para(
        doc,
        (
            "The main P6 question is not whether the whole-economy gap exists - that was established in P4 and P5 - but which sectors explain "
            "the persistence of residual emissions by 2050. The CCC sector comparison is used only as a cautious interpretive bridge because "
            "DESNZ TES categories and CCC7 sector categories are not always one-to-one."
        ),
    )

    add_heading(doc, "2. Residual Emissions Ranking In 2050", 1)
    add_para(
        doc,
        (
            "The 2050 ranking indicates that residual emissions are unevenly distributed. Buildings and product uses are the largest residual source, "
            "followed by domestic transport and industry. Electricity supply and agriculture are also material, while IAS remains important because "
            "it is included in the whole-economy accounting basis used for the DESNZ-CCC benchmark comparison."
        ),
    )

    ranking_rows = []
    for _, row in ranking.iterrows():
        if row["change_2023_2050_MtCO2e"] > 0:
            interpretation = "Higher than 2023"
        elif row["rank_2050_residual"] <= 5:
            interpretation = "Material residual"
        else:
            interpretation = "Smaller residual"
        ranking_rows.append(
            [
                int(row["rank_2050_residual"]),
                row["tes_sector"],
                fmt(row["2050"]),
                fmt(row["share_of_2050_inc_IAS_total_pct"]),
                fmt(row["change_2023_2050_MtCO2e"]),
                interpretation,
            ]
        )
    add_table(
        doc,
        ["Rank", "TES sector", "2050", "Share (%)", "Change", "Use"],
        ranking_rows,
        [650, 2450, 1150, 1200, 1200, 2710],
        font_size=7.0,
    )

    add_picture(
        doc,
        FIGURES_DIR / "p6_desnz_2050_residual_emissions_ranking.png",
        6.2,
        "Figure 1. DESNZ projected residual emissions by TES sector in 2050 under the EEP Reference scenario.",
        "P6 2050 residual emissions ranking",
        "Horizontal bar chart ranking DESNZ TES sectors by projected residual emissions in 2050.",
    )

    add_heading(doc, "3. Sectoral Pathway Patterns", 1)
    add_para(
        doc,
        (
            "The time-pathway results show that high residual emissions in 2050 can arise through different sectoral patterns. "
            "Domestic transport shows a large projected reduction from 2023 to 2050 but remains the second-largest residual sector. "
            "Buildings and product uses are more concerning in this projection because the sector is slightly higher in 2050 than in 2023. "
            "Electricity supply falls sharply in the near term but rises again later in the projection, which matters both directly and because "
            "electricity is an enabling sector for heating, transport and parts of industry."
        ),
    )
    add_picture(
        doc,
        FIGURES_DIR / "p6_desnz_sector_projection_2023_2050.png",
        6.35,
        "Figure 2. DESNZ EEP Reference sectoral emissions projections, 2023-2050.",
        "P6 DESNZ sectoral projection pathways",
        "Line chart showing DESNZ EEP Reference emissions pathways for broad TES sectors from 2023 to 2050.",
    )
    add_picture(
        doc,
        FIGURES_DIR / "p6_sector_change_2023_to_2050.png",
        6.15,
        "Figure 3. Projected change in DESNZ sectoral emissions between 2023 and 2050.",
        "P6 sectoral emissions change 2023 to 2050",
        "Horizontal bar chart showing changes in sectoral emissions from 2023 to 2050.",
    )

    add_heading(doc, "4. Interpretation Of Main Sectoral Drivers", 1)
    sector_text = [
        [
            "Buildings and product uses",
            "84.6 MtCO2e in 2050; +6.1 MtCO2e from 2023.",
            "Largest residual sector. This category needs careful explanation because DESNZ combines buildings and product-use emissions, while CCC separates residential buildings, non-residential buildings and F-gases.",
        ],
        [
            "Domestic transport",
            "49.9 MtCO2e in 2050; -61.9 MtCO2e from 2023.",
            "Shows large projected improvement but remains the second-largest residual source. This is a central P6 finding because transport is both a major historical sector and a persistent residual sector.",
        ],
        [
            "Industry",
            "47.4 MtCO2e in 2050; -5.7 MtCO2e from 2023.",
            "Only limited further reduction is projected from the latest historical year to 2050, so industry remains a major residual source despite being below its 2023 level.",
        ],
        [
            "Electricity supply",
            "41.0 MtCO2e in 2050; -3.0 MtCO2e from 2023.",
            "The pathway declines early and then rises later, so the sector should be interpreted both as a direct residual source and as a system condition for electrification.",
        ],
        [
            "Agriculture and IAS",
            "Agriculture 38.2 MtCO2e; IAS 34.0 MtCO2e in 2050.",
            "Both remain important residual sectors. Agriculture is a direct broad-sector comparison with CCC, while IAS needs explicit accounting-basis language.",
        ],
    ]
    add_table(
        doc,
        ["Sector", "Evidence from DESNZ Reference", "Interpretation for P6 Results"],
        sector_text,
        [1900, 2300, 5160],
        font_size=8.0,
    )

    add_heading(doc, "5. Cautious CCC Sector Alignment", 1)
    add_callout(
        doc,
        "Boundary caution:",
        (
            "The DESNZ-CCC sector comparison should not be presented as a strict sector-by-sector gap decomposition. "
            "It is useful for interpretation, but some categories are only partial matches and CCC engineered removals have no direct DESNZ TES equivalent."
        ),
        fill=CAUTION_FILL,
    )
    add_para(
        doc,
        (
            "A cautious broad-sector comparison suggests that the largest apparent DESNZ-CCC differences are associated with buildings and product uses, "
            "domestic transport, industry and electricity supply. However, these values should support the narrative rather than become the headline P6 metric, "
            "because sector boundaries differ between the DESNZ TES projection and the CCC7 pathway dataset."
        ),
    )

    align_selected = align[
        align["desnz_tes_sector"].isin(
            [
                "Buildings and product uses",
                "Domestic Transport",
                "Industry",
                "Electricity supply",
                "Agriculture",
                "IAS",
                "(no direct DESNZ TES equivalent)",
            ]
        )
    ].copy()
    align_rows = []
    for _, row in align_selected.iterrows():
        desnz = "" if pd.isna(row["desnz_2050_MtCO2e"]) else fmt(row["desnz_2050_MtCO2e"])
        ccc = "" if pd.isna(row["ccc7_2050_MtCO2e"]) else fmt(row["ccc7_2050_MtCO2e"])
        gap = "" if pd.isna(row["gap_DESNZ_minus_CCC_2050_MtCO2e"]) else fmt(row["gap_DESNZ_minus_CCC_2050_MtCO2e"])
        align_rows.append([row["desnz_tes_sector"], row["alignment_type"], desnz, ccc, gap])
    add_table(
        doc,
        ["DESNZ TES sector", "Alignment type", "DESNZ 2050", "CCC7 2050", "Broad gap"],
        align_rows,
        [2600, 2100, 1450, 1450, 1760],
        font_size=7.8,
    )
    add_picture(
        doc,
        FIGURES_DIR / "p6_cautious_desnz_ccc_sector_gap_2050.png",
        6.2,
        "Figure 4. Cautious broad-sector comparison between DESNZ 2050 residual emissions and CCC7 2050 sector values.",
        "P6 cautious DESNZ CCC broad sector gap",
        "Horizontal bar chart showing cautious broad-sector differences between DESNZ 2050 and CCC7 2050.",
    )

    add_heading(doc, "6. Draft Results Text For Dissertation", 1)
    add_para(
        doc,
        (
            "The sectoral results show that the projected 2050 emissions gap is not evenly distributed across the economy. "
            "Under the DESNZ EEP Reference projection, buildings and product uses are the largest residual source in 2050, at 84.6 MtCO2e, "
            "representing 26.1% of projected emissions including IAS. Domestic transport remains the second-largest residual source at 49.9 MtCO2e, "
            "despite a projected reduction of 61.9 MtCO2e between 2023 and 2050. Industry, electricity supply and agriculture also remain material, "
            "with projected 2050 emissions of 47.4, 41.0 and 38.2 MtCO2e respectively."
        ),
    )
    add_para(
        doc,
        (
            "These results refine the whole-economy DESNZ-CCC gap identified in P5 by showing which sectors account for persistent residual emissions. "
            "The largest sectoral residual is not necessarily the sector with the weakest projected improvement: domestic transport falls substantially, "
            "but remains large because of its high starting point, while buildings and product uses increase slightly from the latest historical year to 2050. "
            "This distinction is important for the dissertation's interpretation of policy delivery, because it separates the scale of residual emissions "
            "from the direction of projected change."
        ),
    )
    add_para(
        doc,
        (
            "The sectoral comparison with CCC7 should be interpreted cautiously. Some sectors, such as agriculture and electricity supply, can be compared at a broad level, "
            "but others do not align exactly. DESNZ combines buildings and product uses, whereas CCC separates residential buildings, non-residential buildings and F-gases. "
            "CCC also includes engineered removals as a separate benchmark sector, which has no direct DESNZ TES equivalent. For this reason, the P6 headline should remain "
            "the DESNZ residual-emissions ranking and pathway interpretation, with DESNZ-CCC sector alignment used only as supporting evidence."
        ),
    )

    add_heading(doc, "7. Quality Checks And Remaining Cautions", 1)
    qc_rows = qc.values.tolist()
    add_table(doc, ["Check", "Status", "Details"], qc_rows, [3300, 1000, 5060], font_size=7.7)
    add_para(
        doc,
        (
            "The extraction checks passed: all expected broad TES sectors are present, the 2023-2050 projection years are complete, and the sum of 2050 sector values "
            "matches the DESNZ total including IAS. The main remaining caution is interpretive rather than computational: sector boundaries must be kept explicit, "
            "especially for buildings/product uses, IAS, LULUCF and CCC engineered removals."
        ),
    )

    add_heading(doc, "Source Note", 2)
    add_para(
        doc,
        (
            "Primary data source: DESNZ Energy and emissions projections 2024-2050, Annex A: greenhouse gas emissions by Territorial Emissions Statistics sector, "
            "Reference scenario. Supporting benchmark context: CCC Seventh Carbon Budget dataset and P5 DESNZ-CCC sector alignment table."
        ),
        size=9.5,
        color=MUTED,
    )

    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    path = build_doc()
    with ZipFile(path) as z:
        bad = z.testzip()
    print({"docx": str(path), "zip_bad": bad})
